from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
LOGO_PATH = Path(
    r"C:\Users\digui\Documents\localrise-brain\Claude\Claude Local RIse\assents\logo Local Rise Principal.png"
)

RED = RGBColor(0xE3, 0x1B, 0x23)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)
GRAY = RGBColor(0x5F, 0x63, 0x68)
LIGHT = RGBColor(0xF6, 0xF7, 0xF8)
DARK_RED = RGBColor(0xA0, 0x10, 0x15)

COMPANY = "Capitol Consultoria"
DATE = "09 de abril de 2026"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:%s" % edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:%s" % key), str(edge_data[key]))


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def apply_doc_style(document: Document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(8)


def add_brand_header(document: Document, title: str, subtitle: str | None = None):
    if LOGO_PATH.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(5.4))

    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run("LOCALRISE ADVISORY")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(13)
    run.font.color.rgb = RED

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_border(cell, bottom={"val": "single", "sz": 8, "space": 0, "color": "D9DDE3"})
    cell.text = ""

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title_run.font.size = Pt(18)

    if subtitle:
        subtitle_p = document.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_p.add_run(subtitle)
        subtitle_run.font.name = "Calibri"
        subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        subtitle_run.font.size = Pt(10.5)
        subtitle_run.font.color.rgb = GRAY


def add_section_title(document: Document, text: str):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(12)
    run.font.color.rgb = RED


def add_subsection_title(document: Document, text: str):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(11.5)
    run.font.color.rgb = BLACK


def add_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(11)


def add_paragraph_bold_start(document: Document, bold_part: str, rest: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r1 = p.add_run(bold_part)
    r1.bold = True
    r1.font.name = "Calibri"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r1.font.size = Pt(11)
    r2 = p.add_run(rest)
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r2.font.size = Pt(11)


def add_bullets(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(item)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(11)


def add_numbered(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(item)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(11)


def add_highlight_box(document: Document, text: str, bg_hex: str = "F4F4F4"):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    set_cell_border(
        cell,
        left={"val": "single", "sz": 16, "space": 0, "color": "E31B23"},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(11)
    run.italic = True
    document.add_paragraph()


def add_automation_card(document: Document, number: str, name: str, details: list[tuple[str, str]]):
    """Add a structured card for each automation"""
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(f"AUTOMACAO {number} — {name.upper()}")
    r.bold = True
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r.font.size = Pt(11.5)
    r.font.color.rgb = RED

    table = document.add_table(rows=len(details), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Cm(4.2), Cm(12.0)]
    for i, (label, value) in enumerate(details):
        row = table.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        set_cell_bg(row.cells[0], "F4F4F4")
        set_cell_border(row.cells[0],
                        top={"val": "single", "sz": 4, "space": 0, "color": "D9DDE3"},
                        bottom={"val": "single", "sz": 4, "space": 0, "color": "D9DDE3"})
        set_cell_border(row.cells[1],
                        top={"val": "single", "sz": 4, "space": 0, "color": "D9DDE3"},
                        bottom={"val": "single", "sz": 4, "space": 0, "color": "D9DDE3"})

        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.name = "Calibri"
        lr._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        lr.font.size = Pt(10)
        lr.font.color.rgb = GRAY

        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name = "Calibri"
        vr._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        vr.font.size = Pt(10.5)

    document.add_paragraph()


def add_n8n_flow(document: Document, number: str, name: str, trigger: str, actions: list[str], integrations: list[str]):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(f"FLUXO {number} | {name.upper()}")
    r.bold = True
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r.font.size = Pt(11)
    r.font.color.rgb = RED

    add_paragraph_bold_start(document, "Trigger: ", trigger)
    add_paragraph_bold_start(document, "Acoes: ", "")
    add_bullets(document, actions)
    add_paragraph_bold_start(document, "Integracoes: ", ", ".join(integrations))
    document.add_paragraph()


def add_signature_block(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Sao Paulo, {DATE}.")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(11)

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(7.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            set_cell_border(cell, top={"val": "single", "sz": 8, "space": 0, "color": "AEB4BB"})
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left.add_run("LOCALRISE ADVISORY\n").bold = True
    left.add_run("Nicolas Slater\n")
    left.add_run("Diretoria Comercial")
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right.add_run(f"{COMPANY}\n").bold = True
    right.add_run("Responsavel: ______________________\n")
    right.add_run("Cargo: ___________________________")


def build_doc():
    doc = Document()
    apply_doc_style(doc)

    # ── CAPA ────────────────────────────────────────────────────────────────────
    add_brand_header(
        doc,
        "Estudo de Automacoes",
        "Capitol Consultoria | sistema de aquisicao, conversao e acompanhamento",
    )
    doc.add_page_break()

    # ── INTRODUCAO ──────────────────────────────────────────────────────────────
    add_section_title(doc, "Contexto e objetivo deste estudo")
    add_paragraph(
        doc,
        "A Capitol Consultoria tem algo que a maioria das empresas leva anos para construir: credibilidade real, senioridade institucional e um tema de alta demanda. O problema nao e o produto. O problema e que esse valor nao existe digitalmente. Ele fica restrito a rede proxima, a indicacoes e a conversas que dependem do acaso.",
    )
    add_paragraph(
        doc,
        "Este estudo foi construido para mostrar, com clareza, quais sistemas de automacao precisam existir para que a Capitol deixe de depender de sorte e passe a depender de estrutura. Cada automacao aqui descrita resolve um problema real identificado no diagnostico.",
    )
    add_highlight_box(
        doc,
        "A premissa central e simples: quem nao tem sistema nao tem previsibilidade. E quem nao tem previsibilidade nao cresce com consistencia.",
        "FDF3F3",
    )
    doc.add_page_break()

    # ── BLOCO 1: GARGALOS ───────────────────────────────────────────────────────
    add_section_title(doc, "1. Onde estao os gargalos reais")
    add_paragraph(
        doc,
        "Antes de falar em automacao, precisamos entender onde o negocio perde oportunidades hoje. Identificamos tres gargalos principais — cada um representa receita que nao foi capturada.",
    )

    add_subsection_title(doc, "Gargalo 1 | Nao ha porta de entrada digital")
    add_paragraph(
        doc,
        "A Capitol nao tem site funcional e visivel, nem landing page, nem formulario ativo. Quando alguem ouve falar da empresa e vai pesquisar, nao encontra nada solido. O interessado esfria. O lead se perde. Isso acontece todos os dias, sem que ninguem perceba, porque nao ha sistema para registrar.",
    )
    add_highlight_box(doc, "Impacto direto: leads que chegam por indicacao ou curiosidade nao tem para onde ir. A conversao dependeria de um contato proativo — que muitas vezes nao acontece.")

    add_subsection_title(doc, "Gargalo 2 | Nao ha sequencia de contato")
    add_paragraph(
        doc,
        "Quando alguem entra em contato — seja pelo Instagram, seja por telefone — o que acontece depois depende inteiramente da memoria e disponibilidade da equipe. Nao ha resposta automatica, nao ha sequencia de follow-up, nao ha sistema de qualificacao. O lead fica esperando. E quem espera, muitas vezes, vai embora.",
    )
    add_highlight_box(doc, "Impacto direto: mesmo os leads que chegam por indicacao qualificada podem esfriar por falta de resposta rapida ou acompanhamento estruturado.")

    add_subsection_title(doc, "Gargalo 3 | Nao ha acompanhamento pos-reuniao")
    add_paragraph(
        doc,
        "Quando uma reuniao de diagnostico ou apresentacao acontece, o que vem depois e improviso. Nao ha material automatico de reforcamento, nao ha sequencia de follow-up, nao ha CRM que registre o historico. O fechamento depende da memoria do consultor e da proatividade do prospect.",
    )
    add_highlight_box(doc, "Impacto direto: negociações que poderiam fechar em 2 semanas se arrastam por meses — ou simplesmente somem do radar.")

    doc.add_page_break()

    # ── BLOCO 2: AUTOMACOES ─────────────────────────────────────────────────────
    add_section_title(doc, "2. Automacoes necessarias — visao geral")
    add_paragraph(
        doc,
        "Com base nos gargalos identificados, estruturamos 10 automacoes organizadas em tres frentes: captacao, conversao e acompanhamento. Cada uma resolve um problema especifico do diagnostico.",
    )

    # Tabela visao geral
    headers = ["#", "Automacao", "Problema que resolve", "Prioridade"]
    rows_data = [
        ["1", "Captacao de lead pelo site", "Ausencia de porta de entrada digital", "ESSENCIAL"],
        ["2", "Resposta automatica no WhatsApp", "Sem atendimento imediato", "ESSENCIAL"],
        ["3", "Qualificacao automatica do lead", "Leads sem triagem chegam a reuniao sem preparo", "ESSENCIAL"],
        ["4", "Agendamento automatico de reuniao", "Processo de agendamento e manual e lento", "ESSENCIAL"],
        ["5", "Envio automatico de material de autoridade", "Nenhum aquecimento entre contato e reuniao", "ESSENCIAL"],
        ["6", "Follow-up automatico pos-reuniao", "Negociacoes esfriam por falta de acompanhamento", "INTERMEDIARIA"],
        ["7", "Recuperacao de leads perdidos", "Leads que nao converteram ficam esquecidos", "INTERMEDIARIA"],
        ["8", "Nutricao de leads frios", "Base de contatos nao explorada", "INTERMEDIARIA"],
        ["9", "Prova social automatizada", "Depoimentos e casos nao chegam ao prospect", "INTERMEDIARIA"],
        ["10", "Relatorio automatico de desempenho", "Sem visibilidade do funil para gestao", "AVANCADA"],
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Cm(0.8), Cm(5.2), Cm(7.0), Cm(3.2)]
    hrow = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = hrow.cells[i]
        cell.width = w
        set_cell_bg(cell, "0A0A0A")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, row_vals in enumerate(rows_data):
        row = table.rows[ri + 1]
        bg = "FFFFFF" if ri % 2 == 0 else "F9F9F9"
        for ci, (val, w) in enumerate(zip(row_vals, widths)):
            cell = row.cells[ci]
            cell.width = w
            set_cell_bg(cell, bg)
            set_cell_border(cell, bottom={"val": "single", "sz": 4, "space": 0, "color": "E8E8E8"})
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Calibri"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
            r.font.size = Pt(9.5)
            if ci == 3:
                if val == "ESSENCIAL":
                    r.font.color.rgb = RED
                    r.bold = True
                elif val == "INTERMEDIARIA":
                    r.font.color.rgb = RGBColor(0xD4, 0x7B, 0x00)

    doc.add_paragraph()
    doc.add_page_break()

    # ── BLOCO 3: DETALHAMENTO DAS AUTOMACOES ────────────────────────────────────
    add_section_title(doc, "3. Como funciona cada automacao")

    add_automation_card(doc, "01", "Captacao de lead pelo site", [
        ("O que faz",       "Quando alguem preenche o formulario do site — nome, empresa, cargo e mensagem — o sistema registra automaticamente o lead no CRM, envia um e-mail de confirmacao ao prospect e notifica a equipe da Capitol no WhatsApp."),
        ("Onde entra",      "Topo do funil. E o ponto de entrada de todo o sistema. Sem isso, nenhuma automacao posterior funciona."),
        ("Problema que resolve", "Hoje nao ha site nem formulario. Quando o site existir, todo lead que chegar tera registro, notificacao e primeira resposta automatica — sem depender de memoria ou disponibilidade da equipe."),
        ("Impacto gerado",  "Zero perda de leads qualificados. Toda pessoa que demonstrar interesse tera um primeiro contato imediato e uma ficha no sistema. A Capitol passa a ter controle sobre quem chegou e o que aconteceu depois."),
    ])

    add_automation_card(doc, "02", "Resposta automatica no WhatsApp", [
        ("O que faz",       "Quando um lead envia mensagem direta pelo WhatsApp — seja vindo do formulario do site, do link da bio do Instagram ou de um anuncio — ele recebe uma resposta imediata com apresentacao curta, confirmacao de recebimento e proximos passos."),
        ("Onde entra",      "Ponto de contato inicial. E a transicao entre o interesse e a conversa qualificada."),
        ("Problema que resolve", "Hoje, se alguem manda mensagem fora do horario comercial ou num momento de alta demanda, nao recebe resposta rapida. Leads que nao recebem resposta em ate 5 minutos tem 10 vezes menos chance de converter."),
        ("Impacto gerado",  "A Capitol passa a ter atendimento imediato 24 horas, sete dias por semana, para o primeiro contato — sem contratar ninguem. Isso aumenta conversao logo na primeira interacao."),
    ])

    add_automation_card(doc, "03", "Qualificacao automatica do lead", [
        ("O que faz",       "Apos o primeiro contato, o sistema envia automaticamente 2 ou 3 perguntas por WhatsApp ou e-mail para entender o perfil do lead: qual o setor, qual o desafio atual, qual o porte da empresa. Com base nas respostas, o lead e classificado como quente, morno ou frio."),
        ("Onde entra",      "Entre o primeiro contato e a reuniao. Funciona como filtro inteligente antes de alocar tempo da equipe."),
        ("Problema que resolve", "Hoje nao ha triagem. Qualquer contato pode virar uma reuniao demorada com alguem que nao tem perfil. Isso desperdiça tempo de quem e senior e precioso."),
        ("Impacto gerado",  "A equipe so vai a reuniao com leads que ja demonstraram encaixe. O tempo dos consultores e protegido. A taxa de fechamento de reunioes qualificadas sobe significativamente."),
    ])

    add_automation_card(doc, "04", "Agendamento automatico de reuniao", [
        ("O que faz",       "Quando o lead e qualificado como quente, o sistema envia automaticamente um link de agendamento — integrado com o Google Calendar do consultor responsavel. O lead escolhe o horario disponivel, confirma e recebe os detalhes por e-mail e WhatsApp."),
        ("Onde entra",      "Transicao entre qualificacao e reuniao. Elimina a troca manual de mensagens para combinar horario."),
        ("Problema que resolve", "Hoje o agendamento e feito por troca de mensagens — lento, sujeito a erros, dependente de disponibilidade imediata. Cada passo de fricao reduz a chance de o lead aparecer."),
        ("Impacto gerado",  "O agendamento acontece em minutos, sem intervencao humana. Leads que estao quentes convertem antes de esfriar. O consultor ja recebe a reuniao pronta no calendario."),
    ])

    add_automation_card(doc, "05", "Envio automatico de material de autoridade", [
        ("O que faz",       "Apos o agendamento confirmado, o sistema envia automaticamente um material curto sobre a Capitol: quem somos, como pensamos, um caso ilustrativo ou insight relevante. Pode ser um PDF de 1 pagina, um link para o site ou um video curto."),
        ("Onde entra",      "Entre o agendamento e a reuniao. Funciona como preparo psicologico e validacao de autoridade antes do encontro."),
        ("Problema que resolve", "Hoje o lead chega na reuniao sem ter absorvido nada sobre a Capitol. O consultor precisa gastar os primeiros 15 minutos apresentando a empresa — tempo que poderia ser usado para qualificar e fechar."),
        ("Impacto gerado",  "O lead chega preparado, aquecido e com percepcao de autoridade formada. A reuniao e mais eficiente. O ciclo de fechamento encurta."),
    ])

    add_automation_card(doc, "06", "Follow-up automatico pos-reuniao", [
        ("O que faz",       "Em ate 2 horas apos a reuniao, o sistema envia automaticamente uma mensagem de agradecimento, o resumo dos pontos discutidos e o proximo passo acordado. Nos dias 3, 7 e 14 seguintes, envia lembretes suaves caso nao haja resposta."),
        ("Onde entra",      "Depois da reuniao. E o sistema que sustenta o processo comercial sem depender de memoria da equipe."),
        ("Problema que resolve", "Hoje o follow-up depende inteiramente da proatividade manual. Reunioes que poderiam fechar em 1 semana se arrastam por falta de acompanhamento."),
        ("Impacto gerado",  "Nenhuma negociacao cai por esquecimento. O lead se sente acompanhado e tem mais confianca para tomar decisao. A taxa de fechamento de reunioes realizadas aumenta consideravelmente."),
    ])

    add_automation_card(doc, "07", "Recuperacao de leads perdidos", [
        ("O que faz",       "Leads que chegaram ao sistema, mas nao responderam ou pararam o processo em algum ponto, recebem uma sequencia reativacao automatica — geralmente apos 30 dias. Uma mensagem curta, sem pressao, perguntando se o momento mudou ou oferecendo um conteudo novo."),
        ("Onde entra",      "Base de leads inativos. E uma camada de recuperacao de investimento ja feito."),
        ("Problema que resolve", "Hoje esses leads ficam esquecidos. Nao ha banco de dados, nao ha reativacao. Sao pessoas que demonstraram interesse e simplesmente sumiram do radar."),
        ("Impacto gerado",  "Entre 15% e 30% dos leads que nao converteram na primeira interacao convertem em reativacoes futuras. Isso gera receita a custo zero — nenhum novo lead precisa ser capturado."),
    ])

    add_automation_card(doc, "08", "Nutricao de leads frios", [
        ("O que faz",       "Leads que nao estao prontos para comprar, mas demonstraram interesse, recebem uma sequencia de e-mails ou mensagens com conteudo de autoridade — insights sobre Brasilia, leituras de cenario, alertas regulatorios — a cada 15 dias."),
        ("Onde entra",      "Funil de medio prazo. Mantém a Capitol relevante para decisores que ainda nao tem urgencia."),
        ("Problema que resolve", "Alta demanda no mercado de relacoes governamentais e sazonal — ela aparece quando algo urgente acontece. A Capitol precisa estar na cabeca do decisor quando a urgencia chegar."),
        ("Impacto gerado",  "A Capitol e lembrada quando o momento certo chegar. O lead que nao estava pronto volta quando o cenario muda — e ja tem relacao estabelecida com a marca."),
    ])

    add_automation_card(doc, "09", "Prova social automatizada", [
        ("O que faz",       "Apos o encerramento de um servico, o sistema envia automaticamente uma solicitacao de depoimento para o cliente — por WhatsApp ou e-mail. Com o depoimento autorizado, ele e adicionado automaticamente ao site e disparado na sequencia de nutricao de novos leads."),
        ("Onde entra",      "Pos-venda e funil de novos leads. Transforma experiencia real em prova de mercado."),
        ("Problema que resolve", "A Capitol certamente tem resultados e historias reais. Mas esses resultados ficam apenas na memoria dos clientes, nao viram prova publica. Isso reduz a capacidade de convencer novos prospects."),
        ("Impacto gerado",  "A prova social e o ativo mais poderoso em servicos de alto valor. Um depoimento de CEO ou diretor de empresa regulada vale mais do que qualquer copy. A automacao garante que esse ativo seja capturado sistematicamente."),
    ])

    add_automation_card(doc, "10", "Relatorio automatico de desempenho do funil", [
        ("O que faz",       "Toda semana, o sistema gera e envia automaticamente um relatorio simples com: quantos leads entraram, em qual estagio estao, quantas reunioes foram agendadas, quantas propostas estao abertas e qual e a taxa de conversao."),
        ("Onde entra",      "Gestao estrategica. Nao e uma automacao de vendas — e uma automacao de visibilidade sobre o sistema inteiro."),
        ("Problema que resolve", "Hoje nao ha visibilidade sobre o funil. Nao se sabe quantos leads chegaram, quantos sumiram, onde estao os gargalos. Sem esse dado, e impossivel melhorar o sistema."),
        ("Impacto gerado",  "A Capitol passa a tomar decisoes com base em dados reais, nao em sensacoes. O socio sabe, toda semana, o que esta funcionando e o que precisa de ajuste."),
    ])

    doc.add_page_break()

    # ── BLOCO 4: FLUXO COMPLETO ─────────────────────────────────────────────────
    add_section_title(doc, "4. Fluxo completo do sistema automatizado")
    add_paragraph(
        doc,
        "O fluxo a seguir mostra como todas as automacoes se conectam em uma jornada continua — do primeiro contato ao fechamento e ao pos-venda. Cada etapa foi pensada para eliminar os gargalos identificados no diagnostico.",
    )

    stages = [
        ("ETAPA 1", "PRESENCA DIGITAL ATIVA",
         "Instagram com conteudo de autoridade + site com formulario ativo + Meta Ads (fase 3)",
         "O lead descobre a Capitol, ve credibilidade e decide entrar em contato."),
        ("ETAPA 2", "CAPTURA DO LEAD",
         "Formulario do site ou link da bio do Instagram — lead preenche nome, empresa, cargo e necessidade",
         "O sistema registra o lead, notifica a equipe e dispara a primeira resposta automatica."),
        ("ETAPA 3", "QUALIFICACAO AUTOMATICA",
         "WhatsApp ou e-mail com 2 a 3 perguntas de qualificacao — setor, desafio, urgencia",
         "O lead e classificado como quente, morno ou frio. Apenas os quentes avancam para agendamento imediato."),
        ("ETAPA 4", "AGENDAMENTO E PREPARO",
         "Link de agendamento automatico + envio de material de autoridade antes da reuniao",
         "O lead marca o horario sem fricao e chega preparado para a conversa."),
        ("ETAPA 5", "REUNIAO E PROPOSTA",
         "Reuniao de diagnostico conduzida pela equipe da Capitol — sem distracao operacional",
         "A reuniao e 100% focada em diagnostico e construcao de proposta."),
        ("ETAPA 6", "FOLLOW-UP E FECHAMENTO",
         "Sequencia automatica de acompanhamento com mensagens nos dias 1, 3, 7 e 14 pos-reuniao",
         "Nenhuma negociacao cai por falta de acompanhamento. O fechamento acontece no timing certo."),
        ("ETAPA 7", "POS-VENDA E REPUTACAO",
         "Solicitacao automatica de depoimento + reativacao de leads frios na base",
         "A Capitol transforma clientes satisfeitos em prova publica e ativa leads que nao estavam prontos antes."),
    ]

    for num, title, how, impact in stages:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.rows[0].cells[0].width = Cm(2.5)
        table.rows[0].cells[1].width = Cm(13.7)
        set_cell_bg(table.rows[0].cells[0], "E31B23")
        cell_left = table.rows[0].cells[0]
        cell_right = table.rows[0].cells[1]
        set_cell_border(cell_right, left={"val": "none"})

        lp = cell_left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(num)
        lr.bold = True
        lr.font.name = "Calibri"
        lr._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        vp = cell_right.paragraphs[0]
        vr_title = vp.add_run(title + "\n")
        vr_title.bold = True
        vr_title.font.name = "Calibri"
        vr_title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        vr_title.font.size = Pt(10.5)

        vr_how = vp.add_run(how + "\n")
        vr_how.font.name = "Calibri"
        vr_how._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        vr_how.font.size = Pt(10)
        vr_how.font.color.rgb = GRAY

        vr_impact = vp.add_run("→ " + impact)
        vr_impact.italic = True
        vr_impact.font.name = "Calibri"
        vr_impact._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        vr_impact.font.size = Pt(10)

        doc.add_paragraph()

    doc.add_page_break()

    # ── BLOCO 5: FLUXOS N8N ─────────────────────────────────────────────────────
    add_section_title(doc, "5. Automacoes com n8n — como colocar em pratica")
    add_paragraph(
        doc,
        "O n8n e a ferramenta que conecta todos os sistemas da Capitol em um unico fluxo automatizado. Ele funciona como o cerebro operacional: recebe informacoes de um lugar e dispara acoes em outro, sem intervencao humana. Os fluxos abaixo mostram como implementar as automacoes na pratica.",
    )

    add_n8n_flow(doc, "01", "Captura e notificacao de lead",
        trigger="Formulario do site e submetido com dados do lead",
        actions=[
            "Cria ficha do lead no CRM (Notion, Airtable ou HubSpot)",
            "Envia e-mail de confirmacao ao lead com mensagem personalizada",
            "Dispara notificacao no WhatsApp da equipe com dados do lead",
            "Inicia sequencia de qualificacao automaticamente",
        ],
        integrations=["Site (Webflow ou WordPress)", "CRM", "WhatsApp (Evolution API ou Twilio)", "Gmail"],
    )

    add_n8n_flow(doc, "02", "Qualificacao e roteamento",
        trigger="Lead responde as perguntas de qualificacao no WhatsApp",
        actions=[
            "Analisa respostas e classifica o lead (quente / morno / frio)",
            "Se quente: envia link de agendamento automatico (Calendly ou Cal.com)",
            "Se morno: entra na sequencia de nutricao com cadencia de 15 dias",
            "Se frio: aguarda 30 dias e dispara reativacao automatica",
            "Atualiza o status do lead no CRM em tempo real",
        ],
        integrations=["WhatsApp", "Calendly ou Cal.com", "CRM", "Google Calendar"],
    )

    add_n8n_flow(doc, "03", "Preparo pre-reuniao",
        trigger="Agendamento confirmado no Calendly ou Cal.com",
        actions=[
            "Envia e-mail de confirmacao com data, horario e link da reuniao",
            "24 horas antes: envia lembrete por WhatsApp",
            "1 hora antes: dispara PDF de apresentacao institucional ou link do site",
            "Registra o agendamento no CRM com status 'reuniao marcada'",
        ],
        integrations=["Calendly ou Cal.com", "Gmail", "WhatsApp", "Google Drive", "CRM"],
    )

    add_n8n_flow(doc, "04", "Follow-up pos-reuniao",
        trigger="Status do lead atualizado para 'reuniao realizada' no CRM",
        actions=[
            "Hora 2 pos-reuniao: envia mensagem de agradecimento com resumo e proximo passo",
            "Dia 3: envia mensagem de acompanhamento caso nao haja resposta",
            "Dia 7: envia conteudo de autoridade relevante ao setor do lead",
            "Dia 14: envia mensagem final de reativacao educada",
            "Se resposta positiva: atualiza CRM e notifica equipe para acao comercial",
        ],
        integrations=["CRM", "WhatsApp", "Gmail"],
    )

    add_n8n_flow(doc, "05", "Reativacao de leads inativos",
        trigger="Lead sem interacao ha 30 dias no CRM",
        actions=[
            "Envia mensagem de reativacao personalizada com gatilho de cenario atual",
            "Se responde positivamente: reinicia o fluxo de qualificacao",
            "Se nao responde: move para lista de nutricao de longo prazo (mensal)",
            "Registra todas as interacoes no historico do lead",
        ],
        integrations=["CRM", "WhatsApp", "Gmail"],
    )

    add_n8n_flow(doc, "06", "Relatorio semanal do funil",
        trigger="Toda segunda-feira as 08h (agenda automatica)",
        actions=[
            "Coleta dados do CRM: leads novos, status atual de cada um, reunioes realizadas",
            "Calcula taxa de conversao de cada etapa do funil",
            "Monta relatorio em formato simples e visual",
            "Envia relatorio por e-mail e WhatsApp ao responsavel da Capitol",
        ],
        integrations=["CRM", "Google Sheets (para dados)", "Gmail", "WhatsApp"],
    )

    doc.add_page_break()

    # ── BLOCO 6: PRIORIDADES ─────────────────────────────────────────────────────
    add_section_title(doc, "6. Prioridades de implementacao")
    add_paragraph(
        doc,
        "A implementacao deve seguir uma ordem logica. Nao adianta criar follow-up automatico se nao ha lead capturado. A ordem abaixo garante que cada nova automacao potencializa a anterior.",
    )

    priority_data = [
        ("FASE 1 — ESSENCIAL (primeiros 30 dias)",
         "E0F0FF",
         [
             "1. Site com formulario ativo — sem isso, nenhuma automacao de captacao funciona",
             "2. Resposta automatica no WhatsApp — primeiro contato imediato e essencial",
             "3. Captura e notificacao de lead no CRM — toda entrada precisa ser registrada",
             "4. Agendamento automatico de reuniao — elimina fricao e acelera o ciclo",
             "5. Envio de material pre-reuniao — o lead precisa chegar aquecido",
         ],
         "Resultado esperado em 30 dias: a Capitol deixa de perder leads no primeiro contato e ganha controle total sobre quem chegou e o que aconteceu."),
        ("FASE 2 — INTERMEDIARIA (30 a 60 dias)",
         "FFF8E0",
         [
             "6. Qualificacao automatica do lead — triagem inteligente antes de alocar tempo senior",
             "7. Follow-up automatico pos-reuniao — o ciclo comercial ganha consistencia",
             "8. Recuperacao de leads perdidos — aproveita a base ja construida",
             "9. Nutricao de leads frios — mantém a Capitol relevante no medio prazo",
         ],
         "Resultado esperado em 60 dias: o funil completo esta operando. Nenhuma oportunidade cai por falta de acompanhamento."),
        ("FASE 3 — AVANCADA (60 a 90 dias)",
         "F0FFE8",
         [
             "10. Prova social automatizada — depoimentos viram ativo sistematico de vendas",
             "11. Relatorio automatico do funil — decisoes com base em dados, nao em intuicao",
             "12. Integracao com Meta Ads — trafego pago com base de conversao pronta para receber leads",
         ],
         "Resultado esperado em 90 dias: o sistema esta completo, mensuravel e preparado para escalar com trafego pago sem desperdicar verba."),
    ]

    for phase_title, bg, items, result in priority_data:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_bg(cell, bg)
        set_cell_border(cell,
                        top={"val": "single", "sz": 8, "space": 0, "color": "AEB4BB"},
                        bottom={"val": "single", "sz": 8, "space": 0, "color": "AEB4BB"})
        ph = cell.paragraphs[0]
        pr = ph.add_run(phase_title)
        pr.bold = True
        pr.font.name = "Calibri"
        pr._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        pr.font.size = Pt(11)
        doc.add_paragraph()

        add_bullets(doc, items)
        add_highlight_box(doc, result)

    doc.add_page_break()

    # ── BLOCO 7: IMPACTO ESTRATEGICO ────────────────────────────────────────────
    add_section_title(doc, "7. Impacto estrategico — o que muda para a Capitol")
    add_paragraph(
        doc,
        "Este bloco responde a pergunta mais importante: vale o investimento? E a resposta e sim — mas nao pelo custo da automacao. Pelo custo do que acontece quando nao ha sistema.",
    )

    add_subsection_title(doc, "Reducao da dependencia de indicacao")
    add_paragraph(
        doc,
        "Indicacao e o modelo mais caro de crescimento — nao em dinheiro, mas em controle. A Capitol nao tem como prever quantas indicacoes vao chegar no proximo mes. Com o sistema em operacao, a empresa passa a ter uma fonte previsivel de leads qualificados: Instagram alimentando o topo, site convertendo o meio e automacoes sustentando o fundo.",
    )
    add_highlight_box(doc, "Com o sistema rodando, a Capitol passa de 'dependente de sorte' para 'operando com sistema'. A diferenca e a mesma entre esperar chuver e ter uma rede de irrigacao.")

    add_subsection_title(doc, "Aumento de geracao de leads qualificados")
    add_paragraph(
        doc,
        "Hoje, o volume de leads e basicamente zero — porque nao ha entrada digital. Com site, formulario e Instagram estruturado, a estimativa conservadora e de 8 a 20 leads qualificados por mes no primeiro trimestre. Com Meta Ads ativado, esse numero pode triplicar em 90 dias.",
    )

    add_subsection_title(doc, "Melhora na taxa de conversao")
    add_paragraph(
        doc,
        "Sem sistema, a conversao depende de talento individual e memoria. Com sistema, ela depende de processo. Leads que chegam preparados convertem mais. Follow-up estruturado fecha mais. Prova social reduz objecao. Uma empresa de consultoria B2B com sistema bem estruturado costuma ver sua taxa de fechamento subir de 20% para 40% em 90 dias de operacao — sem contratar ninguem a mais.",
    )

    add_subsection_title(doc, "Previsibilidade operacional e de receita")
    add_paragraph(
        doc,
        "Com relatorio semanal do funil, a lideranca da Capitol consegue ver, com clareza, quantas reunioes estao marcadas para as proximas 2 semanas, quantas propostas estao abertas e qual e a probabilidade de fechamento. Isso transforma a gestao comercial: de reativa para proativa.",
    )
    add_highlight_box(doc, "Previsibilidade nao e luxo — e o que separa uma consultoria que cresce de uma que oscila.")

    doc.add_page_break()

    # ── BLOCO 8: POSICIONAMENTO LOCALRISE ────────────────────────────────────────
    add_section_title(doc, "8. O papel da LocalRise nesse processo")
    add_paragraph(
        doc,
        "A LocalRise nao e uma agencia que entrega conteudo. E uma parceira de sistema. Nosso papel e estruturar a base digital da Capitol — site, Instagram, automacoes e, depois, trafego pago — de forma que cada peca se conecte com a outra e gere resultado composto.",
    )
    add_paragraph(
        doc,
        "Construimos o sistema. A equipe da Capitol faz o que faz de melhor: entender o cenario politico e institucional, construir relacionamentos e entregar resultados reais para os clientes. Enquanto isso, o sistema trabalha para garantir que novas oportunidades continuem chegando.",
    )
    add_highlight_box(
        doc,
        "A Capitol tem o produto certo para o mercado certo. O que falta e o sistema que conecta essa competencia a quem ainda nao a conhece.",
        "FDF3F3",
    )

    add_section_title(doc, "Proximos passos recomendados")
    add_numbered(doc, [
        "Validar o escopo da Fase 1: site com formulario, resposta automatica no WhatsApp e CRM basico.",
        "Definir ferramenta de CRM: Notion, Airtable ou HubSpot — a escolha depende do volume e do orcamento.",
        "Integrar WhatsApp Business API para permitir automacoes de mensagem.",
        "Ativar os primeiros tres fluxos no n8n: captura, notificacao e agendamento.",
        "Revisar resultados em 30 dias e ativar a Fase 2 de follow-up e qualificacao.",
    ])

    # ── ASSINATURA ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_signature_block(doc)

    out = BASE_DIR / "estudo-automacoes-capitol.docx"
    doc.save(out)
    print(f"Documento salvo em: {out}")


if __name__ == "__main__":
    build_doc()
