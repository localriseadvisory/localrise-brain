from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
LOGO_PATH = Path(
    r"C:\Users\digui\Documents\localrise-brain\Claude\Claude Local RIse\assents\logo Local Rise Principal.png"
)

RED    = RGBColor(0xE3, 0x1B, 0x23)
BLACK  = RGBColor(0x0A, 0x0A, 0x0A)
GRAY   = RGBColor(0x5F, 0x63, 0x68)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLUE   = RGBColor(0x1A, 0x56, 0xC8)
GREEN  = RGBColor(0x1A, 0x7A, 0x3C)
ORANGE = RGBColor(0xC8, 0x6A, 0x00)
PURPLE = RGBColor(0x5A, 0x1A, 0x8C)

DATE = "10 de abril de 2026"


# ── primitivos ───────────────────────────────────────────────────────────────────

def set_border(cell, **kw):
    tc = cell._tc; tc_pr = tc.get_or_add_tcPr()
    tb = tc_pr.first_child_found_in("w:tcBorders")
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tc_pr.append(tb)
    for edge, data in kw.items():
        tag = "w:%s" % edge
        el = tb.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); tb.append(el)
        for k, v in data.items():
            el.set(qn("w:%s" % k), str(v))


def bg(cell, hex_color):
    tc = cell._tc; tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tc_pr.append(shd)


def styles(doc):
    s = doc.sections[0]
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    n.font.size = Pt(11)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    n.paragraph_format.space_after = Pt(6)


def r(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color


def jp(doc, text, size=11, bold=False, color=None, italic=False, sa=8):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.space_after = Pt(sa)
    run = par.add_run(text); r(run, size=size, bold=bold, color=color, italic=italic)


def blt(doc, text, size=11, color=None):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.space_after = Pt(3)
    run = par.add_run(text); r(run, size=size, color=color)


def sec(doc, text, accent=RED):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(16)
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(text.upper()); r(run, size=13, bold=True, color=accent)


def sub(doc, text, color=BLACK, sb=10):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(sb)
    par.paragraph_format.space_after = Pt(2)
    run = par.add_run(text); r(run, size=11.5, bold=True, color=color)


def divider(doc, color="E0E0E0"):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_border(c, bottom={"val": "single", "sz": "6", "space": "0", "color": color})
    c.text = ""; doc.add_paragraph()


def callout(doc, text, bg_hex="FDF3F3", border="E31B23", size=11):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); bg(c, bg_hex)
    set_border(c, left={"val": "single", "sz": "18", "space": "0", "color": border})
    par = c.paragraphs[0]; par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = par.add_run(text); r(run, size=size, italic=True)
    doc.add_paragraph()


def brand_header(doc):
    if LOGO_PATH.exists():
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(LOGO_PATH), width=Cm(5.4))
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run("LOCALRISE ADVISORY"); r(run, size=12, bold=True, color=RED)
    divider(doc)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(8)
    run = par.add_run("Inteligencia Regulatoria — Versao Comercial")
    r(run, size=20, bold=True)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run("O que entregamos, como funciona e por que isso importa")
    r(run, size=11, color=GRAY)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(DATE); r(run, size=10, color=GRAY, italic=True)


def signature(doc):
    par = doc.add_paragraph(); run = par.add_run(f"Sao Paulo, {DATE}.")
    r(run, size=11, bold=True)
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for cell in row.cells:
            cell.width = Cm(7.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            set_border(cell, top={"val": "single", "sz": "8", "space": "0", "color": "AEB4BB"})
    lp = t.cell(0, 0).paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run("LOCALRISE ADVISORY\n").bold = True
    lp.add_run("Nicolas Slater\n"); lp.add_run("Diretoria Comercial")


# ── componentes do documento ──────────────────────────────────────────────────────

def product_card(doc, emoji, name, one_liner, problem, what_it_does, value, example, accent_hex, bg_hex):
    # Cabecalho colorido
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc, rc = t.rows[0].cells
    lc.width = Cm(2.0); rc.width = Cm(14.2)
    bg(lc, accent_hex); bg(rc, "F9F9F9")
    lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lp.add_run(emoji + "\n" + name)
    r(run, size=10, bold=True, color=WHITE)
    rp = rc.paragraphs[0]
    run2 = rp.add_run(one_liner); r(run2, size=11.5, bold=True)
    doc.add_paragraph()

    # Os 5 blocos do card
    blocks = [
        ("O PROBLEMA", problem,   "FFF8F8", "E31B23"),
        ("O QUE FAZ",  what_it_does, "F8F8FF", "1A56C8"),
        ("POR QUE VALE", value,   "F8FFF9", "1A7A3C"),
        ("EXEMPLO REAL", example, "FFFAF0", "C86A00"),
    ]
    for label, content, bg2, border in blocks:
        t2 = doc.add_table(rows=1, cols=2); t2.alignment = WD_TABLE_ALIGNMENT.LEFT
        label_cell, content_cell = t2.rows[0].cells
        label_cell.width = Cm(3.5); content_cell.width = Cm(12.7)
        bg(label_cell, "F4F4F4"); bg(content_cell, bg2)
        set_border(label_cell, bottom={"val": "single", "sz": "4", "space": "0", "color": "E8E8E8"})
        set_border(content_cell, bottom={"val": "single", "sz": "4", "space": "0", "color": "E8E8E8"},
                   left={"val": "single", "sz": "10", "space": "0", "color": border})
        lpar = label_cell.paragraphs[0]; lr = lpar.add_run(label)
        r(lr, size=9, bold=True, color=GRAY)
        cpar = content_cell.paragraphs[0]; cr = cpar.add_run(content)
        r(cr, size=10.5)
    doc.add_paragraph()


def automation_card(doc, name, before, after, time_saved, impact):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); bg(c, "F9F9F9")
    set_border(c, left={"val": "single", "sz": "14", "space": "0", "color": "E31B23"})
    hpar = c.paragraphs[0]; hr = hpar.add_run(name)
    r(hr, size=11, bold=True)
    doc.add_paragraph()

    # Antes / Depois / Tempo / Impacto
    rows = [
        ("ANTES", before, "FFF4F4"),
        ("DEPOIS", after,  "F4FFF4"),
        ("TEMPO ECONOMIZADO", time_saved, "F4F4FF"),
        ("IMPACTO NO NEGOCIO", impact, "FFFCF0"),
    ]
    t2 = doc.add_table(rows=len(rows), cols=2); t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, content, color) in enumerate(rows):
        lc, vc = t2.rows[i].cells
        lc.width = Cm(4.5); vc.width = Cm(11.7)
        bg(lc, "EEEEEE"); bg(vc, color)
        set_border(lc, bottom={"val": "single", "sz": "4", "space": "0", "color": "DDDDDD"})
        set_border(vc, bottom={"val": "single", "sz": "4", "space": "0", "color": "DDDDDD"})
        lpar = lc.paragraphs[0]; lr = lpar.add_run(label)
        r(lr, size=9.5, bold=True, color=GRAY)
        cpar = vc.paragraphs[0]; cr = cpar.add_run(content)
        r(cr, size=10.5)
    doc.add_paragraph()


def value_row(doc, icon, category_name, items, accent_hex, bg_hex):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); bg(c, accent_hex)
    hpar = c.paragraphs[0]; hpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hpar.add_run(icon + "  " + category_name.upper())
    r(hr, size=12, bold=True, color=WHITE)
    doc.add_paragraph()
    for item_title, item_desc in items:
        t2 = doc.add_table(rows=1, cols=2); t2.alignment = WD_TABLE_ALIGNMENT.LEFT
        nc, dc = t2.rows[0].cells; nc.width = Cm(0.5); dc.width = Cm(15.7)
        bg(nc, accent_hex); bg(dc, bg_hex)
        dp = dc.paragraphs[0]
        dr1 = dp.add_run(item_title + " — "); r(dr1, size=11, bold=True)
        dr2 = dp.add_run(item_desc); r(dr2, size=10.5)
        doc.add_paragraph()


def pitch_card(doc, headline, body, accent_hex):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); bg(c, accent_hex)
    hp = c.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(headline.upper()); r(hr, size=12, bold=True, color=WHITE)
    doc.add_paragraph()
    jp(doc, body, size=11)
    doc.add_paragraph()


def improvement_item(doc, category, description, category_color):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.space_after = Pt(6)
    cat_run = par.add_run("[" + category + "] "); r(cat_run, size=11, bold=True, color=category_color)
    desc_run = par.add_run(description); r(desc_run, size=11)


# ══════════════════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    styles(doc)
    brand_header(doc)
    doc.add_page_break()

    # ── ABERTURA ─────────────────────────────────────────────────────────────────
    sec(doc, "O que voce vai encontrar aqui")
    jp(doc,
       "Este documento transforma o estudo tecnico da LocalRise em algo que qualquer decisor entende "
       "em poucos minutos. Sem jargao. Sem excesso. So o que importa: o que e, por que vale e como "
       "impacta o seu negocio.")
    callout(doc,
        "A LocalRise nao esta vendendo tecnologia. Está vendendo inteligencia, vantagem competitiva "
        "e eficiencia operacional — construidas sobre tecnologia que a maioria das consultorias "
        "ainda nao tem.",
        bg_hex="F0F4FF", border="1A56C8")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════════
    # BLOCO 1 — PRODUTOS SIMPLIFICADOS
    # ══════════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 1 | Os produtos — explicados de forma simples")
    jp(doc,
       "Cada produto abaixo resolve um problema real que empresas expostas ao ambiente "
       "regulatorio enfrentam todo dia. Nenhum deles exige que o cliente entenda de tecnologia.")
    doc.add_paragraph()

    product_card(doc,
        emoji="01", name="LexRadar",
        one_liner="Voce sabe quando uma lei muda. Antes de todo mundo.",
        problem="Sua equipe precisa acompanhar dezenas de projetos de lei, decretos e normas. "
                "Isso consome horas todo dia — e mesmo assim, mudancas criticas passam despercebidas.",
        what_it_does="Monitora automaticamente tudo que muda no ambiente regulatorio do seu setor. "
                     "Quando algo relevante aparece, voce recebe um resumo direto: o que mudou, "
                     "quem e afetado e qual e a urgencia de agir.",
        value="Voce para de perder mudancas que impactam o negocio. Para de gastar horas lendo diarios "
              "oficiais. E chega nas reunioes com informacao que seus concorrentes ainda nao tem.",
        example="Uma nova instrucao normativa da ANATEL e publicada na quinta-feira. Na sexta de manha, "
                "o CEO da empresa de telecomunicacoes ja recebeu pelo WhatsApp: o que mudou, o impacto "
                "estimado e a acao recomendada — sem precisar pedir nada a ninguem.",
        accent_hex="E31B23", bg_hex="F9F9F9")

    product_card(doc,
        emoji="02", name="PoliticaIQ",
        one_liner="Saiba quem decide, quem influencia quem — e qual e o momento certo de agir.",
        problem="Navegar em Brasilia sem entender quem realmente decide e quem influencia quem e como "
                "entrar numa sala sem conhecer o andar do predio. Tempo e dinheiro desperdicados.",
        what_it_does="Mantém um mapa atualizado dos principais decisores politicos do seu setor: "
                     "como pensam, como votaram, quem os influencia, quando seus mandatos vencem "
                     "e qual o melhor caminho para chegar ate eles.",
        value="Voce nao age no timing errado. Nao aborda a pessoa errada. Nao toma decisoes sem "
              "entender quem esta por tras delas. O acesso que antes dependia de anos de relacionamento "
              "passa a ter estrutura e visibilidade.",
        example="A empresa quer aprovar uma agenda no Senado. Em vez de agir por intuicao, o PoliticaIQ "
                "mostra: quais senadores sao favoraveis, quem os influencia, qual e o caminho mais curto "
                "para articulacao — e qual e o momento certo para agir.",
        accent_hex="1A56C8", bg_hex="F9F9F9")

    product_card(doc,
        emoji="03", name="RegRisk",
        one_liner="Uma nota diaria de risco regulatorio para o seu negocio.",
        problem="Risco regulatorio so e levado a serio quando a crise ja chegou. "
                "Ate la, a empresa operou sem visibilidade — e pagou caro por isso.",
        what_it_does="Calcula diariamente um score de risco do ambiente regulatorio para o seu setor. "
                     "Mostra qual e a sua exposicao atual, o que esta se movendo e qual a probabilidade "
                     "de impacto nos proximos 30, 60 e 90 dias.",
        value="Voce passa de reativo para proativo. A empresa age antes da crise — nao depois. "
              "A diretoria tem um numero claro toda semana: o risco subiu, desceu ou esta estavel.",
        example="O RegRisk indica na segunda-feira que o risco para o setor farmaceutico subiu de 42 "
                "para 67 nos ultimos 7 dias por conta de movimentacoes em duas comissoes do Senado. "
                "A diretoria convoca uma reuniao de strategy antes de qualquer crise estourar.",
        accent_hex="1A7A3C", bg_hex="F9F9F9")

    product_card(doc,
        emoji="04", name="OpportunityMap",
        one_liner="Descubra incentivos fiscais, financiamentos e programas publicos antes do prazo.",
        problem="Bilhoes de reais em incentivos fiscais, linhas de credito e programas de fomento "
                "existem — mas chegam ao conhecimento das empresas por acaso, por indicacao ou "
                "depois do prazo.",
        what_it_does="Monitora continuamente todos os programas, editais e linhas de financiamento "
                     "publico disponiveis. Cruza com o perfil da sua empresa e avisa quando uma "
                     "oportunidade relevante aparece — com prazo, beneficio estimado e proxima acao.",
        value="Voce para de perder dinheiro que ja era seu. Oportunidades que voce nem sabia que "
              "existiam passam a chegar organizadas, com contexto e antes do prazo.",
        example="Uma industria do agronegocio descobre pelo OpportunityMap que ha uma linha do BNDES "
                "com taxa subsidiada especificamente para o seu perfil — com prazo abrindo em 15 dias. "
                "Sem o sistema, essa oportunidade teria passado sem que ninguem soubesse.",
        accent_hex="C86A00", bg_hex="F9F9F9")

    product_card(doc,
        emoji="05", name="CrisisMonitor",
        one_liner="Saber antes que a crise chegue. Agir antes de todo mundo.",
        problem="Crises regulatorias nao aparecem do nada. Elas tem sinais precursores — declaracoes "
                "de parlamentares, movimentacoes em comissoes, mudancas de posicionamento. "
                "Quem nao monitora so descobre quando o estrago ja esta feito.",
        what_it_does="Monitora 5 fontes ao mesmo tempo — diario oficial, pautas legislativas, midia, "
                     "redes sociais de decisores e historico de crises do setor. Quando um padrao "
                     "de crise emergente e identificado, a empresa recebe alerta imediato.",
        value="Um unico alerta antecipado que permite uma mudanca de posicionamento estrategico "
              "pode valer mais do que um ano de contrato. A empresa age quando ainda tem tempo — "
              "nao quando ja e tarde.",
        example="Um senador comeca a fazer declaracoes criticas ao modelo de concessoes de energia. "
                "O CrisisMonitor detecta o padrao em 48 horas. A empresa ja prepara seu posicionamento "
                "antes que o tema chegue a pauta nacional — e antes que os concorrentes percebam.",
        accent_hex="5A1A8C", bg_hex="F9F9F9")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════════
    # BLOCO 2 — AUTOMACOES SIMPLIFICADAS
    # ══════════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 2 | As automacoes — antes e depois")
    jp(doc,
       "Automacao, na pratica, significa uma coisa: o que antes dependia de horas de trabalho manual "
       "passa a acontecer sozinho — com qualidade igual ou superior. "
       "Abaixo esta o impacto concreto de cada automacao em tempo e resultado.")
    doc.add_paragraph()

    automation_card(doc,
        name="Monitoramento automatico de legislacao",
        before="Um analista passa 3 a 4 horas por dia lendo diarios oficiais, pautas e projetos de lei. "
               "Mesmo assim, mudancas relevantes passam despercebidas.",
        after="O sistema lê tudo automaticamente, filtra pelo que importa para o cliente "
              "e entrega um resumo com o que mudou, quem e afetado e o que fazer.",
        time_saved="15 a 25 horas por semana por analista. Com uma equipe de 3 analistas: "
                   "ate 75 horas livres por semana para trabalho estrategico.",
        impact="A consultoria passa a monitorar 10 vezes mais temas sem contratar mais ninguem. "
               "A qualidade e o alcance das analises aumentam. O diferencial fica claro para o cliente.")

    automation_card(doc,
        name="Alerta personalizado por cliente",
        before="Todos os clientes recebem o mesmo boletim generico de noticias. "
               "Cada um precisa filtrar o que e relevante para o seu setor — o que ninguem faz direito.",
        after="Cada cliente recebe apenas o que importa para o seu contexto: setor especifico, "
              "temas monitorados, nivel de risco. Nenhuma informacao desnecessaria, nenhuma lacuna.",
        time_saved="2 a 3 horas por semana por cliente na triagem manual de informacoes.",
        impact="O cliente sente que a consultoria entende profundamente o seu negocio. "
               "O valor percebido aumenta. A retencao de contrato melhora.")

    automation_card(doc,
        name="Relatorio semanal automatico",
        before="Cada relatorio semanal leva entre 3 e 4 horas para ser escrito, formatado e enviado. "
               "Com 10 clientes ativos, isso sao 30 a 40 horas por semana so em relatorios.",
        after="O sistema coleta os eventos da semana, escreve o relatorio em linguagem executiva, "
              "formata no padrao visual do cliente e envia por e-mail — automaticamente toda sexta.",
        time_saved="Cada relatorio sai em 15 minutos. Com 20 clientes: 75 horas recuperadas por mes.",
        impact="Os consultores param de ser maquinas de escrever relatorio e passam a ser "
               "estrategistas. A empresa atende mais clientes sem aumentar a equipe.")

    automation_card(doc,
        name="Analise de impacto de nova lei",
        before="Quando uma nova lei ou normativa aparece, o analista precisa ler, interpretar, "
               "identificar quais clientes sao afetados e escrever uma analise. Isso leva 3 horas.",
        after="O sistema lê a nova lei, identifica impactos por setor, mapeia quais clientes "
              "sao afetados e entrega um rascunho de analise pronto para revisao em 8 minutos.",
        time_saved="De 3 horas para 8 minutos por analise. Uma lei por dia = mais de 50 horas "
                   "recuperadas por mes so nessa tarefa.",
        impact="A consultoria entrega analises de impacto no mesmo dia da publicacao da lei — "
               "antes de qualquer concorrente. Isso e diferencial direto na hora de renovar contrato.")

    automation_card(doc,
        name="Acompanhamento de posicionamento de politicos",
        before="Manter um perfil atualizado de 50, 100 ou 200 decisores politicos exige uma equipe "
               "dedicada a leitura de declaracoes, votacoes e entrevistas — todos os dias.",
        after="O sistema monitora automaticamente tudo que os decisores mapeados falam, votam "
              "e publicam. Quando alguem muda de posicionamento em um tema relevante, o sistema avisa.",
        time_saved="O equivalente a 1 analista dedicado em tempo integral por cada 100 decisores monitorados.",
        impact="A consultoria entrega inteligencia politica atualizada que antes so era possivel "
               "com uma equipe grande. Com isso, pode atender mais setores e mais clientes.")

    automation_card(doc,
        name="Onboarding automatico de novo cliente",
        before="Quando um novo cliente entra, a equipe precisa configurar manualmente tudo: "
               "quais temas monitorar, quais fontes ativar, qual template usar. Isso leva de 4 a 8 horas.",
        after="O cliente preenche um formulario de perfil. O sistema configura tudo automaticamente: "
              "monitoramento, alertas, painel, relatorios. Em 30 minutos, o cliente ja esta ativo.",
        time_saved="4 a 8 horas por novo cliente. Com 5 clientes novos por mes: ate 40 horas recuperadas.",
        impact="A consultoria consegue crescer em velocidade sem gargalo operacional. "
               "A primeira impressao do cliente e de profissionalismo e velocidade.")

    automation_card(doc,
        name="Deteccao precoce de crise",
        before="A empresa so descobre que uma crise regulatoria esta se formando quando ela "
               "aparece na midia — geralmente tarde demais para agir de forma estrategica.",
        after="O sistema cruza 5 fontes simultaneamente e identifica padroes de crise antes "
              "que ela chegue ao mainstream. O alerta chega com antecedencia de 24 a 72 horas.",
        time_saved="Nao e sobre tempo economizado. E sobre decisoes que so sao possiveis "
                   "com antecedencia.",
        impact="Um unico alerta antecipado que evita uma crise ou permite um reposicionamento "
               "estrategico pode valer mais do que todos os outros servicos somados.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════════
    # BLOCO 3 — ORGANIZACAO POR VALOR
    # ══════════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 3 | O que entregamos — organizado por tipo de valor")
    jp(doc, "Tudo que a LocalRise oferece se enquadra em tres tipos de valor. "
            "Cada empresa escolhe onde quer comecar — mas os tres se potencializam.")
    doc.add_paragraph()

    value_row(doc,
        icon="$", category_name="Automacoes que economizam tempo",
        items=[
            ("Monitoramento legislativo automatico",
             "elimina 15 a 25 horas semanais de leitura manual por analista"),
            ("Relatorio semanal automatico",
             "75 horas recuperadas por mes com 20 clientes ativos"),
            ("Analise automatica de nova lei",
             "de 3 horas para 8 minutos por analise"),
            ("Onboarding automatico de cliente",
             "de 8 horas para 30 minutos por novo contrato"),
        ],
        accent_hex="1A7A3C", bg_hex="F4FFF6")

    value_row(doc,
        icon="@", category_name="Produtos que geram vantagem competitiva",
        items=[
            ("LexRadar",
             "voce sabe de mudancas regulatorias antes dos seus concorrentes"),
            ("PoliticaIQ",
             "voce entende quem decide e como chegar la — com base em dados, nao em intuicao"),
            ("OpportunityMap",
             "voce captura oportunidades de financiamento que a maioria das empresas perde"),
            ("Alerta personalizado por cliente",
             "cada cliente recebe inteligencia filtrada para o seu contexto especifico"),
        ],
        accent_hex="1A56C8", bg_hex="F4F7FF")

    value_row(doc,
        icon="!", category_name="Diferenciais avancados — IA e predicao",
        items=[
            ("RegRisk — score de risco diario",
             "um numero que resume sua exposicao regulatoria — como uma previsao do tempo para Brasilia"),
            ("CrisisMonitor — deteccao precoce",
             "alertas 24 a 72 horas antes de uma crise aparecer na midia"),
            ("Scoring preditivo",
             "probabilidade de um evento regulatorio acontecer nos proximos 30, 60 e 90 dias"),
            ("Assistente de inteligencia por IA",
             "responde perguntas regulatorias complexas em segundos, com contexto especifico do setor"),
        ],
        accent_hex="E31B23", bg_hex="FFF8F8")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════════
    # BLOCO 4 — RESUMO COMERCIAL
    # ══════════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 4 | O que a LocalRise esta construindo — em linguagem simples")
    doc.add_paragraph()

    blocks_comercial = [
        ("O que e",
         "A LocalRise esta construindo a infraestrutura de inteligencia regulatoria que as consultorias "
         "de relacoes governamentais ainda nao tem. Enquanto o mercado vende relatorios, a LocalRise "
         "entrega sistemas. Enquanto outros dependem de memoria humana, a LocalRise entrega dados, "
         "alertas e analises automatizadas — todos os dias, para todos os clientes.",
         "F0F4FF", "1A56C8"),
        ("Por que e diferente de uma consultoria comum",
         "Uma consultoria comum vende horas do consultor. O que um consultor sabe, fica com ele. "
         "Quando ele sai, o conhecimento vai junto. A LocalRise transforma esse conhecimento em sistema: "
         "ele roda enquanto o consultor dorme, atende multiplos clientes ao mesmo tempo e melhora "
         "a cada novo dado que entra. E uma consultoria com escala de tecnologia.",
         "F4FFF4", "1A7A3C"),
        ("Por que e escalavel",
         "Hoje, para dobrar o numero de clientes, uma consultoria precisa dobrar a equipe. "
         "Com o sistema da LocalRise, voce pode dobrar de 10 para 20 clientes sem contratar ninguem. "
         "O custo de atender mais um cliente apos o sistema estar funcionando e proximo de zero. "
         "Isso e o que diferencia servico de produto.",
         "FFF8F0", "C86A00"),
        ("Por que o cliente deveria se interessar",
         "Porque o ambiente regulatorio esta ficando mais complexo, mais rapido e mais caro de ignorar. "
         "Empresas que nao tem inteligencia regulatoria estruturada tomam decisoes tarde, perdem "
         "oportunidades e chegam despreparadas para crises. A LocalRise entrega o sistema que "
         "coloca essas empresas sempre um passo a frente.",
         "FDF3F3", "E31B23"),
    ]

    for title, content, bg_hex, border in blocks_comercial:
        t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.cell(0, 0); bg(c, bg_hex)
        set_border(c, left={"val": "single", "sz": "16", "space": "0", "color": border})
        hp = c.paragraphs[0]; hr = hp.add_run(title + "\n"); r(hr, size=11, bold=True)
        cr2 = hp.add_run(content); r(cr2, size=10.5)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════════
    # BLOCO 5 — VERSAO PARA VENDA
    # ══════════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 5 | Versao para venda — 30 a 60 segundos")
    jp(doc,
       "Este e o material que um vendedor precisa memorizar. "
       "E o suficiente para abrir uma conversa qualificada com qualquer decisor.")
    doc.add_paragraph()

    # Elevator pitch 30s
    sub(doc, "Pitch de 30 segundos", color=RED)
    callout(doc,
        "Imagina ter, todo dia de manha, um resumo do que mudou no ambiente regulatorio do seu setor — "
        "com o que isso significa para o seu negocio e o que voce deveria fazer. "
        "Sem precisar pedir para ninguem. Sem esperar o relatorio mensal. "
        "E um painel que mostra em tempo real o seu nivel de risco regulatorio, "
        "as oportunidades disponiveis e quem esta decidindo o que em Brasilia. "
        "Isso e o que a LocalRise entrega. E o que as maiores empresas do setor ainda nao tem.",
        bg_hex="F0F4FF", border="1A56C8")

    # Pitch 60s por perfil
    sub(doc, "Pitch de 60 segundos — por perfil de decisor", color=BLUE)
    doc.add_paragraph()

    perfis = [
        ("Para o CEO",
         "Voce gasta quanto por ano em consultoria de relacoes governamentais? "
         "E quanto dessa informacao chega a voce de forma estruturada, no momento certo, "
         "sem precisar pedir? A LocalRise entrega inteligencia regulatoria do jeito que "
         "um CEO precisa: um numero de risco, alertas antes das crises e oportunidades "
         "identificadas antes do prazo. Tudo automatico. Tudo no seu painel.",
         "F0F4FF", "1A56C8"),
        ("Para o Diretor Juridico",
         "O quanto do tempo da sua equipe vai para leitura de diarios oficiais e pautas? "
         "E quantas mudancas criticas passam mesmo assim? A LocalRise automatiza esse monitoramento, "
         "filtra pelo que e relevante para o seu setor e entrega a analise de impacto em minutos — "
         "nao em dias. A sua equipe passa a atuar estrategicamente, nao operacionalmente.",
         "F4FFF4", "1A7A3C"),
        ("Para o Diretor de Relacoes Institucionais",
         "Quanto tempo voce perderia tentando montar manualmente um mapa atualizado de quem decide "
         "o que em Brasilia — e quem influencia quem? O PoliticaIQ mantém esse mapa vivo, "
         "atualizado a cada nova declaracao ou votacao. Voce age no momento certo, com a pessoa "
         "certa — e chega nas reunioes com informacao que o outro lado nao tem.",
         "FFFAF0", "C86A00"),
    ]

    for perfil_title, perfil_text, bg_hex, border in perfis:
        t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.cell(0, 0); bg(c, bg_hex)
        set_border(c, left={"val": "single", "sz": "12", "space": "0", "color": border},
                   top={"val": "single", "sz": "4", "space": "0", "color": border})
        hp = c.paragraphs[0]; hr = hp.add_run(perfil_title + "\n"); r(hr, size=11, bold=True)
        cr2 = hp.add_run(perfil_text); r(cr2, size=10.5)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════════
    # BLOCO 6 — MELHORIAS
    # ══════════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 6 | Melhorias para o documento original")
    jp(doc,
       "Este bloco e para uso interno da LocalRise. "
       "Traz o que pode ser ajustado no estudo tecnico para tornar o material mais forte comercialmente.")
    doc.add_paragraph()

    sub(doc, "O que esta confuso", color=RED)
    improvements_confuso = [
        ("NOMENCLATURA",
         "Os nomes dos SaaS (LexRadar, PoliticaIQ, etc) sao bons — mas aparecem tarde no documento. "
         "A introducao deveria apresenta-los imediatamente com uma frase de impacto antes de detalhar.",
         RED),
        ("SEPARACAO",
         "O documento mistura produtos externos (para o cliente final) com automacoes internas (para a consultoria). "
         "Isso confunde quem le pela primeira vez. A divisao precisa ser mais explicita.",
         RED),
        ("AGENTES E PIPELINE",
         "O bloco de orquestracao de agentes (Bloco 6 do original) e altamente tecnico. "
         "Para uso em reuniao com cliente, ele deve ser removido ou transformado em linguagem de negocio.",
         RED),
    ]
    for item in improvements_confuso:
        improvement_item(doc, *item)

    sub(doc, "O que pode ser melhorado", color=ORANGE)
    improvements_melhorar = [
        ("TITULOS DAS SECOES",
         "Os titulos sao descritivos, mas pouco comerciais. Exemplo: 'Bloco 1 | Ideias de SaaS para esse mercado' "
         "poderia ser 'O que voce passa a ter acesso' — mais orientado ao cliente.",
         ORANGE),
        ("PROVAS SOCIAIS",
         "O documento nao tem nenhum caso real, estimativa de ROI ou numero de validacao externa. "
         "Adicionar mesmo que um caso hipotetico detalhado aumentaria credibilidade.",
         ORANGE),
        ("CALL TO ACTION",
         "O documento termina bem, mas nao tem um CTA claro. Deveria terminar com: 'Quer ver um piloto com os seus dados?'",
         ORANGE),
        ("VISUALIZACAO DO DASHBOARD",
         "O dashboard e descrito em texto, mas seria muito mais impactante como um mockup visual — "
         "mesmo que simplificado. Uma imagem do painel valeria mais do que 5 paragrafos.",
         ORANGE),
    ]
    for item in improvements_melhorar:
        improvement_item(doc, *item)

    sub(doc, "O que pode ser removido", color=GRAY)
    improvements_remover = [
        ("BLOCO 6 DO ORIGINAL",
         "A descricao tecnica dos 6 agentes (Collector, Classifier, Analyst...) e a stack tecnologica "
         "nao tem lugar em material de reuniao com cliente. Devem ficar em documento interno.",
         GRAY),
        ("DETALHES DE INFRAESTRUTURA",
         "Next.js, pgvector, Fastify, Railway, Supabase — esses termos nao dizem nada ao decisor de negocio. "
         "Remover do material comercial. Manter apenas em documentacao tecnica.",
         GRAY),
        ("REDUNDANCIAS",
         "Alguns conceitos se repetem entre a descricao do produto e a descricao da automacao. "
         "Consolidar para evitar que o leitor perca atencao.",
         GRAY),
    ]
    for item in improvements_remover:
        improvement_item(doc, *item)

    sub(doc, "O que deve ser mais destacado para venda", color=GREEN)
    improvements_destacar = [
        ("IMPACTO EM HORAS",
         "Os numeros de tempo economizado sao os mais poderosos do documento — mas aparecem enterrados. "
         "Devem vir logo no inicio, destacados visualmente.",
         GREEN),
        ("O DIFERENCIAL DE SABER ANTES",
         "A frase 'saber antes que a crise chegue' e o maior diferencial vendavel. "
         "Deve ser o argumento central de toda apresentacao, nao apenas mencionado no CrisisMonitor.",
         GREEN),
        ("A LOGICA DE PRODUTO vs SERVICO",
         "'Enquanto o mercado vende horas de consultor, nos vendemos um sistema que nunca para.' "
         "Essa frase resume o diferencial estrategico e deve abrir o documento.",
         GREEN),
        ("WHITE-LABEL",
         "A oportunidade de white-label para outras consultorias e o maior multiplicador de receita "
         "da LocalRise — mas aparece enterrado no Bloco 5. Merece destaque proprio como oportunidade estrategica.",
         GREEN),
    ]
    for item in improvements_destacar:
        improvement_item(doc, *item)

    doc.add_paragraph()
    callout(doc,
        "Resumo das prioridades de melhoria: remover o tecniquês do material de reunião, "
        "colocar os numeros de impacto logo no inicio, adicionar um mockup do dashboard "
        "e terminar com um CTA especifico. Com essas 4 mudancas, o material passa de 'bom' para 'vendavel'.",
        bg_hex="F0F4FF", border="1A56C8")

    doc.add_page_break()
    signature(doc)

    out = BASE_DIR / "saas-inteligencia-versao-comercial.docx"
    doc.save(out)
    print(f"Documento salvo em: {out}")


if __name__ == "__main__":
    build()
