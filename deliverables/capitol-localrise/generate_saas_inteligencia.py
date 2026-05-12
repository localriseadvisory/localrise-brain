from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
LOGO_PATH = Path(
    r"C:\Users\digui\Documents\localrise-brain\Claude\Claude Local RIse\assents\logo Local Rise Principal.png"
)

RED       = RGBColor(0xE3, 0x1B, 0x23)
BLACK     = RGBColor(0x0A, 0x0A, 0x0A)
GRAY      = RGBColor(0x5F, 0x63, 0x68)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLUE      = RGBColor(0x1A, 0x56, 0xC8)
GREEN     = RGBColor(0x1A, 0x7A, 0x3C)
ORANGE    = RGBColor(0xC8, 0x6A, 0x00)
PURPLE    = RGBColor(0x5A, 0x1A, 0x8C)

DATE    = "10 de abril de 2026"
COMPANY = "Capitol Consultoria"


# ─── primitives ─────────────────────────────────────────────────────────────────

def set_border(cell, **kw):
    tc    = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tb    = tc_pr.first_child_found_in("w:tcBorders")
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tc_pr.append(tb)
    for edge, data in kw.items():
        tag = "w:%s" % edge
        el  = tb.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); tb.append(el)
        for k, v in data.items():
            el.set(qn("w:%s" % k), str(v))


def set_bg(cell, hex_color: str):
    tc    = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def apply_styles(doc: Document):
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.top_margin    = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin   = Cm(2.4)
    s.right_margin  = Cm(2.4)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    n.font.size = Pt(11)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    n.paragraph_format.space_after       = Pt(6)


def fnt(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color: run.font.color.rgb = color


def p(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=11, bold=False,
      color=None, italic=False, sb=0, sa=6):
    par = doc.add_paragraph()
    par.alignment                         = align
    par.paragraph_format.space_before     = Pt(sb)
    par.paragraph_format.space_after      = Pt(sa)
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if text:
        r = par.add_run(text)
        fnt(r, size=size, bold=bold, color=color, italic=italic)
    return par


def jp(doc, text, size=11, bold=False, color=None, italic=False, sa=8):
    par = doc.add_paragraph()
    par.alignment                         = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.space_after      = Pt(sa)
    r = par.add_run(text)
    fnt(r, size=size, bold=bold, color=color, italic=italic)


def blt(doc, text, size=11):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.space_after      = Pt(3)
    r = par.add_run(text)
    fnt(r, size=size)


def nblt(doc, text, size=11):
    par = doc.add_paragraph(style="List Number")
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.space_after      = Pt(3)
    r = par.add_run(text)
    fnt(r, size=size)


def sec(doc, text, color=RED):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(18)
    par.paragraph_format.space_after  = Pt(4)
    r = par.add_run(text.upper())
    fnt(r, size=13, bold=True, color=color)


def sub(doc, text, color=BLACK):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(10)
    par.paragraph_format.space_after  = Pt(2)
    r = par.add_run(text)
    fnt(r, size=11.5, bold=True, color=color)


def divider(doc, color="E0E0E0"):
    t    = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_border(cell, bottom={"val": "single", "sz": 6, "space": "0", "color": color})
    cell.text = ""
    doc.add_paragraph()


def callout(doc, text, bg="F4F4F4", border="E31B23", size=11):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_bg(c, bg)
    set_border(c, left={"val": "single", "sz": 18, "space": "0", "color": border})
    par = c.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = par.add_run(text)
    fnt(r, size=size, italic=True)
    doc.add_paragraph()


# ─── composites ─────────────────────────────────────────────────────────────────

def brand_header(doc):
    if LOGO_PATH.exists():
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(LOGO_PATH), width=Cm(5.4))
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("LOCALRISE ADVISORY")
    fnt(r, size=12, bold=True, color=RED)
    divider(doc)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(8)
    r = par.add_run("Inteligencia Regulatoria como Produto")
    fnt(r, size=20, bold=True)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("SaaS, dashboards, automacoes e produtizacao para consultorias de relacoes governamentais")
    fnt(r, size=11, color=GRAY)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(DATE)
    fnt(r, size=10, color=GRAY, italic=True)


def saas_card(doc, number, name, tagline, problem, audience, features, differential, pricing, accent_hex):
    # cabecalho
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    nc, lc, tc_ = t.rows[0].cells
    nc.width = Cm(1.2);  lc.width = Cm(3.8);  tc_.width = Cm(11.2)
    set_bg(nc, accent_hex)
    set_bg(lc, "1A1A1A")
    np2 = nc.paragraphs[0]; np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = np2.add_run(number); fnt(r, size=13, bold=True, color=WHITE)
    lp = lc.paragraphs[0]
    r2 = lp.add_run(name); fnt(r2, size=11, bold=True, color=WHITE)
    lp2 = lc.add_paragraph(); r3 = lp2.add_run(tagline); fnt(r3, size=8.5, italic=True, color=RGBColor(0xCC,0xCC,0xCC))
    tp = tc_.paragraphs[0]; r4 = tp.add_run(problem); fnt(r4, size=10.5, italic=True, color=GRAY)
    doc.add_paragraph()

    rows = [
        ("Publico-alvo",  audience),
        ("Funcionalidades principais", features),
        ("Diferencial",   differential),
        ("Cobranca",      pricing),
    ]
    for label, content in rows:
        t2 = doc.add_table(rows=1, cols=2)
        t2.alignment = WD_TABLE_ALIGNMENT.LEFT
        lc2, vc = t2.rows[0].cells
        lc2.width = Cm(4.2); vc.width = Cm(12.0)
        set_bg(lc2, "F4F4F4")
        set_border(lc2, bottom={"val": "single","sz":"4","space":"0","color":"E0E0E0"})
        set_border(vc,  bottom={"val": "single","sz":"4","space":"0","color":"E0E0E0"})
        lpar = lc2.paragraphs[0]; lr = lpar.add_run(label); fnt(lr, size=9.5, bold=True, color=GRAY)
        vpar = vc.paragraphs[0];  vr = vpar.add_run(content); fnt(vr, size=10.5)
    doc.add_paragraph()


def dashboard_widget(doc, title, desc, bg="F9F9F9", accent="E31B23"):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    ac, dc = t.rows[0].cells
    ac.width = Cm(0.5); dc.width = Cm(15.7)
    set_bg(ac, accent)
    set_bg(dc, bg)
    dp = dc.paragraphs[0]
    r1 = dp.add_run(title + "\n"); fnt(r1, size=10.5, bold=True)
    r2 = dp.add_run(desc);         fnt(r2, size=10, color=GRAY)
    doc.add_paragraph()


def automation_card(doc, name, what, tools, impact, accent_hex="E31B23"):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    nc, dc = t.rows[0].cells
    nc.width = Cm(0.5); dc.width = Cm(15.7)
    set_bg(nc, accent_hex)
    par = dc.add_paragraph() if dc.paragraphs[0].text == "" else dc.paragraphs[0]
    r1 = par.add_run(name + "\n"); fnt(r1, size=11, bold=True)
    doc.add_paragraph()

    blocks = [("O que faz", what), ("Ferramentas", tools), ("Impacto", impact)]
    for lbl, val in blocks:
        t2 = doc.add_table(rows=1, cols=2)
        t2.alignment = WD_TABLE_ALIGNMENT.LEFT
        lc, vc = t2.rows[0].cells
        lc.width = Cm(3.8); vc.width = Cm(12.4)
        set_bg(lc, "F4F4F4")
        set_border(lc, bottom={"val":"single","sz":"4","space":"0","color":"E0E0E0"})
        set_border(vc, bottom={"val":"single","sz":"4","space":"0","color":"E0E0E0"})
        lp = lc.paragraphs[0]; lr = lp.add_run(lbl); fnt(lr, size=9.5, bold=True, color=GRAY)
        vp = vc.paragraphs[0]; vr = vp.add_run(val); fnt(vr, size=10.5)
    doc.add_paragraph()


def package_card(doc, tier, subtitle, bg, accent, items_auto, items_human, price):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_bg(c, bg)
    set_border(c,
               top={"val":"single","sz":"12","space":"0","color": accent},
               left={"val":"single","sz":"4","space":"0","color": accent})
    hpar = c.paragraphs[0]
    r1 = hpar.add_run(tier + " | "); fnt(r1, size=13, bold=True, color=RGBColor(int(accent[0:2],16),int(accent[2:4],16),int(accent[4:6],16)))
    r2 = hpar.add_run(subtitle); fnt(r2, size=11, bold=True)
    doc.add_paragraph()

    sub(doc, "O que e automatizado")
    for i in items_auto: blt(doc, i)
    sub(doc, "O que e consultoria humana")
    for i in items_human: blt(doc, i)
    callout(doc, "Cobranca: " + price, bg="F7F7F7", border=accent)


def agent_row(doc, name, role, input_from, output_to, accent_hex):
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    ac, nc, ic, oc = t.rows[0].cells
    ac.width = Cm(0.5); nc.width = Cm(4.2); ic.width = Cm(5.5); oc.width = Cm(6.0)
    set_bg(ac, accent_hex)
    set_bg(nc, "1A1A1A")
    set_border(ic, bottom={"val":"single","sz":"4","space":"0","color":"E0E0E0"})
    set_border(oc, bottom={"val":"single","sz":"4","space":"0","color":"E0E0E0"})
    np2 = nc.paragraphs[0]
    nr = np2.add_run(name + "\n"); fnt(nr, size=9.5, bold=True, color=WHITE)
    rr = np2.add_run(role); fnt(rr, size=8, italic=True, color=RGBColor(0xAA,0xAA,0xAA))
    ip = ic.paragraphs[0]; ir = ip.add_run("Recebe: " + input_from); fnt(ir, size=9.5, color=GRAY)
    op = oc.paragraphs[0]; or_ = op.add_run("Entrega: " + output_to); fnt(or_, size=9.5)
    doc.add_paragraph()


def signature(doc):
    par = doc.add_paragraph()
    r = par.add_run(f"Sao Paulo, {DATE}.")
    fnt(r, size=11, bold=True)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for cell in row.cells:
            cell.width = Cm(7.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            set_border(cell, top={"val":"single","sz":"8","space":"0","color":"AEB4BB"})
    lp = t.cell(0,0).paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run("LOCALRISE ADVISORY\n").bold = True
    lp.add_run("Nicolas Slater\n"); lp.add_run("Diretoria Comercial")
    rp = t.cell(0,1).paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.add_run(f"{COMPANY}\n").bold = True
    rp.add_run("Responsavel: ______________________\n")
    rp.add_run("Cargo: ___________________________")


# ─── documento ──────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    apply_styles(doc)
    brand_header(doc)
    doc.add_page_break()

    # ── INTRO ────────────────────────────────────────────────────────────────────
    sec(doc, "Visao geral")
    jp(doc,
       "Este estudo analisa como uma consultoria de relacoes governamentais pode deixar de vender apenas "
       "servico e comecar a vender sistema. O objetivo e identificar onde o conhecimento humano pode virar "
       "produto digital, qual tecnologia viabiliza essa transicao e como a LocalRise pode construir ou "
       "licenciar esse produto para o mercado.")
    jp(doc,
       "O documento esta estruturado em seis blocos: ideias de SaaS, design do dashboard estrategico, "
       "automacoes operacionais, produtizacao do servico, ideias avancadas de diferencial e a arquitetura "
       "de orquestracao de agentes que faria tudo isso funcionar.")
    callout(doc,
        "A premissa central: conhecimento politico e regulatorio tem altissimo valor — mas hoje ele fica "
        "dentro da cabeca dos consultores. O SaaS e o que transforma esse conhecimento em ativo escalavel.",
        bg="F0F4FF", border="1A56C8")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BLOCO 1 — SAAS
    # ═══════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 1 | Ideias de SaaS para esse mercado")
    jp(doc,
       "Cada produto abaixo resolve uma dor especifica do cliente de uma consultoria de relacoes "
       "governamentais. Todos podem ser vendidos como SaaS independente ou como camada de valor "
       "adicional ao contrato de consultoria — o que aumenta ticket, retencao e previsibilidade de receita.")
    doc.add_paragraph()

    saas_card(doc, "01", "LexRadar", "Inteligencia legislativa em tempo real",
        problem="Decisores nao sabem o que esta mudando no ambiente regulatorio ate que ja seja tarde. "
                "Monitorar manualmente diarios oficiais, pautas de comissoes e projetos de lei e caro e ineficiente.",
        audience="Diretorias juridicas, de compliance e de relacoes governamentais em empresas de setores regulados "
                 "(energia, saude, telecomunicacoes, financeiro, agronegocio)",
        features="Monitoramento em tempo real de projetos de lei, decretos e normas | Filtro por setor e palavra-chave | "
                 "Resumo automatico via IA de cada novo ato normativo | Alerta imediato por e-mail e WhatsApp | "
                 "Historico de alteracoes por tema | Exportacao de relatorios para C-suite",
        differential="A maioria dos servicos entrega o dado bruto. O LexRadar entrega o dado ja interpretado: "
                     "o que mudou, qual o impacto para o setor do cliente e qual e a urgencia de acao. "
                     "IA + curadoria humana da consultoria.",
        pricing="Mensalidade por numero de temas monitorados: R$1.800 (ate 5 temas) / R$3.500 (ate 15 temas) / "
                "R$6.500 (ilimitado + alertas prioritarios). Setup unico de R$2.000.",
        accent_hex="E31B23")

    saas_card(doc, "02", "PoliticaIQ", "Mapa de influencia politica e poder",
        problem="Empresas precisam saber quem decide, quem influencia quem e qual e o timing certo para agir. "
                "Esse mapa esta na cabeca do consultor — e desaparece quando o contrato termina.",
        audience="Diretores de relacoes institucionais, CEOs de empresas com agenda regulatoria, "
                 "assessorias juridicas de grandes grupos",
        features="Perfis de decisores-chave por tema e por orgao | Mapa de relacionamentos e influencia | "
                 "Historico de posicionamento e votacoes | Timeline de mandatos e sucessoes | "
                 "Analise de risco politico por agenda | Integracao com calendario legislativo",
        differential="Nao e uma base de dados politicos generica. E um mapa de influencia construido "
                     "para o contexto especifico do cliente — filtrado pelo setor, pela agenda e pelo timing. "
                     "Funciona como um GPS de Brasilia.",
        pricing="Mensalidade por numero de perfis ativos: R$2.500 (50 perfis) / R$5.000 (200 perfis) / "
                "R$9.500 (ilimitado + atualizacoes diarias). Versao white-label disponivel para consultorias parceiras.",
        accent_hex="1A56C8")

    saas_card(doc, "03", "RegRisk", "Gestao de risco regulatorio continuo",
        problem="Compliance e risco regulatorio sao tratados de forma reativa — a empresa so age quando a crise "
                "ja chegou. Nao existe uma visao prospectiva e continua do risco.",
        audience="Diretorias de compliance, risco corporativo e juridico em empresas de medio e grande porte "
                 "com exposicao regulatoria significativa",
        features="Score de risco regulatorio atualizado em tempo real | Matriz de exposicao por tema e por orgao "
                 "| Simulacao de cenarios para propostas em tramitacao | Historico de risco ao longo do tempo | "
                 "Relatorio executivo semanal automatico | Integracao com gestao de crise",
        differential="Enquanto outros produtos monitoram o que ja aconteceu, o RegRisk projeta o que pode "
                     "acontecer. O score de risco combina dados publicos com a interpretacao da consultoria — "
                     "gerando uma nota diaria de exposicao para cada cliente.",
        pricing="Mensalidade por empresa: R$3.500 (basico) / R$7.000 (avancado com simulacoes) / "
                "R$14.000 (premium com consultoria humana integrada). Trial de 30 dias.",
        accent_hex="1A7A3C")

    saas_card(doc, "04", "OpportunityMap", "Radar de oportunidades governamentais",
        problem="Incentivos fiscais, linhas de financiamento publico, programas de fomento e editais existem em "
                "volume imenso — mas chegam ao conhecimento das empresas por acaso ou ja fora do prazo.",
        audience="Diretorias financeiras, de inovacao e de estrategia de empresas de todos os setores, "
                 "especialmente startups, industrias e agro",
        features="Monitoramento de editais, programas e financiamentos publicos | Matching automatico entre "
                 "perfil da empresa e oportunidades disponiveis | Alerta de prazo e janela de inscricao | "
                 "Calculadora de beneficio estimado | Historico de programas por setor | Guia de habilitacao",
        differential="Nao e uma lista de editais — e um sistema de matching inteligente. A empresa cadastra "
                     "seu perfil uma vez e o sistema cruza continuamente com as oportunidades disponiveis, "
                     "priorizando as mais viaveis.",
        pricing="Mensalidade: R$1.200 (basico - alertas) / R$2.800 (com matching e calculadora) / "
                "R$5.500 (premium com suporte de habilitacao). Comissao opcional sobre beneficio captado.",
        accent_hex="C86A00")

    saas_card(doc, "05", "CrisisMonitor", "Alerta precoce de crises politicas e regulatorias",
        problem="Crises regulatorias e politicas que afetam o negocio nao aparecem do nada — elas tem sinais "
                "precursores que passam despercebidos por falta de monitoramento estruturado.",
        audience="C-suite e diretorias de comunicacao, juridico e relacoes institucionais de empresas "
                 "com alta exposicao politica ou regulatoria",
        features="Monitoramento de midia, redes sociais, diario oficial e pautas legislativas em paralelo | "
                 "Deteccao de sinais de crise por IA | Score de probabilidade de impacto | "
                 "Fluxo de resposta recomendado | Historico de crises setoriais | Relatorio de posicionamento "
                 "de stakeholders estrategicos",
        differential="A maioria das ferramentas de monitoramento e reativa. O CrisisMonitor trabalha com "
                     "deteccao de padroes — cruzando sinais fracos de diferentes fontes para gerar alertas "
                     "antes que a crise chegue ao mainstream.",
        pricing="Mensalidade: R$4.500 (monitoramento basico) / R$9.000 (com IA preditiva) / "
                "R$18.000 (com gestao de crise integrada e consultoria on-call).",
        accent_hex="5A1A8C")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BLOCO 2 — DASHBOARD
    # ═══════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 2 | Dashboard estrategico — o painel de inteligencia politica")
    jp(doc,
       "O dashboard e o produto mais visivel para o cliente da consultoria. "
       "E o que transforma a relacao de 'eu pago por relatorio' para 'eu tenho um painel que me mostra "
       "o que esta acontecendo no meu ambiente regulatorio agora mesmo'. "
       "Abaixo esta o design conceitual do painel ideal.")
    doc.add_paragraph()

    sub(doc, "Organizacao geral do painel")
    jp(doc,
       "O painel e dividido em quatro visoes principais — acessadas por abas ou por cards no topo. "
       "Cada visao responde uma pergunta diferente que o decisor faz todo dia.")

    visoes = [
        ("Visao RISCO",       "O que esta ameacando meu negocio agora?",      "E31B23"),
        ("Visao OPORTUNIDADE","O que posso aproveitar antes dos concorrentes?","1A7A3C"),
        ("Visao INFLUENCIA",  "Quem decide e o que esta pensando?",           "1A56C8"),
        ("Visao AGENDA",      "O que vai acontecer nos proximos 30, 60, 90 dias?","C86A00"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, question, color) in enumerate(visoes):
        c = t.rows[0].cells[idx]
        c.width = Cm(3.9)
        set_bg(c, color)
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = cp.add_run(label + "\n"); fnt(r1, size=10, bold=True, color=WHITE)
        r2 = cp.add_run(question); fnt(r2, size=8.5, italic=True, color=RGBColor(0xEE,0xEE,0xEE))
    doc.add_paragraph()

    sub(doc, "Widgets e modulos por visao")

    dashboard_widget(doc, "Score de Risco Regulatorio (0-100)",
        "Numero unico que resume a exposicao atual da empresa ao ambiente regulatorio. "
        "Calculado diariamente com base em: volume de atos normativos monitorados, proximidade de votacoes criticas, "
        "posicionamento dos decisores-chave e historico de impacto setorial. Cor: verde (0-30), amarelo (31-60), vermelho (61-100).",
        bg="FFF8F8", accent="E31B23")

    dashboard_widget(doc, "Radar Legislativo — Proximos 30 dias",
        "Timeline visual com todos os projetos de lei, votacoes e prazos regulatorios relevantes para o setor do cliente. "
        "Cada item tem: status atual, probabilidade de aprovacao, impacto estimado e acao recomendada. "
        "Cores indicam urgencia. Clicavel para ver detalhes completos.",
        bg="F8F8FF", accent="1A56C8")

    dashboard_widget(doc, "Mapa de Stakeholders Ativos",
        "Visualizacao em rede mostrando os decisores relevantes para as agendas do cliente. "
        "Cada no representa um decisor — tamanho proporcional a influencia, cor proporcional ao posicionamento "
        "(favoravel, neutro, desfavoravel). Clicavel para ver perfil, historico e ultimo contato registrado.",
        bg="F8FFF9", accent="1A7A3C")

    dashboard_widget(doc, "Oportunidades Abertas",
        "Cards com as principais oportunidades ativas: incentivos fiscais, financiamentos, programas e editais. "
        "Cada card mostra: nome, beneficio estimado, prazo, status de habilitacao da empresa e proxima acao. "
        "Ordenado por potencial de retorno.",
        bg="FFFAF0", accent="C86A00")

    dashboard_widget(doc, "Feed de Inteligencia — Ultimas 48h",
        "Stream cronologico de eventos monitorados: novas publicacoes no diario oficial, movimentacoes em comissoes, "
        "declaracoes de decisores, mudancas de posicionamento, alertas de midia. "
        "Filtrado por relevancia para o setor. Cada item com tag de impacto (alto / medio / baixo).",
        bg="F8F8F8", accent="5A1A8C")

    dashboard_widget(doc, "KPIs do Relacionamento Institucional",
        "Metricas de resultado da propria consultoria: numero de reunioes realizadas no mes, "
        "agendas acompanhadas, alertas emitidos, oportunidades identificadas, crises evitadas. "
        "Funciona como relatorio de valor entregue — justifica o contrato mensalmente.",
        bg="F0F4FF", accent="1A56C8")

    callout(doc,
        "O dashboard transforma a consultoria de um fornecedor de relatorios em um parceiro estrategico permanente. "
        "O cliente nao espera o relatorio mensal — ele acessa o painel e toma decisao em tempo real.",
        bg="F0F4FF", border="1A56C8")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BLOCO 3 — AUTOMACOES
    # ═══════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 3 | Automacoes que geram eficiencia operacional")
    jp(doc,
       "Cada automacao abaixo pode ser implementada com n8n como orquestrador, APIs publicas como fonte de dados "
       "e LLMs (Claude, GPT-4) como camada de interpretacao. "
       "O objetivo e claro: o que hoje leva horas de trabalho humano passa a ser feito em minutos — "
       "com qualidade superior e custo marginal proximo de zero.")
    doc.add_paragraph()

    automation_card(doc,
        "Monitoramento legislativo automatico e continuo",
        "Scraping diario das APIs da Camara dos Deputados, Senado Federal e Diario Oficial da Uniao. "
        "Filtragem por palavras-chave e setor. LLM gera resumo de cada novo ato em 3 paragrafos: "
        "o que mudou, quem e afetado e qual e o prazo de impacto. Resultado armazenado no banco de dados do painel.",
        "n8n (orquestrador) + APIs: dados.camara.leg.br, legis.senado.leg.br, in.gov.br + Claude API (resumos) "
        "+ banco de dados (Supabase ou PostgreSQL)",
        "Elimina 15 a 25 horas/semana de leitura manual por analista. Permite monitorar dezenas de temas "
        "simultaneamente — algo humanamente inviavel. Qualidade de resumo consistente e rastreavel.",
        accent_hex="E31B23")

    automation_card(doc,
        "Alerta inteligente personalizado por cliente",
        "Com base no perfil cadastrado de cada cliente (setor, temas de interesse, nivel de risco), "
        "o sistema filtra os eventos do dia e envia apenas o que e relevante para cada empresa. "
        "O alerta inclui: evento, impacto estimado e acao recomendada. Enviado via e-mail e WhatsApp.",
        "n8n + CRM (perfil do cliente) + Claude API (personalizacao e contextualizacao) "
        "+ WhatsApp Business API + Gmail",
        "Transforma a consultoria de generalista para hiperpersonalizada — cada cliente recebe "
        "inteligencia filtrada para o seu contexto. Aumenta o valor percebido e reduz churn.",
        accent_hex="1A56C8")

    automation_card(doc,
        "Geracao automatica de relatorio semanal",
        "Toda sexta-feira, o sistema coleta os eventos da semana, filtra por cliente, "
        "gera um relatorio narrativo via LLM (com linguagem executiva, nao tecnica), "
        "formata em PDF com a identidade visual do cliente e envia por e-mail. "
        "O consultor revisa apenas casos de alta criticidade.",
        "n8n + banco de dados de eventos + Claude API (geracao do texto narrativo) "
        "+ Python-docx ou Puppeteer (formatacao PDF) + Gmail",
        "Reduz de 4 horas para 15 minutos o tempo de producao de cada relatorio semanal. "
        "Com 20 clientes ativos, isso representa 75 horas/mes recuperadas para trabalho estrategico.",
        accent_hex="1A7A3C")

    automation_card(doc,
        "Analise automatica de impacto regulatorio",
        "Quando um novo projeto de lei ou normativa e identificado, o sistema envia o texto completo para "
        "um LLM com instrucoes especificas: identificar impacto por setor, estimar timeline, mapear "
        "quais clientes sao afetados e sugerir posicionamento. O analista recebe um rascunho de analise "
        "pronto para revisao e envio.",
        "n8n + APIs legislativas + Claude API (analise de impacto com prompt especializado) "
        "+ CRM (mapeamento de clientes afetados) + e-mail",
        "Uma analise que levava 3 horas e entregue em 8 minutos. "
        "O analista passa de produtor de texto para revisor estrategico — dobrando a capacidade de atendimento.",
        accent_hex="C86A00")

    automation_card(doc,
        "Monitoramento de posicionamento de decisores",
        "Scraping de declaracoes publicas, votacoes, entrevistas e notas oficiais de decisores politicos "
        "mapeados. LLM analisa o posicionamento de cada decisor em relacao aos temas monitorados e "
        "atualiza o mapa de influencia do painel automaticamente. Alerta quando ha mudanca de posicionamento.",
        "n8n + scraping de midia (NewsAPI, Bing News API) + APIs legislativas + "
        "Claude API (analise de sentimento e posicionamento) + banco de dados",
        "Manter um mapa de influencia atualizado manualmente exigiria uma equipe dedicada. "
        "A automacao faz isso de forma continua, para centenas de decisores simultaneamente.",
        accent_hex="5A1A8C")

    automation_card(doc,
        "Qualificacao e onboarding automatico de novo cliente",
        "Quando um novo cliente entra, preenche um formulario de perfil (setor, temas, orgaos, nivel de risco). "
        "O sistema configura automaticamente: quais fontes monitorar, quais palavras-chave ativar, quais "
        "decisores incluir no mapa, qual template de relatorio usar e qual frequencia de alerta. "
        "O painel do cliente e configurado sem intervencao manual.",
        "n8n + formulario (Typeform ou nativo) + CRM + banco de dados + Claude API (geracao "
        "de keywords e perfil de monitoramento a partir do texto livre do cliente)",
        "Elimina de 4 a 8 horas de configuracao manual por novo cliente. "
        "Permite onboarding de multiplos clientes em paralelo sem sobrecarregar a equipe.",
        accent_hex="E31B23")

    automation_card(doc,
        "Deteccao precoce de crise regulatoria",
        "Monitoramento cruzado de 5 fontes simultaneamente: diario oficial, pauta legislativa, "
        "midia especializada, redes sociais de decisores e historico de crises setoriais. "
        "LLM identifica padroes de crise emergente e calcula probabilidade de impacto. "
        "Alerta de alta prioridade enviado ao cliente e ao consultor responsavel antes da crise estourar.",
        "n8n + multiplas APIs de midia e dados publicos + Claude API (analise de padroes e scoring) "
        "+ banco de dados historico + WhatsApp + e-mail urgente",
        "O diferencial mais vendavel da consultoria: saber antes. "
        "Um unico alerta antecipado que permite reposicionamento estrategico pode valer centenas de vezes "
        "o valor do contrato anual.",
        accent_hex="1A56C8")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BLOCO 4 — PRODUTIZACAO
    # ═══════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 4 | Produtizacao — de servico para produto escalavel")
    jp(doc,
       "A produtizacao nao elimina a consultoria humana — ela libera o consultor para o que realmente "
       "exige senioridade. Abaixo esta como estruturar tres niveis de produto, onde cada camada "
       "automatiza mais e entrega mais valor com menos custo marginal.")
    doc.add_paragraph()

    package_card(doc, "ESSENCIAL", "Monitoramento e Alertas",
        bg="F0F8FF", accent="1A56C8",
        items_auto=[
            "Monitoramento diario de legislacao por palavras-chave e setor",
            "Resumos automaticos de novos atos normativos",
            "Alertas por e-mail e WhatsApp com impacto estimado",
            "Relatorio semanal automatico em PDF",
            "Acesso ao painel basico (feed de inteligencia + radar legislativo)",
        ],
        items_human=[
            "Curadoria mensal de temas com consultor designado (1h/mes)",
            "Revisao de alertas de alta criticidade antes do envio",
            "Suporte por e-mail em ate 48h para duvidas pontuais",
        ],
        price="R$2.500 a R$4.500/mes dependendo do numero de temas. Setup de R$1.500.")

    package_card(doc, "ESTRATEGICO", "Inteligencia e Analise",
        bg="F0FFF4", accent="1A7A3C",
        items_auto=[
            "Tudo do plano Essencial",
            "Mapa de stakeholders com atualizacao automatica",
            "Analise automatica de impacto de novos projetos de lei",
            "Score de risco regulatorio diario",
            "Relatorio executivo mensal gerado por IA e revisado pelo consultor",
            "Acesso ao painel completo (4 visoes: risco, oportunidade, influencia, agenda)",
        ],
        items_human=[
            "Reuniao estrategica mensal com consultor senior (2h)",
            "Analise aprofundada de 2 temas criticos por mes",
            "Representacao junto a orgaos em situacoes de alta prioridade",
        ],
        price="R$7.500 a R$12.000/mes. Contrato minimo de 6 meses.")

    package_card(doc, "PREMIUM", "Parceria Estrategica Continua",
        bg="FFF8F0", accent="C86A00",
        items_auto=[
            "Tudo dos planos anteriores",
            "CrisisMonitor ativo com alertas de alta prioridade 24/7",
            "OpportunityMap com matching automatico de incentivos e financiamentos",
            "Relatorios customizados para board e C-suite",
            "Integracao com sistemas internos do cliente (via API)",
        ],
        items_human=[
            "Consultor senior dedicado com disponibilidade prioritaria",
            "Participacao em reunioes internas do cliente quando necessario",
            "Gestao ativa de relacionamento com decisores estrategicos",
            "Posicionamento institucional em crises",
            "Acesso a rede de relacionamentos da consultoria",
        ],
        price="R$18.000 a R$35.000/mes. Contrato anual. Possibilidade de equity em casos especificos.")

    callout(doc,
        "A logica da produtizacao e simples: o que pode ser automatizado deve ser automatizado. "
        "O que exige julgamento humano, relacao e reputacao — esse e o nucleo insubstituivel da consultoria. "
        "O produto escala o primeiro. O consultor aprofunda o segundo.",
        bg="F0FFF4", border="1A7A3C")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BLOCO 5 — IDEIAS AVANCADAS
    # ═══════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 5 | Ideias avancadas — o que ninguem esta fazendo ainda")
    jp(doc,
       "Este bloco traz as ideias que constroem o diferencial de longo prazo. "
       "Nao sao features adicionais — sao apostas estrategicas que, se construidas antes dos concorrentes, "
       "criam barreiras de entrada significativas.")
    doc.add_paragraph()

    advanced = [
        ("Scoring preditivo de risco regulatorio",
         "Em vez de mostrar o que ja aconteceu, o sistema usa historico legislativo, analise de mandatos, "
         "movimentacoes em comissoes e padroes de crise para calcular a probabilidade de um determinado "
         "evento regulatorio ocorrer nos proximos 30, 60 e 90 dias. "
         "O cliente recebe um numero — como uma previsao do tempo, mas para o ambiente politico. "
         "Base tecnica: modelo de ML treinado com historico de projetos de lei + LLM para interpretacao contextual.",
         "1A7A3C"),
        ("Mapa de influencia em rede (Network Intelligence)",
         "Visualizacao interativa das conexoes entre decisores, empresas, entidades setoriais, partidos e "
         "cargos-chave. O mapa mostra caminhos de influencia — como uma empresa pode chegar a um decisor "
         "especifico em menos passos, passando por quem. "
         "Atualizado automaticamente com dados publicos: votacoes, declaracoes, nomeacoes, financiamentos eleitorais. "
         "Base tecnica: grafo de nos e arestas (Neo4j ou similar) + atualizacao automatica via scraping.",
         "1A56C8"),
        ("Simulador de cenarios regulatorios",
         "Ferramenta interativa onde o cliente descreve um cenario hipotetico ('e se esse projeto de lei "
         "for aprovado no formato atual?') e o sistema projeta os impactos: financeiro, operacional, "
         "competitivo e reputacional. "
         "Baseado em analogias com casos historicos similares + modelagem de impacto por LLM treinado "
         "com dados setoriais.",
         "5A1A8C"),
        ("Indice de sentimento politico setorial",
         "Numero publicado mensalmente (ou em tempo real) que mede o sentimento do ambiente politico "
         "em relacao a um determinado setor. "
         "Calculado a partir de: declaracoes de parlamentares, cobertura de midia, movimentacao de projetos, "
         "votacoes recentes. Comparavel a um indice de bolsa — mas para o clima politico. "
         "Produto de dados que pode ser publicado, gerando autoridade de mercado para a consultoria.",
         "C86A00"),
        ("Assistente de inteligencia regulatoria por LLM",
         "Um assistente de IA treinado com o historico regulatorio do setor do cliente, a base de dados "
         "legislativa e o proprio historico de analises da consultoria. O usuario faz perguntas em linguagem "
         "natural: 'Qual e o risco de o projeto X ser aprovado antes do recesso?' ou 'Quais os precedentes "
         "para esse tipo de decreto?' e recebe resposta contextualizada em segundos. "
         "Base tecnica: RAG (retrieval augmented generation) sobre base proprietaria de documentos + Claude API.",
         "E31B23"),
        ("White-label para outras consultorias",
         "Toda essa infraestrutura — LexRadar, dashboard, automacoes — pode ser licenciada como white-label "
         "para consultorias regionais, escritorios de advocacia tributaria ou associacoes setoriais. "
         "A LocalRise se torna a infraestrutura tecnologica do mercado de inteligencia regulatoria no Brasil. "
         "Modelo: royalty por seat ou percentual de receita da consultoria parceira.",
         "1A56C8"),
    ]

    for idea_title, idea_desc, color in advanced:
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        ac, dc = t.rows[0].cells
        ac.width = Cm(0.5); dc.width = Cm(15.7)
        set_bg(ac, color)
        dp = dc.paragraphs[0]
        r1 = dp.add_run(idea_title + "\n"); fnt(r1, size=11, bold=True)
        r2 = dp.add_run(idea_desc); fnt(r2, size=10.5, color=GRAY)
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BLOCO 6 — ORQUESTRACAO DE AGENTES
    # ═══════════════════════════════════════════════════════════════════════════
    sec(doc, "Bloco 6 | Arquitetura de orquestracao — o sistema por dentro")
    jp(doc,
       "Este bloco descreve como os agentes, skills e comandos se orquestram para fazer o sistema "
       "funcionar de forma autonoma. Pense como um pipeline de inteligencia: cada agente tem uma funcao "
       "especifica, recebe um input e entrega um output que o proximo agente usa.")
    doc.add_paragraph()

    sub(doc, "Os 6 agentes do sistema")
    jp(doc,
       "O sistema opera com 6 agentes especializados, cada um responsavel por uma camada do pipeline. "
       "Eles rodam em paralelo onde possivel e em sequencia onde ha dependencia de dados.")
    doc.add_paragraph()

    agent_row(doc, "AGENT COLLECTOR", "Coleta de dados brutos",
        "APIs publicas (Camara, Senado, DOU, NewsAPI), scraping de fontes especificas",
        "Dados estruturados em JSON, categorizados por fonte e data",
        accent_hex="1A56C8")

    agent_row(doc, "AGENT CLASSIFIER", "Classificacao e filtragem",
        "JSON bruto do Agent Collector + perfil de clientes do CRM",
        "Eventos classificados por relevancia, setor e nivel de urgencia para cada cliente",
        accent_hex="1A7A3C")

    agent_row(doc, "AGENT ANALYST", "Analise e interpretacao",
        "Eventos classificados + historico de analises anteriores + base de contexto setorial",
        "Analise textual: o que aconteceu, quem e afetado, qual o impacto, qual a acao recomendada",
        accent_hex="C86A00")

    agent_row(doc, "AGENT RISK SCORER", "Calculo de scores e metricas",
        "Analises do Agent Analyst + historico de risco + mapa de stakeholders",
        "Score de risco atualizado, variacao em relacao ao dia anterior, flag de alertas criticos",
        accent_hex="E31B23")

    agent_row(doc, "AGENT REPORTER", "Geracao de conteudo e relatorios",
        "Scores + analises + template do cliente + instrucoes de formato",
        "Relatorio em linguagem executiva, formatado para PDF ou mensagem, pronto para envio",
        accent_hex="5A1A8C")

    agent_row(doc, "AGENT DISPATCHER", "Distribuicao e notificacao",
        "Conteudo do Agent Reporter + configuracoes de entrega do cliente (horario, canal, frequencia)",
        "Mensagens enviadas por WhatsApp, e-mail ou painel. Log de entrega registrado.",
        accent_hex="0A0A0A")

    doc.add_paragraph()
    sub(doc, "Fluxo do pipeline — como os agentes se encadeiam")

    pipeline_steps = [
        "06:00 — Agent Collector executa: coleta dados das APIs publicas e scraping configurado",
        "06:15 — Agent Classifier filtra e categoriza por cliente e nivel de urgencia",
        "06:30 — Agent Analyst processa eventos criticos (score alto) e gera analises prioritarias",
        "07:00 — Agent Risk Scorer atualiza o score diario de cada cliente no banco de dados",
        "07:15 — Agent Reporter gera relatorio diario e alertas personalizados por cliente",
        "07:30 — Agent Dispatcher envia alertas via WhatsApp/e-mail e atualiza o painel em tempo real",
        "Continuo — Agent Collector roda a cada 2 horas para eventos urgentes (modo de vigilancia)",
        "Sexta 16:00 — Pipeline especial: relatorio semanal completo para todos os clientes",
    ]
    for step in pipeline_steps:
        blt(doc, step)

    doc.add_paragraph()
    sub(doc, "Stack tecnologica recomendada")

    stack = [
        ("Frontend (painel)",     "Next.js + TypeScript + Tailwind CSS + Recharts (grafos e visualizacoes)"),
        ("Backend / API",         "Node.js (Fastify) ou Python (FastAPI) — REST + WebSockets para atualizacoes em tempo real"),
        ("Banco de dados",        "PostgreSQL (dados estruturados) + pgvector (embeddings para busca semantica)"),
        ("Orquestracao (agentes)","n8n self-hosted como orquestrador central de workflows"),
        ("IA / LLM",              "Claude API (Anthropic) — analises, resumos, scoring e geracao de relatorios"),
        ("Busca semantica",       "Embeddings com text-embedding-3-large + pgvector para busca em base historica"),
        ("Mensageria",            "WhatsApp Business API (Evolution API self-hosted) + Resend para e-mail transacional"),
        ("Infraestrutura",        "Vercel (frontend) + Railway ou Render (backend) + Supabase (banco + auth)"),
        ("Monitoramento",         "Grafana + Sentry para observabilidade do sistema"),
    ]

    t = doc.add_table(rows=len(stack), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (layer, tech) in enumerate(stack):
        lc, vc = t.rows[i].cells
        lc.width = Cm(4.5); vc.width = Cm(11.7)
        bg = "F4F4F4" if i % 2 == 0 else "FAFAFA"
        set_bg(lc, bg); set_bg(vc, bg)
        set_border(lc, bottom={"val":"single","sz":"4","space":"0","color":"E8E8E8"})
        set_border(vc, bottom={"val":"single","sz":"4","space":"0","color":"E8E8E8"})
        lp = lc.paragraphs[0]; lr = lp.add_run(layer); fnt(lr, size=10, bold=True, color=GRAY)
        vp = vc.paragraphs[0]; vr = vp.add_run(tech);  fnt(vr, size=10.5)

    doc.add_paragraph()

    callout(doc,
        "A arquitetura inteira pode ser construida com ferramentas de codigo aberto e APIs acessiveis. "
        "O custo de infraestrutura para atender os primeiros 20 clientes e estimado em R$800 a R$1.500/mes. "
        "O custo marginal de adicionar um novo cliente, apos o sistema estruturado, e proximo de zero.",
        bg="F0F4FF", border="1A56C8")

    doc.add_page_break()

    # ─── CONCLUSAO ───────────────────────────────────────────────────────────────
    sec(doc, "Conclusao — o que este estudo representa para a LocalRise")
    jp(doc,
       "Este documento nao e apenas um estudo academico sobre tecnologia. "
       "E um mapa de produto para uma oportunidade real: o mercado de inteligencia regulatoria no Brasil "
       "ainda nao tem um produto de SaaS dominante. As consultorias estabelecidas vendem servico — nao produto. "
       "As ferramentas estrangeiras nao tem contexto do ambiente politico brasileiro.")
    jp(doc,
       "A LocalRise tem uma posicao unica: esta proxima das consultorias (como a Capitol), entende o mercado "
       "e tem capacidade de construir a tecnologia. O caminho mais inteligente e comecar com um cliente "
       "ancoragem — usar o contrato da Capitol para validar o painel e as automacoes — e transformar "
       "esse MVP em um produto replicavel para outras consultorias.")

    callout(doc,
        "A Capitol nao seria apenas um cliente. Seria o primeiro caso de uso real de um produto que a LocalRise "
        "pode escalar para todo o mercado de inteligencia governamental no Brasil. "
        "Isso muda completamente a proposta de valor da conversa.",
        bg="FDF3F3", border="E31B23")

    sub(doc, "Proximos passos concretos")
    nblt(doc, "Validar com a Capitol quais automacoes e funcionalidades de painel fazem mais sentido para o seu contexto imediato")
    nblt(doc, "Construir um MVP do painel com 3 funcionalidades: feed de inteligencia, score de risco e mapa de stakeholders")
    nblt(doc, "Implementar os primeiros 3 fluxos no n8n: coleta legislativa, resumo via IA e alerta por WhatsApp")
    nblt(doc, "Usar a Capitol como caso de uso validado para apresentar o produto a outras consultorias e setores regulados")
    nblt(doc, "Avaliar o modelo de licenciamento white-label como caminho de escala com custo marginal baixo")

    doc.add_paragraph()
    signature(doc)

    out = BASE_DIR / "saas-inteligencia-regulatoria-localrise.docx"
    doc.save(out)
    print(f"Documento salvo em: {out}")


if __name__ == "__main__":
    build()
