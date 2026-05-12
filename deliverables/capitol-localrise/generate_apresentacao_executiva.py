from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
LOGO_PATH = Path(
    r"C:\Users\digui\Documents\localrise-brain\Claude\Claude Local RIse\assents\logo Local Rise Principal.png"
)

RED    = RGBColor(0xE3, 0x1B, 0x23)
BLACK  = RGBColor(0x0A, 0x0A, 0x0A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x5F, 0x63, 0x68)
LGRAY  = RGBColor(0xD9, 0xDD, 0xE3)
BLUE   = RGBColor(0x1A, 0x56, 0xC8)
GREEN  = RGBColor(0x1A, 0x7A, 0x3C)
ORANGE = RGBColor(0xC8, 0x6A, 0x00)
PURPLE = RGBColor(0x5A, 0x1A, 0x8C)

DATE = "10 de abril de 2026"


# ─── helpers ─────────────────────────────────────────────────────────────────────

def set_border(cell, **kw):
    tc    = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tb    = tc_pr.first_child_found_in("w:tcBorders")
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tc_pr.append(tb)
    for edge, data in kw.items():
        tag = "w:%s" % edge
        el  = tb.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); tb.append(el)
        for k, v in data.items():
            el.set(qn("w:%s" % k), str(v))


def remove_all_borders(cell):
    for edge in ("left", "right", "top", "bottom", "insideH", "insideV"):
        set_border(cell, **{edge: {"val": "none", "sz": "0", "space": "0", "color": "auto"}})


def set_bg(cell, hex_color: str):
    tc    = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def set_page_bg(doc, hex_color: str):
    """Adds a full-bleed background table on the current page."""
    bg_t = doc.add_table(rows=1, cols=1)
    bg_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = bg_t.cell(0, 0)
    c.width = Cm(27.7)
    set_bg(c, hex_color)
    remove_all_borders(c)


def apply_styles(doc: Document):
    s              = doc.sections[0]
    # Landscape A4
    s.page_width   = Cm(29.7)
    s.page_height  = Cm(21)
    s.top_margin   = Cm(0.8)
    s.bottom_margin = Cm(0.8)
    s.left_margin  = Cm(1.2)
    s.right_margin = Cm(1.2)

    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    n.font.size = Pt(12)
    n.paragraph_format.space_after       = Pt(0)
    n.paragraph_format.space_before      = Pt(0)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def fnt(run, size=12, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = color


def p_centered(doc, text, size=12, bold=False, color=None, sb=0, sa=0, italic=False):
    par = doc.add_paragraph()
    par.alignment                         = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before     = Pt(sb)
    par.paragraph_format.space_after      = Pt(sa)
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        run = par.add_run(text)
        fnt(run, size=size, bold=bold, color=color, italic=italic)
    return par


def p_left(doc, text, size=12, bold=False, color=None, sb=0, sa=0, italic=False):
    par = doc.add_paragraph()
    par.alignment                         = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_before     = Pt(sb)
    par.paragraph_format.space_after      = Pt(sa)
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        run = par.add_run(text)
        fnt(run, size=size, bold=bold, color=color, italic=italic)
    return par


def spacer(doc, size=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after  = Pt(0)
    run = par.add_run(" ")
    fnt(run, size=size)


def slide_break(doc):
    doc.add_page_break()


# ─── slide builders ───────────────────────────────────────────────────────────────

def full_bg_table(doc, hex_color, height_cm=18.0):
    """Creates a full-width colored background table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    c.width = Cm(27.3)
    set_bg(c, hex_color)
    remove_all_borders(c)
    tr = c._tc.getparent()
    trpr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trHeight.set(qn("w:hRule"), "exact")
    trpr.append(trHeight)
    return c


def slide_1_capa(doc):
    """SLIDE 1 — CAPA"""
    # Faixa vermelha superior
    t_top = doc.add_table(rows=1, cols=1)
    t_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_top = t_top.cell(0, 0); set_bg(c_top, "E31B23"); remove_all_borders(c_top)
    tr_top = c_top._tc.getparent(); trpr = tr_top.get_or_add_trPr()
    trH = OxmlElement("w:trHeight"); trH.set(qn("w:val"), "800"); trH.set(qn("w:hRule"), "exact")
    trpr.append(trH)
    p1 = c_top.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("  LOCALRISE ADVISORY  ×  CAPITOL CONSULTORIA")
    fnt(r1, size=10, bold=True, color=WHITE)

    spacer(doc, 6)

    # Logo centrado
    if LOGO_PATH.exists():
        lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(4)
        lp.add_run().add_picture(str(LOGO_PATH), width=Cm(4.5))

    spacer(doc, 4)

    p_centered(doc,
        "A Nova Estrutura de Inteligencia Regulatoria",
        size=30, bold=True, color=BLACK, sb=4)

    spacer(doc, 6)

    p_centered(doc,
        "Como transformar senioridade institucional em sistema escalavel de inteligencia e vantagem competitiva",
        size=14, color=GRAY, sa=4)

    spacer(doc, 8)

    # Faixa inferior preta
    t_bot = doc.add_table(rows=1, cols=3)
    t_bot.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = t_bot.rows[0].cells
    widths = [Cm(9.1), Cm(9.1), Cm(9.1)]
    labels = ["LocalRise Advisory", "Apresentacao Estrategica", DATE]
    for i, (cell, w, label) in enumerate(zip(cells, widths, labels)):
        cell.width = w; set_bg(cell, "0A0A0A"); remove_all_borders(cell)
        cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(label); fnt(cr, size=10, bold=(i == 0), color=WHITE)

    slide_break(doc)


def slide_header_full(doc, slide_num, title, subtitle=None,
                       bg_hex="0A0A0A", title_color=WHITE, sub_color=None):
    """Full-width dark header band."""
    if sub_color is None:
        sub_color = RGBColor(0xCC, 0xCC, 0xCC)

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); c.width = Cm(27.3)
    set_bg(c, bg_hex); remove_all_borders(c)
    tr = c._tc.getparent(); trpr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), "1300"); trH.set(qn("w:hRule"), "exact")
    trpr.append(trH)

    hp = c.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(2)
    num_run = hp.add_run(f"  {slide_num:02d}  "); fnt(num_run, size=10, color=RED)
    title_run = hp.add_run(title.upper()); fnt(title_run, size=18, bold=True, color=title_color)
    if subtitle:
        hp2 = c.add_paragraph(); hp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp2.paragraph_format.space_before = Pt(0)
        sr = hp2.add_run("     " + subtitle); fnt(sr, size=11, italic=True, color=sub_color)


def slide_two_col(doc, left_items, right_items,
                  left_bg="F4F4F4", right_bg="FFFFFF",
                  left_header=None, right_header=None,
                  left_accent="E31B23", right_accent="1A56C8"):
    """Two-column layout."""
    spacer(doc, 4)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc, rc = t.rows[0].cells
    lc.width = Cm(13.3); rc.width = Cm(13.3)
    set_bg(lc, left_bg); set_bg(rc, right_bg)
    remove_all_borders(lc); remove_all_borders(rc)
    set_border(lc, right={"val": "single", "sz": "6", "space": "0", "color": "DDDDDD"})

    if left_header:
        lhp = lc.paragraphs[0]; lhp.paragraph_format.space_before = Pt(4)
        lhr = lhp.add_run("  " + left_header.upper())
        fnt(lhr, size=10, bold=True, color=RGBColor(
            int(left_accent[0:2], 16), int(left_accent[2:4], 16), int(left_accent[4:6], 16)))
    else:
        lc.paragraphs[0].text = ""

    for item in left_items:
        lp2 = lc.add_paragraph()
        lp2.paragraph_format.space_before = Pt(3)
        lp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp2.add_run("  • " + item); fnt(lr, size=12)

    if right_header:
        rhp = rc.paragraphs[0]; rhp.paragraph_format.space_before = Pt(4)
        rhr = rhp.add_run("  " + right_header.upper())
        fnt(rhr, size=10, bold=True, color=RGBColor(
            int(right_accent[0:2], 16), int(right_accent[2:4], 16), int(right_accent[4:6], 16)))
    else:
        rc.paragraphs[0].text = ""

    for item in right_items:
        rp2 = rc.add_paragraph()
        rp2.paragraph_format.space_before = Pt(3)
        rp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = rp2.add_run("  • " + item); fnt(rr, size=12)


def slide_statement(doc, statement, author=None,
                     bg_hex="E31B23", text_color=WHITE):
    """Big quote/statement slide."""
    spacer(doc, 6)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); c.width = Cm(24)
    set_bg(c, bg_hex); remove_all_borders(c)
    tr = c._tc.getparent(); trpr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight"); trH.set(qn("w:val"), "3500"); trH.set(qn("w:hRule"), "exact")
    trpr.append(trH)
    sp = c.paragraphs[0]; sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(8)
    sr = sp.add_run(statement)
    fnt(sr, size=22, bold=True, color=text_color)
    if author:
        ap = c.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = ap.add_run("— " + author)
        fnt(ar, size=11, italic=True, color=RGBColor(0xEE, 0xEE, 0xEE))


def slide_cards_row(doc, cards):
    """Row of colored mini-cards (up to 4)."""
    spacer(doc, 6)
    n = len(cards)
    width = Cm(27.3 / n - 0.2)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (title, body, bg_hex, text_color_hex) in enumerate(cards):
        c = t.rows[0].cells[i]; c.width = width
        set_bg(c, bg_hex); remove_all_borders(c)
        tr = c._tc.getparent(); trpr = tr.get_or_add_trPr()
        trH = OxmlElement("w:trHeight"); trH.set(qn("w:val"), "4500"); trH.set(qn("w:hRule"), "exact")
        trpr.append(trH)
        tp = c.paragraphs[0]; tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(6)
        tr2 = tp.add_run(title + "\n")
        fnt(tr2, size=13, bold=True, color=RGBColor(
            int(text_color_hex[0:2], 16), int(text_color_hex[2:4], 16), int(text_color_hex[4:6], 16)))
        bp = c.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(4)
        br = bp.add_run(body)
        fnt(br, size=10.5, color=RGBColor(
            int(text_color_hex[0:2], 16), int(text_color_hex[2:4], 16), int(text_color_hex[4:6], 16)))


def slide_big_numbers(doc, items):
    """Row of big KPI numbers."""
    spacer(doc, 8)
    n = len(items)
    width = Cm(27.3 / n - 0.3)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (number, label, bg_hex, num_color) in enumerate(items):
        c = t.rows[0].cells[i]; c.width = width
        set_bg(c, bg_hex); remove_all_borders(c)
        if i < n - 1:
            set_border(c, right={"val": "single", "sz": "6", "space": "0", "color": "DDDDDD"})
        np2 = c.paragraphs[0]; np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np2.paragraph_format.space_before = Pt(8)
        nr = np2.add_run(number)
        fnt(nr, size=34, bold=True, color=RGBColor(
            int(num_color[0:2], 16), int(num_color[2:4], 16), int(num_color[4:6], 16)))
        lp = c.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(4)
        lr = lp.add_run(label); fnt(lr, size=11, color=GRAY)


def slide_arrow_flow(doc, steps):
    """Left-to-right flow with arrows."""
    spacer(doc, 8)
    n = len(steps)
    w_item = Cm(27.3 / n - 0.2)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, desc, bg_hex, txt_hex) in enumerate(steps):
        c = t.rows[0].cells[i]; c.width = w_item
        set_bg(c, bg_hex); remove_all_borders(c)
        tp = c.paragraphs[0]; tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(6)
        tr2 = tp.add_run(label)
        fnt(tr2, size=12, bold=True, color=RGBColor(
            int(txt_hex[0:2], 16), int(txt_hex[2:4], 16), int(txt_hex[4:6], 16)))
        dp = c.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.paragraph_format.space_before = Pt(3)
        dr = dp.add_run(desc)
        fnt(dr, size=10, color=RGBColor(
            int(txt_hex[0:2], 16), int(txt_hex[2:4], 16), int(txt_hex[4:6], 16)))


def slide_before_after(doc, rows_data):
    """Before/After comparison table."""
    spacer(doc, 6)
    t = doc.add_table(rows=1 + len(rows_data), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(9.1), Cm(9.1), Cm(9.1)]

    # Header row
    headers = ["O QUE E HOJE", "DEPOIS DA LOCALRISE", "GANHO REAL"]
    colors  = ["4A4A4A", "1A7A3C", "E31B23"]
    for i, (hdr, clr, w) in enumerate(zip(headers, colors, widths)):
        c = t.rows[0].cells[i]; c.width = w
        set_bg(c, clr); remove_all_borders(c)
        hp = c.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(hdr); fnt(hr, size=10, bold=True, color=WHITE)

    # Data rows
    bg_alts = [("F9F9F9", "F0FFF4", "FFF4F4"), ("F4F4F4", "E8FFE8", "FFE8E8")]
    for ri, (before, after, gain) in enumerate(rows_data):
        row = t.rows[ri + 1]; alt = bg_alts[ri % 2]
        for ci, (val, bg_hex, w) in enumerate(zip([before, after, gain], alt, widths)):
            cell = row.cells[ci]; cell.width = w
            set_bg(cell, bg_hex); remove_all_borders(cell)
            set_border(cell, bottom={"val": "single", "sz": "4", "space": "0", "color": "E0E0E0"})
            vp = cell.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            vp.paragraph_format.space_before = Pt(2)
            vr = vp.add_run("  " + val); fnt(vr, size=11)


def slide_saas_grid(doc, products):
    """Grid de produtos SaaS."""
    spacer(doc, 4)
    n = 5
    w = Cm(27.3 / n - 0.15)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (name, tagline, body, bg_hex, txt_hex) in enumerate(products):
        c = t.rows[0].cells[i]; c.width = w
        set_bg(c, bg_hex); remove_all_borders(c)
        tr = c._tc.getparent(); trpr = tr.get_or_add_trPr()
        trH = OxmlElement("w:trHeight"); trH.set(qn("w:val"), "5500"); trH.set(qn("w:hRule"), "exact")
        trpr.append(trH)
        np2 = c.paragraphs[0]; np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np2.paragraph_format.space_before = Pt(6)
        nr = np2.add_run(name)
        fnt(nr, size=14, bold=True, color=RGBColor(
            int(txt_hex[0:2], 16), int(txt_hex[2:4], 16), int(txt_hex[4:6], 16)))
        tgp = c.add_paragraph(); tgp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tgp.paragraph_format.space_before = Pt(2)
        tgr = tgp.add_run(tagline)
        fnt(tgr, size=8.5, italic=True, color=RGBColor(
            int(txt_hex[0:2], 16), int(txt_hex[2:4], 16), int(txt_hex[4:6], 16)))
        bp = c.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(8)
        br2 = bp.add_run(body)
        fnt(br2, size=9.5, color=RGBColor(
            int(txt_hex[0:2], 16), int(txt_hex[2:4], 16), int(txt_hex[4:6], 16)))


def slide_cta(doc):
    """SLIDE 15 — CTA final."""
    spacer(doc, 4)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); c.width = Cm(22)
    set_bg(c, "E31B23"); remove_all_borders(c)
    tr = c._tc.getparent(); trpr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight"); trH.set(qn("w:val"), "4000"); trH.set(qn("w:hRule"), "exact")
    trpr.append(trH)
    hp = c.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(8)
    hr = hp.add_run("Vamos construir isso com base no seu contexto?")
    fnt(hr, size=22, bold=True, color=WHITE)
    sp = c.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(6)
    sr = sp.add_run("A proxima reuniao e para mapear, junto com a Capitol, por onde comecar.")
    fnt(sr, size=13, italic=True, color=RGBColor(0xEE, 0xEE, 0xEE))

    spacer(doc, 8)

    t2 = doc.add_table(rows=1, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    steps_cta = [
        ("1", "Definir os temas mais criticos\npara monitorar agora"),
        ("2", "Escolher por onde comecar:\nLexRadar ou RegRisk"),
        ("3", "Mapear o primeiro fluxo\nde automacao em 30 dias"),
    ]
    for i, (num, label) in enumerate(steps_cta):
        c2 = t2.rows[0].cells[i]; c2.width = Cm(9.1)
        set_bg(c2, "0A0A0A"); remove_all_borders(c2)
        np2 = c2.paragraphs[0]; np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np2.paragraph_format.space_before = Pt(4)
        nr = np2.add_run(num); fnt(nr, size=20, bold=True, color=RED)
        lp = c2.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(2)
        lr = lp.add_run(label); fnt(lr, size=11, color=WHITE)


# ══════════════════════════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    apply_styles(doc)

    # ─── SLIDE 1 — CAPA ──────────────────────────────────────────────────────────
    slide_1_capa(doc)

    # ─── SLIDE 2 — CONTEXTO DO MERCADO ───────────────────────────────────────────
    slide_header_full(doc, 2,
        "O ambiente regulatorio esta ficando mais complexo",
        "O modelo tradicional de consultoria nao acompanha mais essa velocidade",
        bg_hex="0A0A0A")
    spacer(doc, 4)
    slide_two_col(doc,
        left_header="O CENARIO ATUAL",
        left_items=[
            "Leis e normas mudam com frequencia crescente",
            "Decisoes em Brasilia impactam negócios em horas",
            "Quem age tarde, paga o preco",
            "O volume de informacao relevante dobrou nos ultimos 5 anos",
            "Setores regulados vivem em estado de alerta permanente",
        ],
        right_header="A LIMITACAO DO MODELO ATUAL",
        right_items=[
            "Consultorias dependem da memoria dos seus consultores",
            "Relatorios chegam depois que o dano ja aconteceu",
            "Monitorar tudo manualmente e humanamente inviavel",
            "Escalar exige contratar mais — e o custo explode",
            "A informacao fica dispersa, nao sistematizada",
        ],
        left_bg="F8F8F8", right_bg="FFFFFF",
        left_accent="E31B23", right_accent="1A56C8")
    slide_break(doc)

    # ─── SLIDE 3 — O PROBLEMA CENTRAL ────────────────────────────────────────────
    slide_header_full(doc, 3,
        "Quatro problemas que toda consultoria carrega",
        "Nenhum deles e resolvido com mais pessoas — e sim com sistema",
        bg_hex="E31B23")
    slide_cards_row(doc, [
        ("Sem escala",
         "Para dobrar de clientes,\nvocê dobra a equipe.\nNao ha outro caminho.",
         "0A0A0A", "FFFFFF"),
        ("Dependencia de pessoas",
         "O conhecimento esta na\ncabeca do consultor.\nQuando ele sai, vai junto.",
         "1A1A1A", "FFFFFF"),
        ("Informacao dispersa",
         "Dados em e-mail, WhatsApp,\nPDF e anotacoes.\nNada centralizado.",
         "2A2A2A", "FFFFFF"),
        ("Reatividade",
         "A crise e descoberta quando\njá chegou na midia.\nTarde demais para agir.",
         "3A3A3A", "FFFFFF"),
    ])
    slide_break(doc)

    # ─── SLIDE 4 — MUDANCA DE MODELO ─────────────────────────────────────────────
    slide_header_full(doc, 4,
        "A mudanca que separa o passado do futuro",
        "Nao e sobre tecnologia. E sobre como a consultoria opera e cresce",
        bg_hex="1A56C8")
    spacer(doc, 6)
    slide_arrow_flow(doc, [
        ("ANTES\nConsultoria",
         "Servico vendido por horas\nRelatorio mensal\nMemoria do consultor\nReativo\nEscala limitada",
         "F0F0F0", "4A4A4A"),
        ("→", "", "FFFFFF", "CCCCCC"),
        ("DEPOIS\nSistema",
         "Inteligencia continua\nAlertas em tempo real\nBase de dados estruturada\nPreditivo\nEscala ilimitada",
         "0A0A0A", "FFFFFF"),
        ("→", "", "FFFFFF", "CCCCCC"),
        ("RESULTADO\nVantagem",
         "Cliente sabe antes\nDecisao com dados\nRisco controlado\nOportunidades capturadas\nLideranca de mercado",
         "E31B23", "FFFFFF"),
    ])
    slide_break(doc)

    # ─── SLIDE 5 — O QUE A LOCALRISE CONSTROI ────────────────────────────────────
    slide_header_full(doc, 5,
        "O que a LocalRise esta construindo",
        "Uma infraestrutura de inteligencia regulatoria que o mercado ainda nao tem",
        bg_hex="0A0A0A")
    spacer(doc, 6)
    slide_cards_row(doc, [
        ("Dashboard\nEstrategico",
         "Um painel que mostra, em tempo\nreal, o seu nivel de risco,\nas oportunidades abertas e\nquem esta decidindo o que.",
         "F4F4F4", "0A0A0A"),
        ("Automacoes\nInteligentes",
         "Processos que antes levavam\nhoras passam a acontecer\nautonomaticamente — com\nqualidade superior.",
         "F0F4FF", "1A56C8"),
        ("Alertas\nPersonalizados",
         "Cada cliente recebe so o que\ne relevante para o seu setor,\nno momento certo,\nsem precisar pedir.",
         "FFF8F0", "C86A00"),
        ("Predicao\ne IA",
         "Antecipar o que vai acontecer\nno ambiente regulatorio antes\nque qualquer concorrente\nperceba.",
         "FDF3F3", "E31B23"),
    ])
    spacer(doc, 6)
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); c.width = Cm(20)
    set_bg(c, "F0F4FF"); remove_all_borders(c)
    set_border(c, left={"val": "single", "sz": "18", "space": "0", "color": "1A56C8"})
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = pp.add_run(
        "Tudo isso integrado em um unico sistema — que roda enquanto o consultor faz o que mais importa.")
    fnt(pr, size=12, italic=True, color=BLUE)
    slide_break(doc)

    # ─── SLIDE 6 — VISAO GERAL DOS SAAS ──────────────────────────────────────────
    slide_header_full(doc, 6,
        "Cinco produtos que mudam como a Capitol opera",
        "Cada um resolve um problema real. Todos se integram.",
        bg_hex="0A0A0A")
    slide_saas_grid(doc, [
        ("LexRadar",
         "Inteligencia legislativa",
         "Voce sabe quando\numa lei muda.\nAntes de todo mundo.",
         "E31B23", "FFFFFF"),
        ("PoliticaIQ",
         "Mapa de influencia",
         "Quem decide.\nQuem influencia quem.\nQual o momento certo.",
         "1A56C8", "FFFFFF"),
        ("RegRisk",
         "Score de risco diario",
         "Um numero que resume\nsua exposicao\nregulatorio hoje.",
         "1A7A3C", "FFFFFF"),
        ("OpportunityMap",
         "Radar de oportunidades",
         "Incentivos, programas\ne financiamentos\nantes do prazo.",
         "C86A00", "FFFFFF"),
        ("CrisisMonitor",
         "Alerta precoce",
         "Saber 48h antes\nque a crise chegue\nna midia.",
         "5A1A8C", "FFFFFF"),
    ])
    slide_break(doc)

    # ─── SLIDE 7 — IMPACTO DOS SAAS ──────────────────────────────────────────────
    slide_header_full(doc, 7,
        "O que muda quando esses produtos entram em operacao",
        "Vantagem competitiva concreta — nao hipotetica",
        bg_hex="1A7A3C")
    spacer(doc, 6)
    slide_two_col(doc,
        left_header="PARA A CAPITOL COMO CONSULTORIA",
        left_items=[
            "Entrega inteligencia que a maioria nao consegue replicar",
            "Diferencia-se como parceira estrategica — nao apenas operacional",
            "Retém clientes com valor percebido crescente",
            "Atende mais empresas sem aumentar a equipe",
            "Gera novo modelo de receita recorrente via produto",
        ],
        right_header="PARA O CLIENTE FINAL DA CAPITOL",
        right_items=[
            "Sabe de mudancas regulatorias antes dos concorrentes",
            "Decide com base em dados — nao em intuicao",
            "Antecipa crises e age enquanto ainda ha tempo",
            "Captura oportunidades que antes passavam despercebidas",
            "Tem visibilidade de quem decide o que em Brasilia",
        ],
        left_bg="F4FFF6", right_bg="F0F8FF",
        left_accent="1A7A3C", right_accent="1A56C8")
    slide_break(doc)

    # ─── SLIDE 8 — ANTES VS DEPOIS ───────────────────────────────────────────────
    slide_header_full(doc, 8,
        "O antes e o depois em numeros reais",
        "O que muda na operacao diaria — com impacto direto no resultado",
        bg_hex="0A0A0A")
    slide_before_after(doc, [
        ("Ler diarios oficiais manualmente\n3 a 4h por analista por dia",
         "Monitoramento automatico\nem tempo real",
         "15 a 25h/semana por analista recuperadas"),
        ("Relatorio semanal leva 3 a 4h\npor cliente para produzir",
         "Relatorio gerado automaticamente\nem 15 minutos",
         "75h/mes recuperadas com 20 clientes"),
        ("Analise de nova lei leva 3h\npara escrever e enviar",
         "Analise entregue em 8 minutos\npelo sistema",
         "De 3 horas para 8 minutos por analise"),
        ("Onboarding de novo cliente\nleva 4 a 8 horas",
         "Cliente ativo em\n30 minutos",
         "40h recuperadas por mes com 5 novos clientes"),
        ("Crise regulatoria descoberta\nquando ja esta na midia",
         "Alerta 24 a 72 horas\nantes da midia",
         "Tempo real de reposicionamento estrategico"),
    ])
    slide_break(doc)

    # ─── SLIDE 9 — IMPACTO NA CAPITOL ────────────────────────────────────────────
    slide_header_full(doc, 9,
        "O que isso muda especificamente para a Capitol",
        "Tres dimensoes de impacto direto no negocio",
        bg_hex="E31B23")
    spacer(doc, 6)
    slide_cards_row(doc, [
        ("Crescimento sem custo proporcional",
         "Hoje: dobrar de clientes = dobrar equipe.\nCom o sistema: atender 2x mais\nsem contratar ninguem.\nO custo marginal e proximo de zero.",
         "0A0A0A", "FFFFFF"),
        ("Valor percebido muito maior",
         "O cliente passa a receber inteligencia\ndiaria — nao relatorio mensal.\nA relacao muda de fornecedor\npara parceiro estrategico permanente.",
         "1A1A1A", "FFFFFF"),
        ("Diferenciacao real de mercado",
         "Nenhuma outra consultoria de\nrelacoes governamentais no Brasil\ntem esse sistema.\nQuem entra primeiro, lidera.",
         "E31B23", "FFFFFF"),
    ])
    slide_break(doc)

    # ─── SLIDE 10 — DIFERENCIAL COMPETITIVO ──────────────────────────────────────
    slide_header_full(doc, 10,
        "O diferencial que nao se copia facilmente",
        "Nao e a tecnologia. E o que ela permite fazer.",
        bg_hex="0A0A0A")
    spacer(doc, 8)
    slide_big_numbers(doc, [
        ("10x", "mais temas monitorados\nsem aumentar equipe", "F9F9F9", "E31B23"),
        ("48h", "de antecedencia\nna deteccao de crises", "F9F9F9", "1A56C8"),
        ("75h", "recuperadas por mes\npor analista", "F9F9F9", "1A7A3C"),
        ("0", "custo marginal por\nnovo cliente no sistema", "F9F9F9", "C86A00"),
    ])
    spacer(doc, 8)
    slide_statement(doc,
        "Saber antes nao e sorte. E sistema.",
        bg_hex="E31B23", text_color=WHITE)
    slide_break(doc)

    # ─── SLIDE 11 — POSICIONAMENTO DE MERCADO ────────────────────────────────────
    slide_header_full(doc, 11,
        "Como a Capitol se posiciona com isso",
        "De consultoria de relacoes governamentais para intelligence partner",
        bg_hex="1A56C8")
    spacer(doc, 6)
    slide_two_col(doc,
        left_header="SEM O SISTEMA",
        left_items=[
            "Vende horas de consultores",
            "Entrega relatorios periodicos",
            "Depende da agenda e memoria da equipe",
            "Crescimento linear com custo linear",
            "Percebida como fornecedor de servico",
            "Vantagem: relacionamento e reputacao offline",
        ],
        right_header="COM O SISTEMA",
        right_items=[
            "Vende inteligencia e sistema",
            "Entrega monitoramento continuo e alertas em tempo real",
            "Opera com escala tecnologica",
            "Crescimento exponencial com custo controlado",
            "Percebida como parceira estrategica indispensavel",
            "Vantagem: dados + relacionamento + tecnologia",
        ],
        left_bg="F9F9F9", right_bg="F0F8FF",
        left_accent="4A4A4A", right_accent="1A56C8")
    slide_break(doc)

    # ─── SLIDE 12 — OPORTUNIDADE ESTRATEGICA ─────────────────────────────────────
    slide_header_full(doc, 12,
        "Uma oportunidade maior do que parece",
        "O sistema abre portas que hoje estao fechadas",
        bg_hex="C86A00")
    spacer(doc, 4)
    slide_cards_row(doc, [
        ("Novo modelo de receita",
         "Alem da consultoria por hora,\numa mensalidade recorrente\npor acesso ao sistema.\nReceita que nao depende de projeto.",
         "F9F9F9", "0A0A0A"),
        ("White-label para o mercado",
         "A Capitol pode licenciar\no sistema para outras consultorias\ne associacoes setoriais.\nSe torna infraestrutura do mercado.",
         "0A0A0A", "FFFFFF"),
        ("Novo perfil de cliente",
         "Alem das empresas ja atendidas,\nqualquer empresa com exposicao\nregulatorio vira cliente potencial\ndo produto.",
         "F4F4F4", "0A0A0A"),
        ("Barreira de entrada",
         "Uma vez com o sistema rodando,\nos concorrentes precisam de meses\npara chegar no mesmo nivel.\nJanela estrategica agora.",
         "E31B23", "FFFFFF"),
    ])
    slide_break(doc)

    # ─── SLIDE 13 — VISAO DE FUTURO ──────────────────────────────────────────────
    slide_header_full(doc, 13,
        "Como isso evolui com o tempo",
        "O sistema aprende. A vantagem cresce.",
        bg_hex="5A1A8C")
    spacer(doc, 6)
    slide_arrow_flow(doc, [
        ("30 DIAS",
         "Monitoramento ativo\nAlertas funcionando\nPrimeiros relatorios automaticos\nPainel basico no ar",
         "F4F4F4", "0A0A0A"),
        ("60 DIAS",
         "Score de risco calibrado\nMapa de stakeholders ativo\nFluxos de follow-up rodando\nPrimeira analise de impacto automatica",
         "F0F4FF", "1A56C8"),
        ("90 DIAS",
         "CrisisMonitor ativo\nOpportunityMap integrado\nPrimeiro cliente piloto validado\nMVP do produto pronto",
         "F0FFF4", "1A7A3C"),
        ("6 MESES+",
         "Sistema preditivo rodando\nIA aprendendo com historico\nLicenciamento white-label\nNovo padrao de mercado",
         "FDF3F3", "E31B23"),
    ])
    spacer(doc, 6)
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); c.width = Cm(22)
    set_bg(c, "F9F0FF"); remove_all_borders(c)
    set_border(c, left={"val": "single", "sz": "18", "space": "0", "color": "5A1A8C"})
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = pp.add_run(
        "O sistema fica mais inteligente a cada novo cliente e a cada nova lei monitorada. "
        "Quem entra primeiro constroi a vantagem que os outros nao conseguem copiar.")
    fnt(pr, size=11, italic=True, color=PURPLE)
    slide_break(doc)

    # ─── SLIDE 14 — FECHAMENTO ESTRATEGICO ───────────────────────────────────────
    slide_header_full(doc, 14,
        "O que a Capitol tem que a maioria nao tem",
        "A combinacao certa — no momento certo",
        bg_hex="0A0A0A")
    spacer(doc, 6)
    slide_two_col(doc,
        left_header="O QUE A CAPITOL JA TEM",
        left_items=[
            "Credibilidade institucional consolidada",
            "Rede de relacionamentos em Brasilia",
            "Senioridade e leitura politica real",
            "Reputacao construida ao longo do tempo",
            "Conhecimento profundo do ambiente regulatorio",
        ],
        right_header="O QUE O SISTEMA ADICIONA",
        right_items=[
            "Escala que hoje nao existe",
            "Inteligencia continua — 24 horas por dia",
            "Predicao e antecipacao de movimentos",
            "Um produto que vai alem do contrato de consultoria",
            "Posicionamento de lideranca tecnologica no mercado",
        ],
        left_bg="F9F9F9", right_bg="FDF3F3",
        left_accent="4A4A4A", right_accent="E31B23")
    spacer(doc, 6)
    slide_statement(doc,
        "A Capitol nao precisa escolher entre ser consultoria\nou ter tecnologia.\nEla pode ser as duas coisas — e liderar as duas.",
        bg_hex="0A0A0A", text_color=WHITE)
    slide_break(doc)

    # ─── SLIDE 15 — CTA ──────────────────────────────────────────────────────────
    slide_header_full(doc, 15,
        "Proximo passo",
        "Uma conversa. Sem compromisso. Com foco no seu contexto.",
        bg_hex="E31B23")
    slide_cta(doc)

    spacer(doc, 6)
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_items = [
        "LocalRise Advisory", "Nicolas Slater | Diretoria Comercial", DATE
    ]
    for i, (cell, label) in enumerate(zip(t.rows[0].cells, footer_items)):
        cell.width = Cm(9.1); set_bg(cell, "F4F4F4"); remove_all_borders(cell)
        set_border(cell, top={"val": "single", "sz": "4", "space": "0", "color": "CCCCCC"})
        cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(label); fnt(cr, size=9, color=GRAY, bold=(i == 0))

    out = BASE_DIR / "apresentacao-executiva-inteligencia-regulatoria.docx"
    doc.save(out)
    print(f"Documento salvo em: {out}")


if __name__ == "__main__":
    build()
