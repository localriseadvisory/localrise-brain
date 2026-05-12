from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
LOGO_PATH = Path(
    r"C:\Users\digui\Documents\localrise-brain\Claude\Claude Local RIse\assents\logo Local Rise Principal.png"
)

RED = RGBColor(0xE3, 0x1B, 0x23)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)
GRAY = RGBColor(0x5F, 0x63, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

COMPANY = "Capitol Consultoria"
DATE = "09 de abril de 2026"


# ── helpers ─────────────────────────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:%s" % edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:%s" % key), str(edge_data[key]))


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def apply_doc_style(doc: Document):
    s = doc.sections[0]
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.4)
    s.right_margin = Cm(2.4)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    n.font.size = Pt(11)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    n.paragraph_format.space_after = Pt(6)


def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=11, bold=False,
         color=None, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if text:
        r = p.add_run(text)
        font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def justify(doc, text, size=11, bold=False, color=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, size=size, bold=bold, color=color, italic=italic)


def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    font(r, size=size)


def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    font(r, size=13, bold=True, color=RED)


def divider(doc):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_border(cell, bottom={"val": "single", "sz": 6, "space": 0, "color": "E0E0E0"})
    cell.text = ""
    doc.add_paragraph()


# ── componentes visuais ──────────────────────────────────────────────────────────

def add_brand_header(doc):
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Cm(5.4))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LOCALRISE ADVISORY")
    font(r, size=12, bold=True, color=RED)

    divider(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Como a Capitol pode parar de depender de indicacao")
    font(r, size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Estudo de automacoes e sistema de aquisicao | Capitol Consultoria")
    font(r, size=11, color=GRAY)


def callout(doc, text, bg="FDF3F3", border_color="E31B23"):
    """Caixa de destaque — usada para pontos-chave de reuniao."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_cell_bg(c, bg)
    set_cell_border(c, left={"val": "single", "sz": 18, "space": 0, "color": border_color})
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    font(r, size=11, italic=True)
    doc.add_paragraph()


def problem_solution_card(doc, number, title, problem, impact, solution, gain):
    """Bloco visual: Problema → Impacto → Solucao → Ganho."""
    # Cabecalho do card
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    num_cell = t.rows[0].cells[0]
    ttl_cell = t.rows[0].cells[1]
    num_cell.width = Cm(1.4)
    ttl_cell.width = Cm(14.8)
    set_cell_bg(num_cell, "E31B23")
    np = num_cell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run(number)
    font(nr, size=13, bold=True, color=WHITE)

    tp = ttl_cell.paragraphs[0]
    tr = tp.add_run(title)
    font(tr, size=12, bold=True)

    doc.add_paragraph()

    # Os 4 blocos internos
    blocks = [
        ("O PROBLEMA", problem, "F9F9F9", "AEB4BB"),
        ("O IMPACTO NO NEGOCIO", impact, "FFF8F8", "E31B23"),
        ("A SOLUCAO", solution, "F9F9F9", "AEB4BB"),
        ("O QUE A CAPITOL GANHA", gain, "F2FFF4", "2E8B57"),
    ]

    for label, content, bg, border in blocks:
        t2 = doc.add_table(rows=1, cols=1)
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        c2 = t2.cell(0, 0)
        set_cell_bg(c2, bg)
        set_cell_border(c2, left={"val": "single", "sz": 12, "space": 0, "color": border})
        p2 = c2.paragraphs[0]
        label_run = p2.add_run(label + "\n")
        font(label_run, size=8, bold=True, color=GRAY)
        content_run = p2.add_run(content)
        font(content_run, size=10.5)
        doc.add_paragraph()


def flow_step(doc, number, label, description):
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    n_cell = t.rows[0].cells[0]
    l_cell = t.rows[0].cells[1]
    d_cell = t.rows[0].cells[2]
    n_cell.width = Cm(1.0)
    l_cell.width = Cm(4.5)
    d_cell.width = Cm(10.7)

    set_cell_bg(n_cell, "E31B23")
    np = n_cell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run(number)
    font(nr, size=11, bold=True, color=WHITE)

    set_cell_bg(l_cell, "0A0A0A")
    lp = l_cell.paragraphs[0]
    lr = lp.add_run(label)
    font(lr, size=10, bold=True, color=WHITE)

    dp = d_cell.paragraphs[0]
    dr = dp.add_run(description)
    font(dr, size=10)
    doc.add_paragraph()


def priority_block(doc, phase, bg, items, summary):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_cell_bg(c, bg)
    set_cell_border(c,
                    top={"val": "single", "sz": 6, "space": 0, "color": "CCCCCC"},
                    bottom={"val": "single", "sz": 6, "space": 0, "color": "CCCCCC"})
    p = c.paragraphs[0]
    r = p.add_run(phase)
    font(r, size=11, bold=True)
    doc.add_paragraph()

    for item in items:
        bullet(doc, item)
    callout(doc, summary, bg="F7F7F7", border_color="888888")


def add_signature(doc):
    p = doc.add_paragraph()
    r = p.add_run(f"Sao Paulo, {DATE}.")
    font(r, size=11, bold=True)

    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for cell in row.cells:
            cell.width = Cm(7.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            set_cell_border(cell, top={"val": "single", "sz": 8, "space": 0, "color": "AEB4BB"})

    lp = t.cell(0, 0).paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run("LOCALRISE ADVISORY\n").bold = True
    lp.add_run("Nicolas Slater\n")
    lp.add_run("Diretoria Comercial")

    rp = t.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.add_run(f"{COMPANY}\n").bold = True
    rp.add_run("Responsavel: ______________________\n")
    rp.add_run("Cargo: ___________________________")


# ── conteudo ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    apply_doc_style(doc)

    # CAPA
    add_brand_header(doc)
    doc.add_page_break()

    # ── ABERTURA ────────────────────────────────────────────────────────────────
    section_title(doc, "Como usar este material na reuniao")
    justify(doc,
        "Este documento foi criado para ser a base de uma conversa estrategica com a Capitol Consultoria. "
        "Nao e um relatorio para ser lido. E um roteiro para ser explicado. "
        "Cada bloco mostra um problema real, o que isso custa para o negocio e como existe uma solucao clara para isso.")
    callout(doc,
        "A logica e simples: mostrar que o problema existe, que tem impacto real e que ha um caminho estruturado. "
        "O cliente nao precisa entender de tecnologia — precisa entender de negocio.",
        bg="F0F4FF", border_color="3A6FD8")
    doc.add_page_break()

    # ── BLOCO 1: CONTEXTO ───────────────────────────────────────────────────────
    section_title(doc, "Onde a Capitol esta hoje")
    justify(doc,
        "A Capitol tem algo que a maioria das empresas leva anos para construir: credibilidade real, senioridade e um tema "
        "que o mercado precisa. O problema nao e o produto. O problema e que esse valor nao chega ate quem ainda nao conhece a empresa.")
    justify(doc,
        "Quando fizemos o diagnostico, encontramos tres realidades concretas. Nao sao criticas — sao oportunidades. "
        "E e exatamente aqui que a conversa comeca.")

    # os 3 problemas principais — formato simples para falar em reuniao
    problems = [
        ("Nao ha porta de entrada digital",
         "Quando alguem ouve falar da Capitol e vai pesquisar — pelo Google, pelo Instagram — nao encontra uma estrutura "
         "que transmita confianca e convide ao contato. O lead esfria antes de entrar em contato. "
         "Isso acontece todos os dias, sem que ninguem perceba."),
        ("Nao ha sequencia de contato",
         "Quando alguem entra em contato, o que acontece depois depende da memoria e disponibilidade da equipe. "
         "Nao ha primeira resposta automatica, nao ha acompanhamento estruturado. "
         "Um lead que nao recebe resposta rapida perde interesse — e vai embora silenciosamente."),
        ("Nao ha acompanhamento pos-reuniao",
         "Depois de uma reuniao de apresentacao ou diagnostico, o processo de fechamento e totalmente manual. "
         "Negociacoes que poderiam fechar em 1 semana se arrastam por meses — ou simplesmente desaparecem do radar."),
    ]

    for i, (title, desc) in enumerate(problems, 1):
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        nc = t.rows[0].cells[0]
        dc = t.rows[0].cells[1]
        nc.width = Cm(1.2)
        dc.width = Cm(15.0)
        set_cell_bg(nc, "E31B23")
        np2 = nc.paragraphs[0]
        np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr2 = np2.add_run(str(i))
        font(nr2, size=13, bold=True, color=WHITE)

        dp2 = dc.paragraphs[0]
        tr2 = dp2.add_run(title + "\n")
        font(tr2, size=11, bold=True)
        dr2 = dp2.add_run(desc)
        font(dr2, size=10.5, color=GRAY)
        doc.add_paragraph()

    callout(doc,
        "A consequencia direta desses tres problemas e uma so: a Capitol depende de indicacao para crescer. "
        "E quem depende de indicacao nao controla o proprio crescimento.",
        bg="FDF3F3")
    doc.add_page_break()

    # ── BLOCO 2: AS AUTOMACOES ──────────────────────────────────────────────────
    section_title(doc, "O que precisa existir — e por que")
    justify(doc,
        "Automacao, na pratica, nao e tecnologia. E processo que roda sozinho. "
        "E o que garante que toda pessoa que demonstrar interesse em trabalhar com a Capitol "
        "seja atendida, qualificada e acompanhada — sem depender de memoria ou disponibilidade da equipe.")
    justify(doc,
        "Cada automacao abaixo resolve um problema real identificado no diagnostico. "
        "A logica de cada uma e sempre a mesma: qual e o problema, o que custa nao resolver, "
        "qual e a solucao e o que a Capitol ganha.")
    doc.add_paragraph()

    automations = [
        (
            "01", "Primeiro contato automatico",
            "Quando alguem envia mensagem pelo WhatsApp ou preenche o formulario do site, "
            "nao ha garantia de resposta rapida. O lead fica esperando.",
            "Um lead que nao recebe resposta em ate 5 minutos tem 10 vezes menos chance de avancar. "
            "Cada hora de silencio e uma oportunidade que esfria.",
            "Assim que o lead entrar em contato, o sistema envia automaticamente uma mensagem de apresentacao, "
            "confirma o recebimento e informa o proximo passo — 24 horas por dia, sem intervencao da equipe.",
            "A Capitol passa a ter atendimento imediato para o primeiro contato. Nenhum lead esfria por falta de resposta rapida. "
            "A percepcao de profissionalismo aumenta desde o primeiro segundo.",
        ),
        (
            "02", "Captura de lead com registro automatico",
            "Hoje, quando alguem demonstra interesse, nao ha registro sistematico. O historico fica na memoria "
            "ou em conversas de WhatsApp espalhadas.",
            "Sem registro, nao ha visibilidade. Sem visibilidade, nao ha gestao. "
            "Leads se perdem sem que ninguem perceba — e a Capitol nao tem como saber quantas oportunidades passaram.",
            "Quando o lead preenche o formulario do site ou entra em contato, o sistema registra automaticamente "
            "nome, empresa, cargo e necessidade em um painel centralizado. A equipe recebe uma notificacao imediata.",
            "A Capitol passa a saber, em tempo real, quem chegou, quando chegou e o que precisa. "
            "Zero perda de lead por falta de registro. A base de contatos cresce de forma organizada.",
        ),
        (
            "03", "Qualificacao antes da reuniao",
            "Hoje qualquer pessoa que entrar em contato pode virar uma reuniao. "
            "Nao ha filtro. Nao ha triagem. O tempo dos consultores seniors e alocado sem criterio.",
            "Uma reuniao com alguem sem perfil custa caro: tempo de preparacao, deslocamento ou videoconferencia, "
            "followup posterior. E tempo que poderia estar em uma oportunidade real.",
            "Apos o primeiro contato, o sistema envia automaticamente 2 ou 3 perguntas simples por WhatsApp: "
            "qual o setor, qual o desafio atual, qual o porte da empresa. Com base nas respostas, o lead e classificado "
            "como pronto para reuniao, em amadurecimento ou fora do perfil.",
            "A equipe so vai a reuniao com quem tem perfil. O tempo dos consultores e protegido. "
            "A taxa de fechamento de reunioes realizadas sobe porque o publico ja foi filtrado.",
        ),
        (
            "04", "Agendamento sem troca de mensagens",
            "Hoje combinar uma reuniao exige troca de varias mensagens — para alinhar datas, horarios, formato. "
            "Esse vai-e-vem e lento e cria friccao desnecessaria.",
            "Cada passo de friccao no processo de agendamento aumenta a chance do lead desistir ou perder o interesse. "
            "Um prospect que esta quente hoje pode esfriar em 48 horas de silencio.",
            "Quando o lead e qualificado, recebe automaticamente um link de agendamento com os horarios disponiveis do consultor. "
            "Ele escolhe, confirma e recebe os detalhes por WhatsApp e e-mail — sem nenhuma intervencao da equipe.",
            "O agendamento acontece em minutos. A reuniao e marcada enquanto o interesse ainda esta alto. "
            "O consultor ja recebe a reuniao organizada no calendario.",
        ),
        (
            "05", "Preparo do lead antes da reuniao",
            "O lead chega na reuniao sem ter absorvido nada sobre a Capitol. "
            "O consultor passa os primeiros 15 minutos apresentando a empresa — tempo que poderia ser mais estrategico.",
            "Uma reuniao de apresentacao tem menor taxa de fechamento do que uma reuniao de diagnostico. "
            "Quando o lead ja conhece a Capitol, a conversa comeca mais avancada e vai mais longe.",
            "Apos o agendamento confirmado, o sistema envia automaticamente um material curto sobre a Capitol: "
            "quem somos, como pensamos, um insight sobre relacoes governamentais. O lead chega preparado.",
            "A reuniao comeca com o lead ja convencido da credibilidade da Capitol. "
            "O consultor usa o tempo para diagnosticar e construir proposta — nao para apresentar. "
            "O ciclo de fechamento encurta.",
        ),
        (
            "06", "Acompanhamento automatico pos-reuniao",
            "Depois de uma reuniao, o followup depende completamente da proatividade da equipe. "
            "Quando a agenda aperta, o acompanhamento atrasa — e o lead esfria.",
            "Estudos de vendas B2B mostram que mais de 60% das negociacoes que nao fecham e por falta de followup adequado. "
            "Nao e que o cliente nao queria. E que ninguem foi atras.",
            "Em ate 2 horas apos a reuniao, o sistema envia automaticamente uma mensagem de agradecimento com o resumo "
            "do que foi discutido e o proximo passo. Nos dias 3, 7 e 14, envia mensagens de acompanhamento caso nao haja resposta.",
            "Nenhuma negociacao cai por esquecimento. O lead se sente acompanhado e tem mais confianca para decidir. "
            "A taxa de fechamento de reunioes realizadas aumenta de forma consistente.",
        ),
        (
            "07", "Reativacao de quem sumiu",
            "Leads que chegaram, demonstraram interesse, mas pararam em algum ponto do processo — "
            "ficam esquecidos. Nao ha reativacao. Nao ha segundo momento.",
            "Esses leads ja conhecem a Capitol. Ja tiveram algum nivel de interesse. "
            "Reativá-los custa muito menos do que capturar um lead novo — mas hoje essa receita e simplesmente perdida.",
            "Leads sem interacao ha 30 dias recebem automaticamente uma mensagem curta e sem pressao: "
            "uma pergunta sobre se o momento mudou, ou um insight relevante sobre o cenario atual. "
            "Se responderem, o fluxo de qualificacao e reativado.",
            "Entre 15% e 30% dos leads reativados convertem. E receita gerada a custo zero — "
            "sem precisar capturar nenhum lead novo. A Capitol aproveita o investimento ja feito.",
        ),
        (
            "08", "Nutricao de quem ainda nao esta pronto",
            "A demanda por relacoes governamentais e sazonal — ela aparece quando algo urgente acontece. "
            "Decisores que nao tem urgencia hoje podem ter amanha. Mas se a Capitol nao estiver presente, "
            "outro nome vai estar.",
            "Um lead que conhece a Capitol hoje mas nao fecha, se nao for nutrido, vai esquecer a empresa em 90 dias. "
            "Quando a urgencia chegar, vai procurar quem estiver mais presente na memoria.",
            "Leads que estao em amadurecimento recebem, a cada 15 dias, um conteudo de autoridade por e-mail ou WhatsApp: "
            "leitura de cenario, alerta regulatorio, insight sobre Brasilia. Sem pressao de venda — apenas presenca.",
            "A Capitol e lembrada quando o momento certo chegar. "
            "O lead que nao estava pronto volta como cliente qualificado — ja com relacao estabelecida com a marca.",
        ),
    ]

    for args in automations:
        problem_solution_card(doc, *args)
        divider(doc)

    doc.add_page_break()

    # ── BLOCO 3: FLUXO COMPLETO ─────────────────────────────────────────────────
    section_title(doc, "Como tudo isso se conecta — o fluxo completo")
    justify(doc,
        "Isoladas, cada automacao ja resolve um problema. Conectadas, elas formam um sistema de crescimento. "
        "Abaixo esta a jornada completa de um lead — do primeiro contato ao fechamento — com automacao em cada etapa.")
    doc.add_paragraph()

    steps = [
        ("1", "DESCOBERTA",
         "O lead ve um post no Instagram, clica em um anuncio ou ouve uma indicacao. "
         "Vai ao site e encontra uma estrutura profissional com CTA claro."),
        ("2", "PRIMEIRO CONTATO",
         "Preenche o formulario ou envia mensagem no WhatsApp. "
         "Em segundos, recebe confirmacao e apresentacao automatica da Capitol."),
        ("3", "QUALIFICACAO",
         "Responde 2 ou 3 perguntas rapidas. O sistema classifica o perfil "
         "e direciona: reuniao imediata, nutricao ou reativacao futura."),
        ("4", "AGENDAMENTO",
         "Recebe link de agendamento. Em minutos, marca a reuniao sozinho. "
         "Recebe material de preparo antes do encontro."),
        ("5", "REUNIAO",
         "Chega preparado e com percepcao de autoridade formada. "
         "O consultor foca em diagnostico e proposta — nao em apresentacao."),
        ("6", "FOLLOWUP",
         "Recebe mensagem de agradecimento automatica em 2 horas. "
         "Nos dias seguintes, acompanhamento estruturado sem depender de memoria."),
        ("7", "FECHAMENTO OU NUTRICAO",
         "Se fechar: entra no fluxo de pos-venda e solicitacao de depoimento. "
         "Se nao fechar: entra na nutricao de longo prazo e pode converter no momento certo."),
    ]

    for args in steps:
        flow_step(doc, *args)

    callout(doc,
        "Cada etapa elimina um dos gargalos identificados no diagnostico. "
        "O resultado e um sistema onde nenhum lead se perde — e a Capitol cresce com previsibilidade.",
        bg="F0FFF4", border_color="2E8B57")
    doc.add_page_break()

    # ── BLOCO 4: PRIORIDADES ────────────────────────────────────────────────────
    section_title(doc, "Por onde comecar — ordem de implementacao")
    justify(doc,
        "Nao e necessario implementar tudo de uma vez. A logica e construir por camadas: "
        "primeiro o que e essencial para nao perder leads, depois o que melhora a conversao, "
        "depois o que escala o sistema.")
    doc.add_paragraph()

    priority_block(doc,
        "FASE 1 — Primeiros 30 dias | Nao perder mais nenhum lead",
        "E8F0FE",
        [
            "Site com formulario ativo — sem isso, nenhuma automacao funciona",
            "Resposta automatica no WhatsApp — primeiro contato imediato",
            "Registro automatico do lead no painel central",
            "Agendamento automatico com link de calendario",
            "Envio de material de preparo antes da reuniao",
        ],
        "Resultado esperado: a Capitol deixa de perder leads no primeiro contato e ganha controle "
        "total sobre quem chegou e o que aconteceu depois.")

    priority_block(doc,
        "FASE 2 — De 30 a 60 dias | Converter mais o que ja chega",
        "FFF9E8",
        [
            "Qualificacao automatica — filtrar antes de alocar tempo senior",
            "Followup automatico pos-reuniao — sustentar o ciclo comercial",
            "Reativacao de leads inativos — aproveitar a base ja construida",
            "Nutricao de longo prazo — manter presenca com quem ainda nao esta pronto",
        ],
        "Resultado esperado: o funil completo esta operando. "
        "Nenhuma oportunidade cai por falta de acompanhamento.")

    priority_block(doc,
        "FASE 3 — De 60 a 90 dias | Escalar com controle",
        "F0FFF4",
        [
            "Prova social automatizada — depoimentos viram ativo sistematico",
            "Relatorio semanal do funil — decisoes com base em dados reais",
            "Meta Ads integrado — trafego pago com base de conversao pronta",
        ],
        "Resultado esperado: o sistema e completo, mensuravel e preparado para escalar. "
        "A Capitol deixa de depender de indicacao e passa a ter aquisicao previsivel.")

    doc.add_page_break()

    # ── BLOCO 5: IMPACTO FINAL ──────────────────────────────────────────────────
    section_title(doc, "O que muda para a Capitol com esse sistema")
    justify(doc,
        "Este e o bloco mais importante. Nao estamos falando de tecnologia. Estamos falando do que muda, "
        "na pratica, para o negocio.")

    impacts = [
        ("Menos dependencia de indicacao",
         "Hoje o crescimento da Capitol depende de quem indica. Com o sistema, ha uma fonte previsivel "
         "de leads qualificados: Instagram no topo, site convertendo no meio, automacoes sustentando o fundo. "
         "A empresa passa a controlar o proprio crescimento."),
        ("Mais leads qualificados chegando",
         "Com site ativo, formulario funcionando e Instagram estruturado, a estimativa conservadora "
         "e de 8 a 20 novos contatos qualificados por mes. Com Meta Ads ativado, esse numero pode triplicar em 90 dias."),
        ("Taxa de fechamento maior",
         "Leads que chegam preparados fecham mais rapido. Followup estruturado fecha mais. "
         "Prova social reduz objecao. Uma consultoria B2B com sistema bem estruturado costuma ver "
         "a taxa de fechamento dobrar em 90 dias de operacao — sem contratar ninguem."),
        ("Previsibilidade de receita",
         "Com relatorio semanal do funil, a lideranca da Capitol consegue ver, com clareza, "
         "quantas reunioes estao marcadas, quantas propostas estao abertas e qual e a probabilidade de fechamento. "
         "A gestao comercial passa de reativa para proativa."),
    ]

    for title, desc in impacts:
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        ic = t.rows[0].cells[0]
        dc2 = t.rows[0].cells[1]
        ic.width = Cm(0.6)
        dc2.width = Cm(15.6)
        set_cell_bg(ic, "E31B23")
        ip = ic.paragraphs[0]
        ir = ip.add_run(" ")
        font(ir, size=11)

        dp3 = dc2.paragraphs[0]
        tr3 = dp3.add_run(title + "\n")
        font(tr3, size=11, bold=True)
        dr3 = dp3.add_run(desc)
        font(dr3, size=10.5)
        doc.add_paragraph()

    callout(doc,
        "A pergunta nao e se a Capitol precisa desse sistema. A pergunta e quanto custa continuar sem ele. "
        "Cada mes sem estrutura e uma lista de leads que chegaram, esfriaram e foram embora sem que ninguem percebesse.",
        bg="FDF3F3")

    doc.add_page_break()

    # ── BLOCO 6: PROXIMOS PASSOS ────────────────────────────────────────────────
    section_title(doc, "Proximos passos recomendados")
    justify(doc,
        "O caminho mais direto para a Capitol sair da dependencia de indicacao e comecar pela base: "
        "ter uma porta de entrada digital que funcione. Tudo o mais e construido em cima disso.")
    doc.add_paragraph()

    next_steps = [
        ("Passo 1", "Ativar o site com formulario de captura — essa e a fundacao de tudo"),
        ("Passo 2", "Configurar a resposta automatica no WhatsApp — primeiro contato imediato"),
        ("Passo 3", "Definir o CRM — painel central onde todos os leads ficam registrados"),
        ("Passo 4", "Ativar os primeiros fluxos automaticos — agendamento e preparo pre-reuniao"),
        ("Passo 5", "Revisar resultados em 30 dias e ativar a Fase 2 — followup e qualificacao"),
    ]

    for label, desc in next_steps:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r_label = p.add_run(label + ": ")
        font(r_label, size=11, bold=True, color=RED)
        r_desc = p.add_run(desc)
        font(r_desc, size=11)

    doc.add_paragraph()
    callout(doc,
        "A LocalRise nao entrega uma lista de ferramentas. Entrega o sistema funcionando — "
        "da captura do primeiro lead ao acompanhamento do fechamento. "
        "A equipe da Capitol faz o que faz de melhor. O sistema cuida do resto.",
        bg="F0F4FF", border_color="3A6FD8")

    doc.add_paragraph()
    add_signature(doc)

    out = BASE_DIR / "estudo-automacoes-capitol-reuniao.docx"
    doc.save(out)
    print(f"Documento salvo em: {out}")


if __name__ == "__main__":
    build()
