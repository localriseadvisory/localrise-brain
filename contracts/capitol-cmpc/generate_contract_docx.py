from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "contrato_capitol_cmpc.docx"
LOGO_PATH = Path(r"C:\Users\digui\Downloads\1Artboard 1@4x.png")

CONTRACTOR_NAME = "CAPITOL CONSULTORIA LTDA"
CONTRACTEE_NAME = "CMPC"
DATE_LINE = "Brasilia, 07 de abril de 2026."
TITLE = "CONTRATO DE PRESTACAO DE SERVICOS DE ASSESSORIA EM RELACOES GOVERNAMENTAIS"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:%s" % edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:%s" % key), str(edge_data[key]))


def apply_base_style(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(8)


def add_runs(paragraph, chunks):
    for text, bold in chunks:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_body_paragraph(document: Document, chunks):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(8)
    add_runs(paragraph, chunks)
    return paragraph


def add_clause_title(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(12)
    return paragraph


def add_roman_list(document: Document, items):
    numerals = [
        "I",
        "II",
        "III",
        "IV",
        "V",
        "VI",
        "VII",
        "VIII",
        "IX",
        "X",
    ]
    for index, item in enumerate(items):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        add_runs(paragraph, [(f"{numerals[index]} - {item}", False)])


def add_header_block(document: Document):
    if LOGO_PATH.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(6.5))

    company = document.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company.paragraph_format.space_after = Pt(4)
    run = company.add_run(CONTRACTOR_NAME)
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(12)

    line_table = document.add_table(rows=1, cols=1)
    line_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    line_table.autofit = False
    cell = line_table.cell(0, 0)
    cell.width = Cm(16)
    set_cell_border(cell, bottom={"val": "single", "sz": 6, "space": 0, "color": "B7B7B7"})
    cell.text = ""

    document.add_paragraph()


def add_title(document: Document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(TITLE)
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(13)


def add_signatures(document: Document):
    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_paragraph.paragraph_format.space_before = Pt(12)
    add_runs(date_paragraph, [(DATE_LINE, True)])

    document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ASSINATURAS")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(12)

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(7.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            set_cell_border(cell, top={"val": "single", "sz": 8, "space": 0, "color": "7A7A7A"})

    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(
        left,
        [
            (f"{CONTRACTEE_NAME}\n", True),
            ("Nome: ________________________________\n", False),
            ("Cargo: ________________________________", False),
        ],
    )

    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(
        right,
        [
            (f"{CONTRACTOR_NAME}\n", True),
            ("Luiz Carlos Saraiva\n", False),
            ("Socio Diretor", False),
        ],
    )

    document.add_paragraph()

    witness_table = document.add_table(rows=1, cols=2)
    witness_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    witness_table.autofit = False
    for row in witness_table.rows:
        for cell in row.cells:
            cell.width = Cm(7.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            set_cell_border(cell, top={"val": "single", "sz": 8, "space": 0, "color": "7A7A7A"})

    witness_1 = witness_table.cell(0, 0).paragraphs[0]
    witness_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(
        witness_1,
        [
            ("TESTEMUNHA 1\n", True),
            ("Nome: ________________________________\n", False),
            ("CPF: __________________________________", False),
        ],
    )

    witness_2 = witness_table.cell(0, 1).paragraphs[0]
    witness_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(
        witness_2,
        [
            ("TESTEMUNHA 2\n", True),
            ("Nome: ________________________________\n", False),
            ("CPF: __________________________________", False),
        ],
    )


def build_document():
    document = Document()
    apply_base_style(document)
    add_header_block(document)
    add_title(document)

    add_body_paragraph(
        document,
        [
            ("Pelo presente instrumento particular, de um lado, ", False),
            (f"{CONTRACTEE_NAME}", True),
            (
                ", pessoa juridica inscrita no CNPJ sob o no ____________, com sede em "
                "______________________________, neste ato representada por ______________________________, "
                "na qualidade de ______________________________, doravante denominada ",
                False,
            ),
            ("CONTRATANTE", True),
            (
                "; e, de outro lado, ",
                False,
            ),
            (CONTRACTOR_NAME, True),
            (
                ", pessoa juridica de direito privado, inscrita no CNPJ sob o no "
                "10.248.895/0001-31, com sede no Setor de Radio e TV Sul (SRTVS), Quadra 701, Conjunto L, "
                "no 38, Bloco 01, Sala 608, Edificio Assis Chateaubriand, Asa Sul, Brasilia/DF, CEP 70.340-906, "
                "neste ato representada por seu Socio Diretor, ",
                False,
            ),
            ("Luiz Carlos Saraiva", True),
            (", doravante denominada ", False),
            ("CONTRATADA", True),
            (
                ", tem entre si justo e acordado o presente contrato de prestacao de servicos, que se regera "
                "pelas clausulas e condicoes seguintes.",
                False,
            ),
        ],
    )

    add_body_paragraph(
        document,
        [
            (
                "Considerando a relevancia institucional e regulatoria da atuacao da CONTRATANTE no Brasil, "
                "bem como a necessidade de acompanhamento qualificado e permanente de agendas estrategicas em "
                "Brasilia, as partes resolvem firmar o presente instrumento para disciplinar a assessoria em "
                "relacoes governamentais e institucionais.",
                False,
            )
        ],
    )

    add_clause_title(document, "CLAUSULA PRIMEIRA - DO OBJETO")
    add_body_paragraph(
        document,
        [
            (
                "O presente contrato tem por objeto a prestacao, pela CONTRATADA, de servicos especializados "
                "de assessoria em relacoes governamentais e institucionais, com atuacao em Brasilia/DF, voltados "
                "a antecipacao de cenarios, defesa de interesses legitimos da CONTRATANTE e fortalecimento do "
                "dialogo com o Poder Publico.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "A atuacao da CONTRATADA constitui obrigacao de meio, nao implicando garantia de resultado "
                "especifico, aprovacao de medidas, edicao de atos normativos ou deliberacoes por terceiros.",
                False,
            )
        ],
    )

    add_clause_title(document, "CLAUSULA SEGUNDA - DO ESCOPO DOS SERVICOS")
    add_body_paragraph(document, [("Os servicos compreenderao, de forma nao exaustiva:", False)])
    add_roman_list(
        document,
        [
            "monitoramento do Poder Executivo Federal, do Congresso Nacional e de temas regulatorios de interesse da CONTRATANTE;",
            "gestao e acompanhamento de stakeholders estrategicos;",
            "atuacao institucional junto a orgaos federais, ministerios, autarquias, agencias e demais entidades publicas pertinentes;",
            "elaboracao de notas tecnicas, subsidios, memorandos executivos e sugestoes de posicionamento institucional, quando demandado;",
            "estruturacao e manutencao de agenda institucional em Brasilia/DF;",
            "acompanhamento de reunioes, audiencias e interlocucoes estrategicas relacionadas aos interesses da CONTRATANTE;",
            "emissao de relatorios estrategicos e atualizacoes executivas sobre temas regulatorios, tributarios e setoriais;",
            "apoio institucional relacionado ao ambiente de negocios da CONTRATANTE no setor de celulose, papel, base florestal e economia renovavel.",
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "Paragrafo unico. Nao se incluem no escopo, salvo ajuste escrito especifico entre as partes, "
                "a execucao de servicos juridicos contenciosos, a emissao de parecer juridico formal, a "
                "representacao judicial ou administrativa com poderes de mandato e atos privativos de outras "
                "profissoes regulamentadas.",
                False,
            )
        ],
    )

    add_clause_title(document, "CLAUSULA TERCEIRA - DA METODOLOGIA")
    add_body_paragraph(
        document,
        [
            (
                "A CONTRATADA prestara os servicos de forma personalizada, com presenca ativa em Brasilia/DF, "
                "articulacao institucional de alto nivel e alinhamento continuo com a CONTRATANTE, observadas "
                "as prioridades estrategicas definidas em comum acordo.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "As partes poderao realizar reunioes presenciais ou remotas para avaliacao de cenario, definicao "
                "de posicionamento institucional e priorizacao de temas, agendas e interlocucoes.",
                False,
            )
        ],
    )

    add_clause_title(document, "CLAUSULA QUARTA - DA VIGENCIA")
    add_body_paragraph(
        document,
        [
            ("O presente contrato vigorara da data de sua assinatura ate ", False),
            ("31 de dezembro de 2026", True),
            (", podendo ser prorrogado mediante termo aditivo escrito firmado entre as partes.", False),
        ],
    )

    add_clause_title(document, "CLAUSULA QUINTA - DO VALOR E DAS CONDICOES DE PAGAMENTO")
    add_body_paragraph(
        document,
        [
            ("Pelos servicos objeto deste contrato, a CONTRATANTE pagara a CONTRATADA o valor mensal de ", False),
            ("R$ 74.000,00 (setenta e quatro mil reais)", True),
            (", mediante apresentacao da respectiva nota fiscal.", False),
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "Paragrafo primeiro. O pagamento devera ser efetuado ate o ____ dia util de cada mes, por meio "
                "de deposito, transferencia bancaria ou outro meio previamente ajustado entre as partes.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "Paragrafo segundo. Os tributos incidentes sobre a prestacao dos servicos observarao a legislacao "
                "aplicavel, cabendo a cada parte o cumprimento das obrigacoes legais a ela atribuidas.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "Paragrafo terceiro. Na hipotese de atraso no pagamento, incidirao correcao monetaria, juros "
                "legais e demais encargos cabiveis, observada a legislacao vigente.",
                False,
            )
        ],
    )

    add_clause_title(document, "CLAUSULA SEXTA - DAS OBRIGACOES DA CONTRATADA")
    add_roman_list(
        document,
        [
            "executar os servicos com zelo, diligencia, etica, discricao e observancia da legislacao aplicavel;",
            "atuar tecnicamente no campo institucional, mantendo alinhamento previo com a CONTRATANTE quanto a pautas sensiveis, posicionamentos e prioridades;",
            "manter a CONTRATANTE informada, em nivel executivo, sobre fatos relevantes ligados ao objeto deste contrato;",
            "resguardar o sigilo das informacoes estrategicas, empresariais, regulatorias e institucionais a que tiver acesso;",
            "abster-se de praticar qualquer ato que possa comprometer a imagem, a reputacao ou os interesses legitimos da CONTRATANTE.",
        ],
    )

    add_clause_title(document, "CLAUSULA SETIMA - DAS OBRIGACOES DA CONTRATANTE")
    add_roman_list(
        document,
        [
            "fornecer, em tempo habil, as informacoes, documentos e diretrizes necessarias a adequada execucao dos servicos;",
            "designar interlocutor ou interlocutores para alinhamento das demandas estrategicas e tomada de decisao;",
            "efetuar os pagamentos nos prazos e condicoes estabelecidos neste contrato;",
            "manifestar-se, sempre que necessario, sobre posicionamentos institucionais, prioridades, agendas e encaminhamentos relevantes.",
        ],
    )

    add_clause_title(document, "CLAUSULA OITAVA - DA CONFIDENCIALIDADE")
    add_body_paragraph(
        document,
        [
            (
                "As partes comprometem-se a manter, sob absoluto sigilo, todas as informacoes, documentos, dados, "
                "estrategias, materiais, planos, posicionamentos, analises, agendas, tratativas e demais conteudos "
                "a que tiverem acesso em razao deste contrato, independentemente do meio em que se encontrem registrados.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "A obrigacao de confidencialidade abrange, inclusive, informacoes obtidas verbalmente, documentos "
                "preparatorios, materiais internos, informacoes de negocio, dados pessoais eventualmente tratados "
                "e quaisquer informacoes que, por sua natureza ou contexto, devam ser consideradas reservadas.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "A obrigacao prevista nesta clausula permanecera valida durante toda a vigencia contratual e por ",
                False,
            ),
            ("5 (cinco) anos", True),
            (
                " apos o seu termino, ressalvadas apenas as hipoteses de divulgacao exigida por lei, ordem de "
                "autoridade competente ou informacao que ja seja comprovadamente publica sem violacao deste instrumento.",
                False,
            ),
        ],
    )

    add_clause_title(document, "CLAUSULA NONA - DO COMPLIANCE E INTEGRIDADE")
    add_body_paragraph(
        document,
        [
            (
                "As partes declaram que conduzirao sua atuacao no ambito deste contrato em estrita observancia "
                "as normas de etica, integridade, transparencia, prevencao a corrupcao, prevencao a conflitos "
                "de interesse e conformidade legal aplicaveis as suas atividades, incluindo, quando cabivel, a "
                "Lei no 12.846/2013 e demais normas correlatas.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "A CONTRATADA compromete-se a nao prometer, oferecer, autorizar ou conceder vantagem indevida "
                "a agente publico ou a terceiro a ele relacionado, tampouco adotar conduta incompativel com "
                "padroes legitimos de representacao institucional.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "A ocorrencia de violacao grave comprovada a esta clausula autorizara a parte inocente a rescindir "
                "imediatamente o contrato, sem prejuizo da apuracao de perdas e danos e das responsabilidades legais cabiveis.",
                False,
            )
        ],
    )

    add_clause_title(document, "CLAUSULA DECIMA - DA RESCISAO")
    add_body_paragraph(
        document,
        [
            (
                "O presente contrato podera ser rescindido, imotivadamente, por qualquer das partes, mediante "
                "comunicacao escrita a outra parte com antecedencia minima de ",
                False,
            ),
            ("30 (trinta) dias", True),
            (".", False),
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "Em qualquer hipotese de rescisao, permanecerao exigiveis os valores devidos ate a data da efetiva "
                "extincao contratual, assim como as obrigacoes de confidencialidade, integridade e demais disposicoes "
                "que, por sua natureza, devam subsistir ao termino do contrato.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "O descumprimento contratual relevante, nao sanado no prazo razoavel indicado pela parte adimplente, "
                "tambem podera ensejar a rescisao motivada, sem prejuizo das medidas cabiveis.",
                False,
            )
        ],
    )

    add_clause_title(document, "CLAUSULA DECIMA PRIMEIRA - DA INEXISTENCIA DE VINCULO")
    add_body_paragraph(
        document,
        [
            (
                "O presente contrato nao estabelece, nem podera ser interpretado como gerador de vinculo "
                "empregaticio, societario, associativo, representativo com poderes permanentes, mandato, "
                "exclusividade ou subordinacao entre as partes, seus socios, empregados, prepostos ou colaboradores.",
                False,
            )
        ],
    )
    add_body_paragraph(
        document,
        [
            (
                "Cada parte permanecera integralmente responsavel por sua estrutura organizacional e por seus "
                "encargos trabalhistas, previdenciarios, fiscais e operacionais, nao havendo solidariedade ou "
                "subsidiariedade entre elas, salvo imposicao legal.",
                False,
            )
        ],
    )

    add_clause_title(document, "CLAUSULA DECIMA SEGUNDA - DO FORO")
    add_body_paragraph(
        document,
        [
            (
                "Fica eleito o foro da Comarca de Brasilia/DF, com renuncia expressa a qualquer outro, por mais "
                "privilegiado que seja, para dirimir quaisquer controversias oriundas deste contrato.",
                False,
            )
        ],
    )

    add_signatures(document)
    document.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
