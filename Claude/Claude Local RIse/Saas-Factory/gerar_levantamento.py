"""
Gerador do Levantamento Estratégico — LocalRise Advisory
Gera um arquivo .docx com design fiel ao HTML original.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paleta de cores ──────────────────────────────────────────────────────────
DARK       = RGBColor(0x0F, 0x0F, 0x18)   # fundo escuro
RED        = RGBColor(0xE3, 0x1B, 0x23)   # vermelho LocalRise
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF7, 0xF6, 0xF3)
MED_GRAY   = RGBColor(0xA0, 0xA0, 0xC0)
DIM_GRAY   = RGBColor(0x60, 0x60, 0xA0)
TEXT_DARK  = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_MID   = RGBColor(0x44, 0x44, 0x44)
TEXT_LIGHT = RGBColor(0x88, 0x88, 0x88)
BORDER_CLR = RGBColor(0xE8, 0xE4, 0xDF)
INTRO_BG   = RGBColor(0xF9, 0xF8, 0xF6)

# ── Helpers XML ──────────────────────────────────────────────────────────────
def hex_color(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(rgb))
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, color="E8E4DF", size="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        if val == 'none':
            el.set(qn('w:val'), 'nil')
        elif val:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), size)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_table_width(table, width_cm):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_cm * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def set_para_indent(para, left=0, first=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left))
    if first:
        ind.set(qn('w:firstLine'), str(first))
    pPr.append(ind)

def add_left_border(para, color="E31B23", size="18"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def add_bottom_border_para(para, color="F0EDE8", size="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_shading(para, fill: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(fill))
    pPr.append(shd)

def set_page_margins(section, top=2.0, bottom=2.0, left=2.5, right=2.5):
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)

def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    set_para_spacing(para, 0, 0)

# ── Funções de conteúdo ──────────────────────────────────────────────────────
def add_run_styled(para, text, bold=False, italic=False, size=12,
                   color=TEXT_DARK, font='Segoe UI'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = font
    return run

def add_para(doc, text='', align=WD_ALIGN_PARAGRAPH.LEFT,
             before=0, after=80, bold=False, italic=False,
             size=12, color=TEXT_DARK, font='Georgia'):
    para = doc.add_paragraph()
    para.alignment = align
    set_para_spacing(para, before, after)
    if text:
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color
        run.font.name  = font
    return para

# ── CAPA ─────────────────────────────────────────────────────────────────────
def build_cover(doc):
    # Fundo escuro via tabela de largura total
    tbl = doc.add_table(rows=1, cols=1)
    set_table_width(tbl, 16.5)
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, DARK)
    set_cell_borders(cell, top='none', bottom='none', left='none', right='none')

    def cp(text='', bold=False, italic=False, size=12, color=WHITE,
           font='Segoe UI', align=WD_ALIGN_PARAGRAPH.LEFT,
           before=0, after=60):
        para = cell.add_paragraph()
        para.alignment = align
        set_para_spacing(para, before, after)
        if text:
            r = para.add_run(text)
            r.bold   = bold
            r.italic = italic
            r.font.size  = Pt(size)
            r.font.color.rgb = color
            r.font.name  = font
        return para

    # Remover parágrafo vazio inicial da célula
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    # Eyebrow
    p = cp('DOCUMENTO CONFIDENCIAL', bold=True, size=9, color=RED,
           before=240, after=160)

    # Logo mark + nome
    p_logo = cp(before=0, after=60)
    add_run_styled(p_logo, '  LR  ', bold=True, size=13, color=WHITE,
                   font='Segoe UI')
    add_run_styled(p_logo, '   LocalRise Advisory', bold=True, size=15,
                   color=WHITE, font='Segoe UI')

    # Tagline
    cp('Gestão · Automação · Crescimento', size=10, color=DIM_GRAY,
       font='Segoe UI', before=0, after=260)

    # Título principal
    p_title = cp(before=0, after=60)
    add_run_styled(p_title, 'Levantamento Estratégico\n', size=28,
                   color=WHITE, font='Georgia')
    add_run_styled(p_title, 'Sistema de Gestão e Automação', bold=True,
                   size=28, color=WHITE, font='Georgia')

    # Sub
    cp('Diagnóstico de negócio · Mapeamento de processos · Estruturação de sistema personalizado',
       size=12, color=MED_GRAY, font='Segoe UI', before=80, after=280)

    # Meta row
    p_meta = cp(before=0, after=240)
    add_run_styled(p_meta, 'Preparado por:  ', size=9, color=DIM_GRAY, font='Segoe UI')
    add_run_styled(p_meta, 'LocalRise Advisory', size=11, color=MED_GRAY, font='Segoe UI')
    add_run_styled(p_meta, '     |     Tipo:  ', size=9, color=DIM_GRAY, font='Segoe UI')
    add_run_styled(p_meta, 'Levantamento de Requisitos', size=11, color=MED_GRAY, font='Segoe UI')
    add_run_styled(p_meta, '     |     Data:  ', size=9, color=DIM_GRAY, font='Segoe UI')
    add_run_styled(p_meta, 'Abril de 2026', size=11, color=MED_GRAY, font='Segoe UI')

# ── SEÇÃO HEADER ──────────────────────────────────────────────────────────────
def section_header(doc, number, title):
    # Número (eyebrow vermelho)
    p_num = doc.add_paragraph()
    set_para_spacing(p_num, 280, 40)
    r = p_num.add_run(number.upper())
    r.bold = True
    r.font.size  = Pt(9)
    r.font.color.rgb = RED
    r.font.name  = 'Segoe UI'

    # Título
    p_title = doc.add_paragraph()
    set_para_spacing(p_title, 0, 80)
    r2 = p_title.add_run(title)
    r2.bold = True
    r2.font.size  = Pt(17)
    r2.font.color.rgb = TEXT_DARK
    r2.font.name  = 'Segoe UI'

    # Linha divisória (borda inferior com acento vermelho)
    # Simulamos com parágrafo de borda vermelha
    p_div = doc.add_paragraph()
    set_para_spacing(p_div, 0, 120)
    pPr = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), hex_color(RED))
    pBdr.append(bot)
    pPr.append(pBdr)

# ── INTRO BOX ────────────────────────────────────────────────────────────────
def intro_box(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, 0, 120)
    set_para_indent(para, left=200)
    add_left_border(para, color=hex_color(RED), size='24')
    set_para_shading(para, INTRO_BG)
    r = para.add_run(text)
    r.italic = True
    r.font.size  = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r.font.name  = 'Georgia'

# ── PERGUNTA ─────────────────────────────────────────────────────────────────
def add_question(doc, number, text, answer_lines=2):
    tbl = doc.add_table(rows=1, cols=2)
    set_table_width(tbl, 16.5)
    remove_table_borders(tbl)

    # Col widths
    tbl.columns[0].width = Cm(1.2)
    tbl.columns[1].width = Cm(15.3)

    row = tbl.rows[0]
    row.height = None

    # Círculo numerado (simulado com célula quadrada escura)
    c_num = row.cells[0]
    set_cell_bg(c_num, DARK)
    set_cell_borders(c_num, top='none', bottom='none', left='none', right='none')
    c_num.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for p in c_num.paragraphs:
        p._element.getparent().remove(p._element)
    p_n = c_num.add_paragraph()
    p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_n, 120, 0)
    r = p_n.add_run(str(number))
    r.bold = True
    r.font.size  = Pt(10)
    r.font.color.rgb = WHITE
    r.font.name  = 'Segoe UI'

    # Conteúdo da pergunta
    c_txt = row.cells[1]
    set_cell_borders(c_txt, top='none', left='none', right='none',
                     bottom='single', color="F0EDE8", size="6")
    c_txt.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for p in c_txt.paragraphs:
        p._element.getparent().remove(p._element)

    # Texto da pergunta
    p_q = c_txt.add_paragraph()
    set_para_spacing(p_q, 80, 60)
    r = p_q.add_run(text)
    r.font.size  = Pt(11)
    r.font.color.rgb = TEXT_DARK
    r.font.name  = 'Georgia'

    # Linhas de resposta
    for i in range(answer_lines):
        p_ans = c_txt.add_paragraph()
        after = 80 if i < answer_lines - 1 else 120
        set_para_spacing(p_ans, 0, after)
        pPr = p_ans._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'D1CCC7')
        pBdr.append(bot)
        pPr.append(pBdr)
        r = p_ans.add_run(' ')
        r.font.size = Pt(11)

# ── CHECKBOX GRID ────────────────────────────────────────────────────────────
def add_checkbox_grid(doc, items):
    # 2 colunas
    pairs = [(items[i], items[i+1] if i+1 < len(items) else '') for i in range(0, len(items), 2)]
    tbl = doc.add_table(rows=len(pairs), cols=2)
    set_table_width(tbl, 16.5)
    remove_table_borders(tbl)

    for row_idx, (left, right) in enumerate(pairs):
        row = tbl.rows[row_idx]
        for col_idx, item_text in enumerate([(left, 0), (right, 1)]):
            txt, ci = item_text
            cell = row.cells[ci]
            set_cell_bg(cell, LIGHT_GRAY)
            set_cell_borders(cell, top='none', bottom='none', left='none', right='none')

            for p in cell.paragraphs:
                p._element.getparent().remove(p._element)

            if txt:
                p = cell.add_paragraph()
                set_para_spacing(p, 60, 60)
                set_para_indent(p, left=120)
                add_run_styled(p, '☐  ', size=12, color=TEXT_MID, font='Segoe UI')
                add_run_styled(p, txt, size=11, color=RGBColor(0x33,0x33,0x33), font='Georgia')
            else:
                p = cell.add_paragraph()
                set_para_spacing(p, 60, 60)

    doc.add_paragraph()  # espaço após grid

# ── HIGHLIGHT BOX ────────────────────────────────────────────────────────────
def add_highlight_box(doc, label, question, hint=''):
    tbl = doc.add_table(rows=1, cols=1)
    set_table_width(tbl, 16.5)
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, DARK)
    set_cell_borders(cell, top='none', bottom='none', left='none', right='none')

    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    def cp(text='', bold=False, italic=False, size=12, color=WHITE,
           font='Segoe UI', before=0, after=60):
        para = cell.add_paragraph()
        set_para_spacing(para, before, after)
        set_para_indent(para, left=200)
        if text:
            r = para.add_run(text)
            r.bold   = bold
            r.italic = italic
            r.font.size  = Pt(size)
            r.font.color.rgb = color
            r.font.name  = font
        return para

    cp(label.upper(), bold=True, size=9, color=RED, before=180, after=80)
    cp(question, italic=True, size=14, color=RGBColor(0xE0,0xE0,0xF0),
       font='Georgia', before=0, after=60)

    # Linhas de resposta
    for i in range(3):
        p_ans = cell.add_paragraph()
        set_para_spacing(p_ans, 0, 80)
        set_para_indent(p_ans, left=200)
        pPr = p_ans._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '404060')
        pBdr.append(bot)
        pPr.append(pBdr)
        r = p_ans.add_run(' ')
        r.font.size = Pt(11)

    if hint:
        cp(hint, italic=True, size=10, color=DIM_GRAY, before=80, after=180)

# ── CLOSING BOX ──────────────────────────────────────────────────────────────
def add_closing_box(doc, paragraphs_text):
    tbl = doc.add_table(rows=1, cols=1)
    set_table_width(tbl, 16.5)
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, INTRO_BG)
    set_cell_borders(cell, top='single', bottom='single', left='single', right='single',
                     color=hex_color(BORDER_CLR), size='6')

    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    for text in paragraphs_text:
        para = cell.add_paragraph()
        set_para_spacing(para, 80, 100)
        set_para_indent(para, left=200)
        r = para.add_run(text)
        r.font.size  = Pt(11)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        r.font.name  = 'Georgia'

    # Assinatura
    p_sig = cell.add_paragraph()
    set_para_spacing(p_sig, 160, 80)
    set_para_indent(p_sig, left=200)
    add_run_styled(p_sig, ' LR ', bold=True, size=11, color=WHITE, font='Segoe UI')
    add_run_styled(p_sig, '   LocalRise Advisory', bold=True, size=13,
                   color=TEXT_DARK, font='Segoe UI')

    p_role = cell.add_paragraph()
    set_para_spacing(p_role, 0, 160)
    set_para_indent(p_role, left=200)
    add_run_styled(p_role, '       Gestão · Automação · Crescimento',
                   size=10, color=TEXT_LIGHT, font='Segoe UI')

# ── RODAPÉ ────────────────────────────────────────────────────────────────────
def build_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para, 0, 0)
    add_run_styled(para, 'Documento confidencial — uso interno LocalRise e cliente',
                   size=9, color=TEXT_LIGHT, font='Segoe UI')
    add_run_styled(para, '     |     LocalRise Advisory · localrise.com.br',
                   size=9, color=TEXT_LIGHT, font='Segoe UI')


# ════════════════════════════════════════════════════════════════════════════
#  CONSTRUÇÃO DO DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════
def build_document():
    doc = Document()

    # Margens
    for section in doc.sections:
        set_page_margins(section, top=1.5, bottom=1.5, left=1.5, right=1.5)

    # Estilos base
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)

    # ── CAPA ────────────────────────────────────────────────────────────────
    build_cover(doc)
    page_break(doc)

    # ── INTRODUÇÃO ──────────────────────────────────────────────────────────
    section_header(doc, 'Introdução', 'Por que esse documento existe?')
    intro_box(doc,
        'Antes de construir qualquer sistema, precisamos entender como o seu negócio funciona '
        'de verdade — não do jeito que está no papel, mas do jeito que acontece no dia a dia.')

    add_para(doc,
        'Este documento foi preparado para que você possa nos contar, com suas próprias palavras, '
        'como está sua operação hoje. Com isso, vamos conseguir montar um sistema feito sob medida '
        'para o seu negócio — sem desperdício de tempo e sem precisar refazer depois.',
        size=11, color=TEXT_MID, font='Georgia', before=0, after=100)

    add_para(doc,
        'Não precisa responder tudo de forma técnica ou perfeita. Escreva da forma que for mais '
        'natural para você. Se não souber exatamente uma resposta, coloque uma estimativa ou '
        'explique como funciona atualmente. Qualquer informação já nos ajuda muito.',
        size=11, color=TEXT_MID, font='Georgia', before=0, after=100)

    add_para(doc,
        'Se preferir responder por áudio ou chamada, podemos adaptar. O importante é que o '
        'resultado final funcione exatamente para o que você precisa.',
        size=11, color=TEXT_MID, font='Georgia', before=0, after=120)

    # ── BLOCO 1: VISÃO GERAL ────────────────────────────────────────────────
    section_header(doc, 'Bloco 1', 'Visão Geral do Negócio')
    intro_box(doc,
        'Queremos entender o que você faz, como funciona e de onde vem o dinheiro. '
        'Nada técnico — só nos conte como é o seu dia a dia.')
    add_question(doc, 1,
        'O que sua empresa faz hoje? Descreva com suas palavras o que você vende ou entrega.', 2)
    add_question(doc, 2,
        'Quais são seus principais produtos ou serviços? (Liste os mais importantes.)', 2)
    add_question(doc, 3,
        'Como você ganha dinheiro hoje? (Exemplo: venda direta, recorrência, marketplace, atacado…)', 1)

    # ── BLOCO 2: PROCESSO ATUAL ─────────────────────────────────────────────
    section_header(doc, 'Bloco 2', 'Como você Gerencia Hoje?')
    intro_box(doc,
        'Queremos entender como a operação funciona na prática agora — '
        'sem julgar, só para entender o ponto de partida.')
    add_question(doc, 4,
        'Como você controla seu financeiro hoje? (Planilha, sistema, caderno, aplicativo…)', 1)
    add_question(doc, 5,
        'Como você acompanha o que entra e o que sai de dinheiro? Existe alguma rotina de controle?', 2)
    add_question(doc, 6,
        'Como você toma a decisão de quando pode gastar mais e quando precisa segurar? '
        'Isso é feito de forma estruturada ou no feeling?', 2)

    # ── BLOCO 3: FINANCEIRO ─────────────────────────────────────────────────
    section_header(doc, 'Bloco 3', 'Estrutura Financeira')
    intro_box(doc,
        'Esse é um dos blocos mais importantes. Quanto mais você compartilhar aqui, melhor '
        'conseguiremos estruturar o fluxo de caixa e os alertas automáticos.')
    add_question(doc, 7,
        'Quais são seus principais custos fixos mensais? '
        '(Exemplos: aluguel, folha, assinaturas, energia…)', 2)
    add_question(doc, 8,
        'Quais são seus principais custos variáveis? '
        '(Exemplos: insumos, frete, comissões, embalagens…)', 2)
    add_question(doc, 9,
        'Você trabalha com prazos de pagamento? '
        '(Exemplo: paga fornecedor em 30 dias.) Qual é o prazo médio?', 1)
    add_question(doc, 10,
        'Você trabalha com prazos de recebimento? '
        '(Exemplo: vende parcelado, recebe em 15 dias.) Qual é o prazo médio?', 1)
    add_question(doc, 11,
        'Hoje você sabe com antecedência quanto vai ter no caixa nos próximos 15 ou 30 dias? '
        'Isso é visível para você de alguma forma?', 1)

    # ── BLOCO 4: VENDAS / SHOPIFY ───────────────────────────────────────────
    section_header(doc, 'Bloco 4', 'Vendas e Canal Shopify')
    intro_box(doc,
        'Queremos entender como funciona o processo de venda para garantir que o sistema '
        'integre corretamente com a sua plataforma.')
    add_question(doc, 12,
        'Você já utiliza o Shopify? Desde quando? É o seu canal principal de vendas ou um dos canais?', 1)
    add_question(doc, 13,
        'Como funciona o processo de venda do início ao fim? '
        '(Desde o pedido até o cliente receber.)', 2)
    add_question(doc, 14,
        'Quando você considera que uma venda está "concluída"? '
        '(Exemplo: quando é pago, quando é entregue…)', 1)
    add_question(doc, 15,
        'As vendas são à vista, parceladas ou ambas? '
        'Quais formas de pagamento você aceita hoje?', 1)

    # ── BLOCO 5: OPERAÇÃO / PRODUÇÃO ────────────────────────────────────────
    section_header(doc, 'Bloco 5', 'Operação e Produção')
    intro_box(doc,
        'Para sistemas de controle de produção, precisamos entender o ciclo produtivo '
        '— do insumo ao produto pronto.')
    add_question(doc, 16,
        'Você produz seus próprios produtos ou apenas revende? '
        'Se produz, o que é produzido internamente?', 1)
    add_question(doc, 17,
        'Quanto tempo leva para produzir um lote ou um pedido? '
        'Existe variação dependendo do produto?', 1)
    add_question(doc, 18,
        'Você trabalha com estoque? Qual é a lógica atual — produz sob demanda, '
        'mantém estoque fixo ou misto?', 1)
    add_question(doc, 19,
        'Existe uma quantidade mínima de produção por lote? '
        '(Exemplo: produz sempre em múltiplos de 100 unidades.)', 1)
    add_question(doc, 20,
        'Você sabe quanto custa produzir cada lote ou unidade hoje? '
        'Esse cálculo é feito de alguma forma?', 1)

    # ── BLOCO 6: INSUMOS / COMPRAS ──────────────────────────────────────────
    section_header(doc, 'Bloco 6', 'Insumos e Gestão de Compras')
    intro_box(doc,
        'Essa parte nos ajuda a criar alertas automáticos para que você nunca fique sem '
        'insumo e sempre compre na hora certa.')
    add_question(doc, 21,
        'Quais são os principais insumos que você precisa comprar para sua operação funcionar?', 2)
    add_question(doc, 22,
        'Existe quantidade mínima de compra por fornecedor? '
        '(Exemplo: compra mínima de R$500 ou 50 kg.)', 1)
    add_question(doc, 23,
        'Como você decide quando é hora de comprar mais insumos? '
        'Existe alguma lógica ou é feito na observação?', 1)
    add_question(doc, 24,
        'Já aconteceu de você ficar sem insumo em um momento crítico? '
        'Com que frequência isso ocorre?', 1)

    # ── BLOCO 7: LOGÍSTICA ──────────────────────────────────────────────────
    section_header(doc, 'Bloco 7', 'Logística e Entrega')
    intro_box(doc,
        'Entender a logística é fundamental para integrar rastreamento de pedidos '
        'e automatizar a comunicação com o cliente.')
    add_question(doc, 25,
        'Como funciona a entrega dos seus produtos hoje? '
        '(Correios, transportadora, frota própria, retirada…)', 1)
    add_question(doc, 26,
        'Você utiliza algum sistema ou plataforma para gerenciar os envios hoje? Qual?', 1)
    add_question(doc, 27,
        'Como você acompanha o status dos pedidos enviados? '
        'Isso é feito manualmente ou de forma automatizada?', 1)

    # ── BLOCO 8: ALERTAS ────────────────────────────────────────────────────
    section_header(doc, 'Bloco 8', 'Alertas e Controles Desejados')
    intro_box(doc,
        'Um bom sistema te avisa antes do problema acontecer. '
        'Marque abaixo o que faz sentido para o seu negócio.')

    add_para(doc, 'Que tipo de alertas você gostaria de receber automaticamente?',
             size=11, color=RGBColor(0x55,0x55,0x55), font='Georgia', before=0, after=100)

    add_checkbox_grid(doc, [
        'Saldo baixo no caixa',
        'Contas a pagar próximas do vencimento',
        'Estoque de insumo abaixo do mínimo',
        'Necessidade de iniciar nova produção',
        'Pedido parado sem movimentação',
        'Meta de vendas não atingida',
        'Recebimento atrasado',
        'Outros (descreva abaixo)',
    ])

    add_question(doc, 28,
        'Você tem algum outro alerta ou controle que gostaria de receber e que não foi listado acima?', 2)
    add_question(doc, 29,
        'Hoje você tem algum tipo de controle ou alerta sobre sua operação? Como ele funciona?', 1)

    # ── BLOCO 9: SISTEMAS / INTEGRAÇÕES ─────────────────────────────────────
    section_header(doc, 'Bloco 9', 'Sistemas e Integrações Existentes')
    intro_box(doc,
        'Para conectar tudo corretamente, precisamos saber com quais sistemas '
        'você já trabalha ou pretende trabalhar.')
    add_question(doc, 30,
        'Você utiliza algum ERP hoje? Qual? '
        '(Exemplo: Bling, Tiny, SAP, Totvs, OMIE, outro…)', 1)
    add_question(doc, 31,
        'Além do Shopify e do ERP, existe alguma outra ferramenta que faz parte da sua operação hoje? '
        '(Exemplo: app de logística, gateway de pagamento, WhatsApp Business, plataforma de e-mail…)', 2)
    add_question(doc, 32,
        'Existe alguma ferramenta que você gostaria de eliminar com o novo sistema?', 1)

    # ── BLOCO 10: PRIORIDADE ────────────────────────────────────────────────
    section_header(doc, 'Bloco 10', 'O Que Mais Importa Para Você')

    add_highlight_box(doc,
        'A Pergunta Mais Importante',
        'Se você pudesse resolver apenas um problema hoje com esse sistema, qual seria?',
        'Não precisa ser técnico. Escreva como você falaria para um sócio ou amigo.')

    doc.add_paragraph()  # espaçamento

    add_question(doc, 33,
        'Quais são os 3 maiores problemas ou travamentos que você enfrenta na gestão hoje?', 2)
    add_question(doc, 34,
        'Como seria sua rotina ideal com um sistema funcionando perfeitamente? '
        'O que mudaria no seu dia?', 2)

    # ── ENCERRAMENTO ────────────────────────────────────────────────────────
    section_header(doc, 'Finalização', 'Muito Obrigado')

    add_closing_box(doc, [
        'Suas respostas são o insumo mais valioso que temos para construir um sistema que realmente '
        'funcione para o seu negócio. Não existe resposta certa ou errada aqui — o que importa é '
        'entender a realidade da sua operação.',

        'Com base neste levantamento, vamos estruturar a proposta técnica, mapear as automações '
        'mais relevantes e apresentar um plano de implementação com prioridades claras.',

        'Qualquer dúvida ou ponto que prefira explicar em uma conversa, é só nos chamar. '
        'Podemos agendar uma sessão de alinhamento para complementar este documento.',
    ])

    # ── RODAPÉ ──────────────────────────────────────────────────────────────
    build_footer(doc)

    # ── SALVAR ──────────────────────────────────────────────────────────────
    output = r'c:\Users\digui\Documents\localrise-brain\Claude\Claude Local RIse\Saas-Factory\LEVANTAMENTO-ESTRATEGICO-SISTEMA-GESTAO.docx'
    doc.save(output)
    print(f'Documento salvo: {output}')


if __name__ == '__main__':
    build_document()
